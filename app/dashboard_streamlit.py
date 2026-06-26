# HOTFIX_15F_10_34_ANALISE_INDIVIDUAL_VM_SEM_ERROS
# HOTFIX_15F_10_33_TOTAL_IA_GRAFICOS_PUSH
# HOTFIX_15F_10_32_MARKDOWN_RENDER_WORD_PDF_HTML
# HOTFIX_15F_10_31_RELATORIO_EXECUTIVO_PDF_WORD
# HOTFIX_15F_10_30_IA_RELATORIO_COMERCIAL
# HOTFIX_15F_10_29_ANALISE_COMPLETA_V1
# HOTFIX_15F_10_28_ANALISE_SOLICITACAO_INDEX
# HOTFIX_15F_10_25_LLM_STARTER_CONTINUO
# HOTFIX_15F_10_24_IMPORTS_LLM_FUP
# HOTFIX_15F_10_23_LLM_THRESHOLD_REPORTS
# MARCO_15F_10_22_FUP_CLICK_VM - estabiliza FUP; clique na VM abre histórico e previsão
# HOTFIX_15F_10_19_NAV_AMBIENTES_SEPARADOS
# HOTFIX_15F_10_12_VSPHERE_DATASTORE_AUTHORITATIVE
# HOTFIX_15F_10_11_SCHEMA_REAL_ALOCACOES
# HOTFIX_15F_10_9_CARDS_FORMATACAO_SAFE: merge de cards parciais, infinitos e width stretch.
# HOTFIX_15F_10_8_FORMATACAO_2_CASAS: tabelas com 2 casas decimais.
# HOTFIX_15F_10_2_COLETA_ALOCACOES_DATASTORES: coleta vCPU/pCPU/over ratios/datastores.
# HOTFIX_15F_10_1_VCPU_OVERRATIO_CLUSTER: vCPU alocada e over ratio CPU no cluster/host.
# HOTFIX_15F_10_COLETA_MANUAL_DASHBOARD: botão de coleta, prazo histórico e prazo forecast.
# HOTFIX_15F_9_SELETOR_COLETA_RESOURCE_RUN: adiciona seletor de coleta vROps/run_id para histórico.
# HOTFIX_15F_8_ALOCACOES_CARDS_ND: integra alocações para remover N/D dos cards.
# HOTFIX_15F_7_1_FORECAST_PARTICOES_FUP: adiciona forecast 30/60/90 por partição de disco para FUP.
# HOTFIX_15F_7_PARTICOES_DISCO_HISTORICO: adiciona gráfico histórico por partição de disco.
# HOTFIX_15F_6_7_DASHBOARD_HISTORICO_REAL_FALLBACK_VM: helper DuckDB tenta run atual, sem run e forecast-only.
# HOTFIX_15F_6_4_HISTORICO_REAL_DUCKDB_IDEMPOTENTE: gráfico de histórico usa vm_resource_timeseries diretamente no DuckDB.
# HOTFIX_15F_6_3_HISTORICO_REAL_DUCKDB_POR_ANCORA: PASSO 1 usa vm_resource_timeseries diretamente no gráfico.

from pathlib import Path
import re
import sys
import os
import math
import subprocess
import json
import urllib.request
import urllib.error
import io
import zipfile
import html

import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st

try:
    import duckdb
except Exception:
    duckdb = None


# =============================================================================
# Path do projeto
# =============================================================================
PROJECT_ROOT = Path(__file__).resolve().parents[1]

if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from rmc_copilot.dashboard.historico_real_duckdb import criar_grafico_historico_forecast_escopo_duckdb
from rmc_copilot.dashboard.particoes_duckdb import criar_grafico_historico_particoes_duckdb
from rmc_copilot.dashboard.alocacoes_duckdb import calcular_alocacoes_cards, carregar_datastores_escopo
from rmc_copilot.dashboard.coletas_duckdb import listar_resource_collection_runs_dashboard



# =============================================================================
# Imports internos
# =============================================================================
from rmc_copilot.etl import carregar_aba
from rmc_copilot.parser_rmc import consolidar_metricas_vm
from rmc_copilot.analyzer import analisar_capacity_vms
from rmc_copilot.forecasting import consolidar_forecasts, adicionar_risco_futuro
from rmc_copilot.prioritization import aplicar_priorizacao_v04
from rmc_copilot.database import (
    salvar_execucao,
    listar_execucoes,
    carregar_execucoes_com_label,
    carregar_analise_por_execucao,
    carregar_resumo_cluster_por_execucao,
    obter_ultima_execucao,
    comparar_execucoes_vms,
    resumir_comparacao_execucoes,
    comparar_resumo_cluster,
)


# =============================================================================
# Configurações
# =============================================================================
CAMINHO_BANCO = PROJECT_ROOT / "data/database/rmc_copilot.duckdb"

ARQUIVO_PADRAO = (
    PROJECT_ROOT
    / "data/raw/rmc_outputs/RMC_Recursos_VM_v5_10_2_20260610_094958.xlsx"
)

st.set_page_config(
    page_title="BV | RMC Copilot",
    page_icon="📊",
    layout="wide",
)


# =============================================================================
# Estilo visual BV
# =============================================================================
def aplicar_estilo_bv():
    st.markdown(
        """
        <style>
        .stApp {
            background-color: #F5F8FC;
        }

        .block-container {
            padding-top: 1.3rem;
            padding-bottom: 2rem;
        }

        .bv-header {
            background: linear-gradient(90deg, #1F3B73 0%, #2E5B9A 55%, #46B5E5 100%);
            padding: 20px 26px;
            border-radius: 18px;
            margin-bottom: 20px;
            color: white;
            box-shadow: 0 4px 16px rgba(31, 59, 115, 0.18);
        }

        .bv-title {
            font-size: 32px;
            font-weight: 800;
            margin: 0;
            letter-spacing: 0.2px;
        }

        .bv-subtitle {
            font-size: 14px;
            opacity: 0.95;
            margin-top: 5px;
        }

        .section-title {
            font-size: 22px;
            font-weight: 800;
            color: #1F3B73;
            margin-top: 26px;
            margin-bottom: 12px;
        }

        .section-subtitle {
            font-size: 14px;
            color: #5B6472;
            margin-top: -6px;
            margin-bottom: 16px;
        }

        .bv-card {
            background: #FFFFFF;
            border-radius: 18px;
            padding: 18px 18px 14px 18px;
            box-shadow: 0 4px 14px rgba(31, 59, 115, 0.08);
            border: 1px solid #E3ECF7;
            margin-bottom: 16px;
        }

        .bv-card-title {
            color: #1F3B73;
            font-weight: 800;
            font-size: 16px;
            margin-bottom: 8px;
        }

        .kpi-card {
            border-radius: 18px;
            padding: 14px 14px 12px 14px;
            color: white;
            box-shadow: 0 6px 16px rgba(31, 59, 115, 0.16);
            min-height: 122px;
            display: flex;
            flex-direction: column;
            justify-content: space-between;
            border: 1px solid rgba(255, 255, 255, 0.14);
        }

        .kpi-blue {
            background: linear-gradient(135deg, #1F3B73, #2E5B9A);
        }

        .kpi-cyan {
            background: linear-gradient(135deg, #2E86C1, #46B5E5);
        }

        .kpi-green {
            background: linear-gradient(135deg, #4FAE5A, #78C27B);
        }

        .kpi-yellow {
            background: linear-gradient(135deg, #D9A441, #E5B95C);
        }

        .kpi-orange {
            background: linear-gradient(135deg, #E67E22, #F39C12);
        }

        .kpi-red {
            background: linear-gradient(135deg, #D9534F, #E57373);
        }

        .kpi-label {
            font-size: 13px;
            font-weight: 800;
            opacity: 0.96;
            margin-bottom: 8px;
            letter-spacing: 0.1px;
        }

        .kpi-value {
            font-size: 20px;
            font-weight: 900;
            line-height: 1.08;
            word-break: break-word;
        }

        .kpi-foot {
            font-size: 11px;
            margin-top: 8px;
            opacity: 0.92;
            text-transform: uppercase;
            letter-spacing: 0.4px;
        }

        .summary-box {
            background: #FFFFFF;
            border-radius: 18px;
            padding: 18px 22px;
            border-left: 6px solid #2E5B9A;
            box-shadow: 0 4px 14px rgba(31, 59, 115, 0.08);
            color: #1F2937;
            margin-bottom: 20px;
        }

        section[data-testid="stSidebar"] {
            background: linear-gradient(180deg, #DDF0FB 0%, #C7E8F9 100%);
            border-right: 1px solid #D8E6F2;
        }

        section[data-testid="stSidebar"] h1,
        section[data-testid="stSidebar"] h2,
        section[data-testid="stSidebar"] h3,
        section[data-testid="stSidebar"] label,
        section[data-testid="stSidebar"] p,
        section[data-testid="stSidebar"] span {
            color: #1F3B73 !important;
        }

        div[data-testid="stDataFrame"] {
            background: #FFFFFF;
            border-radius: 16px;
            border: 1px solid #E3ECF7;
            padding: 8px;
            box-shadow: 0 3px 12px rgba(31, 59, 115, 0.06);
        }

        .stButton > button,
        .stDownloadButton > button {
            background: #1F3B73;
            color: white;
            border: none;
            border-radius: 11px;
            padding: 0.45rem 1rem;
            font-weight: 700;
        }

        .stButton > button:hover,
        .stDownloadButton > button:hover {
            background: #2E5B9A;
            color: white;
        }

        div[data-baseweb="notification"] {
            border-radius: 12px;
        }

        .stSelectbox,
        .stMultiSelect,
        .stTextInput,
        .stFileUploader {
            border-radius: 12px;
        }

        hr {
            border: none;
            height: 1px;
            background: #B9D9EC;
            margin: 16px 0;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def render_header():
    st.markdown(
        """
        <div class="bv-header">
            <div class="bv-title">BV | Capacity Dashboard</div>
            <div class="bv-subtitle">
                Capacity Planning VMware • Priorização Operacional • Forecast 30/60/90 dias • Histórico DuckDB
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_kpi_card(label, value, css_class="kpi-blue", foot=""):
    st.markdown(
        f"""
        <div class="kpi-card {css_class}">
            <div class="kpi-label">{label}</div>
            <div class="kpi-value">{value}</div>
            <div class="kpi-foot">{foot}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def formatar_numero(valor):
    try:
        return f"{int(valor):,}".replace(",", ".")
    except Exception:
        return str(valor)


# =============================================================================
# Nomes amigáveis para métricas
# =============================================================================
NOME_PICO_USO = "Pico de uso típico"
NOME_FORECAST = "Previsão de uso"

COLUNAS_VISUAIS = {
    "cpu_p95_pct": "CPU — pico típico de uso (%)",
    "mem_p95_pct": "Memória — pico típico de uso (%)",
    "disk_p95_pct": "Disco — pico típico de uso (%)",
    "cpu_p95_medio": "CPU — pico médio por cluster (%)",
    "mem_p95_medio": "Memória — pico médio por cluster (%)",
    "disk_p95_medio": "Disco — pico médio por cluster (%)",
    "cpu_forecast_30d": "CPU — previsão 30 dias (%)",
    "cpu_forecast_60d": "CPU — previsão 60 dias (%)",
    "cpu_forecast_90d": "CPU — previsão 90 dias (%)",
    "mem_forecast_30d": "Memória — previsão 30 dias (%)",
    "mem_forecast_60d": "Memória — previsão 60 dias (%)",
    "mem_forecast_90d": "Memória — previsão 90 dias (%)",
    "disk_forecast_30d": "Disco — previsão 30 dias (%)",
    "disk_forecast_60d": "Disco — previsão 60 dias (%)",
    "disk_forecast_90d": "Disco — previsão 90 dias (%)",
    "recurso": "Recurso",
    "valor_atual": "Uso atual — pico típico (%)",
    "forecast_30d": "Previsão 30 dias (%)",
    "forecast_60d": "Previsão 60 dias (%)",
    "forecast_90d": "Previsão 90 dias (%)",
    "delta_90d": "Variação até 90 dias (p.p.)",
    "vms_acima_85_90d": "VMs previstas acima de 85% em 90 dias",
    "vms_acima_95_90d": "VMs previstas acima de 95% em 90 dias",
}


COLUNAS_VISUAIS.update(
    {
        "partition_name": "Partição",
        "last_used_pct": "Último uso (%)",
        "max_used_pct": "Maior uso no período (%)",
        "avg_used_pct": "Uso médio no período (%)",
        "first_date": "Primeira data",
        "last_date": "Última data",
        "pontos": "Pontos coletados",
        "series_label": "Série",
    }
)


COLUNAS_VISUAIS.update(
    {
        "forecast_30d": "Previsão 30 dias (%)",
        "forecast_60d": "Previsão 60 dias (%)",
        "forecast_90d": "Previsão 90 dias (%)",
        "delta_90d": "Variação 90 dias (p.p.)",
        "trend_pp_day": "Tendência (p.p./dia)",
        "dias_ate_85": "Dias até 85%",
        "dias_ate_95": "Dias até 95%",
        "status_fup": "Status FUP",
        "acao_fup": "Ação FUP",
        "evidencia_fup": "Evidência para FUP",
    }
)

COLUNAS_VISUAIS.update(
    {
        "fonte_alocacao": "Fonte da alocação",
        "campos_faltantes": "Campos de alocação faltantes",
        "alocacao_status": "Status da alocação",
    }
)


COLUNAS_VISUAIS.update(
    {
        "pcpus": "pCPU / cores físicos",
        "cpu_over_ratio": "CPU over ratio",
        "host_memoria_gb": "Memória física host/cluster (GB)",
        "memoria_over_ratio": "Memória over ratio",
        "host_count": "Hosts",
    }
)


COLUNAS_VISUAIS.update(
    {
        "datastore_total_gb": "Datastore — capacidade total (GB)",
        "datastore_used_gb": "Datastore — usado (GB)",
        "datastore_free_gb": "Datastore — livre (GB)",
        "datastore_used_pct": "Datastore — ocupação (%)",
        "datastore_count": "Datastores",
        "disk_over_ratio": "Disco over ratio",
        "memoria_over_ratio": "Memória over ratio",
        "cpu_over_ratio": "CPU over ratio",
        "host_memoria_gb": "Memória física host/cluster (GB)",
    }
)



COLUNAS_VISUAIS.update(
    {
        "cluster": "Cluster",
        "host": "Host",
        "datastore": "Datastore",
        "motivo_fup": "Motivo FUP",
        "partition_name": "Partição",
    }
)


def formatar_dataframe_visual_2_casas(df: pd.DataFrame) -> pd.DataFrame:
    """
    HOTFIX 15F.10.9.
    Formata colunas numéricas para no máximo 2 casas, padrão pt-BR.
    Protege contra NaN, +inf e -inf para não quebrar o Streamlit.
    """
    if df is None or df.empty:
        return df

    out = df.copy()

    for col in out.columns:
        if pd.api.types.is_bool_dtype(out[col]):
            continue
        if pd.api.types.is_numeric_dtype(out[col]):
            s = pd.to_numeric(out[col], errors="coerce")

            def _fmt_cell(v):
                if pd.isna(v):
                    return ""
                try:
                    n = float(v)
                    if not math.isfinite(n):
                        return ""
                    if abs(n - round(n)) < 0.005:
                        return f"{int(round(n)):,}".replace(",", ".")
                    return f"{n:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                except Exception:
                    return ""

            out[col] = s.apply(_fmt_cell)

    return out


COLUNAS_VISUAIS.update(
    {
        "cluster": "Cluster",
        "total_vms": "VMs",
        "p0_acao_imediata": "P0 — ação imediata",
        "p1_alta": "P1 — alta prioridade",
        "vms_prioritarias": "VMs prioritárias",
        "pct_vms_prioritarias": "% VMs prioritárias",
        "score_max": "Score máximo",
        "CPU — pico médio por cluster (%)": "CPU pico médio (%)",
        "Memória — pico médio por cluster (%)": "Memória pico médio (%)",
        "Disco — pico médio por cluster (%)": "Disco pico médio (%)",
        "vcpu_count": "vCPU alocada",
        "memory_gb": "Memória alocada (GB)",
        "provisioned_gb": "Disco provisionado (GB)",
        "physical_cores": "pCPU / cores físicos",
        "memory_gb_total": "Memória física (GB)",
        "capacity_gb": "Capacidade datastore (GB)",
        "used_gb": "Usado datastore (GB)",
        "free_gb": "Livre datastore (GB)",
    }
)


def renomear_colunas_visual(df: pd.DataFrame) -> pd.DataFrame:
    return df.rename(columns={c: COLUNAS_VISUAIS.get(c, c) for c in df.columns})


# =============================================================================
# Pipeline
# =============================================================================
@st.cache_data(show_spinner=True)
def processar_excel(caminho_arquivo: str) -> pd.DataFrame:
    arquivo = Path(caminho_arquivo)

    df_vms = carregar_aba(arquivo, "VMS_SELECIONADAS")
    df_cpu = carregar_aba(arquivo, "HIST_CPU")
    df_mem = carregar_aba(arquivo, "HIST_MEM")
    df_disk = carregar_aba(arquivo, "HIST_DISK")

    df_consolidado = consolidar_metricas_vm(
        df_vms=df_vms,
        df_cpu=df_cpu,
        df_mem=df_mem,
        df_disk=df_disk,
    )

    df_v2 = analisar_capacity_vms(df_consolidado)

    df_v3 = consolidar_forecasts(
        df_base=df_v2,
        df_cpu=df_cpu,
        df_mem=df_mem,
        df_disk=df_disk,
    )

    df_v3 = adicionar_risco_futuro(df_v3)

    df_v4 = aplicar_priorizacao_v04(df_v3)

    return df_v4


@st.cache_data(show_spinner=True)
def carregar_execucao_banco(caminho_banco: str, execution_id: str):
    df = carregar_analise_por_execucao(caminho_banco, execution_id)
    resumo = carregar_resumo_cluster_por_execucao(caminho_banco, execution_id)

    return df, resumo


# =============================================================================
# Resumos
# =============================================================================
def criar_resumo_cluster(df: pd.DataFrame) -> pd.DataFrame:
    colunas_resumo = [
        "cluster",
        "total_vms",
        "p0_acao_imediata",
        "p1_alta",
        "p2_media",
        "p3_baixa",
        "p4_monitorar",
        "criticas",
        "risco_atual",
        "atencao",
        "otimizacao",
        "ok",
        "risco_futuro_30d",
        "risco_futuro_60d",
        "risco_futuro_90d",
        "score_medio",
        "score_max",
        "cpu_p95_medio",
        "mem_p95_medio",
        "disk_p95_medio",
        "cpu_forecast_90d_medio",
        "mem_forecast_90d_medio",
        "disk_forecast_90d_medio",
        "disk_forecast_90d_max",
        "vms_prioritarias",
        "pct_vms_prioritarias",
    ]

    if df.empty:
        return pd.DataFrame(columns=colunas_resumo)

    resumo = (
        df.groupby("cluster", dropna=False)
        .agg(
            total_vms=("vm", "count"),
            p0_acao_imediata=("prioridade_final", lambda x: (x == "P0_ACAO_IMEDIATA").sum()),
            p1_alta=("prioridade_final", lambda x: (x == "P1_ALTA").sum()),
            p2_media=("prioridade_final", lambda x: (x == "P2_MEDIA").sum()),
            p3_baixa=("prioridade_final", lambda x: (x == "P3_BAIXA").sum()),
            p4_monitorar=("prioridade_final", lambda x: (x == "P4_MONITORAR").sum()),
            criticas=("status_geral", lambda x: (x == "CRITICO").sum()),
            risco_atual=("status_geral", lambda x: (x == "RISCO").sum()),
            atencao=("status_geral", lambda x: (x == "ATENCAO").sum()),
            otimizacao=("status_geral", lambda x: (x == "OTIMIZACAO").sum()),
            ok=("status_geral", lambda x: (x == "OK").sum()),
            risco_futuro_30d=("criticidade_futura", lambda x: (x == "RISCO_FUTURO_30D").sum()),
            risco_futuro_60d=("criticidade_futura", lambda x: (x == "RISCO_FUTURO_60D").sum()),
            risco_futuro_90d=("criticidade_futura", lambda x: (x == "RISCO_FUTURO_90D").sum()),
            score_medio=("score_prioridade", "mean"),
            score_max=("score_prioridade", "max"),
            cpu_p95_medio=("cpu_p95_pct", "mean"),
            mem_p95_medio=("mem_p95_pct", "mean"),
            disk_p95_medio=("disk_p95_pct", "mean"),
            cpu_forecast_90d_medio=("cpu_forecast_90d", "mean"),
            mem_forecast_90d_medio=("mem_forecast_90d", "mean"),
            disk_forecast_90d_medio=("disk_forecast_90d", "mean"),
            disk_forecast_90d_max=("disk_forecast_90d", "max"),
        )
        .reset_index()
    )

    colunas_numericas = [c for c in resumo.columns if c != "cluster"]

    for coluna in colunas_numericas:
        resumo[coluna] = pd.to_numeric(resumo[coluna], errors="coerce").fillna(0)

    resumo["vms_prioritarias"] = (
        resumo["p0_acao_imediata"] + resumo["p1_alta"]
    )

    resumo["pct_vms_prioritarias"] = 0.0
    mask = resumo["total_vms"] > 0

    resumo.loc[mask, "pct_vms_prioritarias"] = (
        resumo.loc[mask, "vms_prioritarias"]
        / resumo.loc[mask, "total_vms"]
        * 100
    )

    resumo = resumo.sort_values(
        ["pct_vms_prioritarias", "p0_acao_imediata", "p1_alta", "score_max"],
        ascending=False,
    )

    return resumo


def gerar_resumo_textual(df: pd.DataFrame, resumo_cluster: pd.DataFrame) -> str:
    total = len(df)

    prioridade_counts = df["prioridade_final"].value_counts()
    status_counts = df["status_geral"].value_counts()
    risco_counts = df["criticidade_futura"].value_counts()

    p0 = int(prioridade_counts.get("P0_ACAO_IMEDIATA", 0))
    p1 = int(prioridade_counts.get("P1_ALTA", 0))

    critico = int(status_counts.get("CRITICO", 0))
    risco = int(status_counts.get("RISCO", 0))
    atencao = int(status_counts.get("ATENCAO", 0))
    otimizacao = int(status_counts.get("OTIMIZACAO", 0))
    ok = int(status_counts.get("OK", 0))
    sem_dados = int(status_counts.get("SEM_DADOS", 0))

    risco_30 = int(risco_counts.get("RISCO_FUTURO_30D", 0))
    risco_60 = int(risco_counts.get("RISCO_FUTURO_60D", 0))
    risco_90 = int(risco_counts.get("RISCO_FUTURO_90D", 0))

    linhas = []
    linhas.append(
        f"Foram analisadas **{formatar_numero(total)} VMs**. "
        f"A fila prioritária contém **{formatar_numero(p0 + p1)} VMs P0/P1**, "
        f"sendo **{formatar_numero(p0)} P0** e **{formatar_numero(p1)} P1**."
    )

    linhas.append(
        f"No status atual, há **{formatar_numero(critico)} críticas**, "
        f"**{formatar_numero(risco)} em risco**, "
        f"**{formatar_numero(atencao)} em atenção**, "
        f"**{formatar_numero(otimizacao)} candidatas a otimização**, "
        f"**{formatar_numero(ok)} OK** e **{formatar_numero(sem_dados)} sem dados**."
    )

    linhas.append(
        f"O forecast indica **{formatar_numero(risco_30)} VMs com risco em 30 dias**, "
        f"**{formatar_numero(risco_60)} em 60 dias** e "
        f"**{formatar_numero(risco_90)} em 90 dias**."
    )

    if not resumo_cluster.empty:
        linhas.append("")
        linhas.append("**Clusters mais relevantes:**")

        for _, row in resumo_cluster.head(3).iterrows():
            linhas.append(
                f"- **{row['cluster']}**: "
                f"{int(row['p0_acao_imediata'])} P0, "
                f"{int(row['p1_alta'])} P1, "
                f"{row['pct_vms_prioritarias']:.1f}% de VMs prioritárias."
            )

    return "\n".join(linhas)


# =============================================================================
# Filtros
# =============================================================================
def filtrar_dataframe(df: pd.DataFrame):
    st.sidebar.markdown("---")
    st.sidebar.markdown("### Filtros")

    recurso_selecionado = st.sidebar.radio(
        "Recurso principal",
        options=["Todos", "CPU", "Memória", "Disco"],
        index=0,
    )

    clusters = sorted(df["cluster"].dropna().unique().tolist())
    categorias = sorted(df["categoria_vm"].dropna().unique().tolist())
    prioridades = sorted(df["prioridade_final"].dropna().unique().tolist())
    acoes = sorted(df["acao_final"].dropna().unique().tolist())
    riscos = sorted(df["criticidade_futura"].dropna().unique().tolist())

    filtro_cluster = st.sidebar.multiselect("Cluster", options=clusters, default=[])
    filtro_categoria = st.sidebar.multiselect("Categoria VM", options=categorias, default=[])
    filtro_prioridade = st.sidebar.multiselect("Prioridade", options=prioridades, default=[])
    filtro_acao = st.sidebar.multiselect("Ação final", options=acoes, default=[])
    filtro_risco = st.sidebar.multiselect("Risco futuro", options=riscos, default=[])

    busca_vm = st.sidebar.text_input("Buscar VM", "")

    if not filtro_cluster:
        filtro_cluster = clusters

    if not filtro_categoria:
        filtro_categoria = categorias

    if not filtro_prioridade:
        filtro_prioridade = prioridades

    if not filtro_acao:
        filtro_acao = acoes

    if not filtro_risco:
        filtro_risco = riscos

    df_filtrado = df[
        df["cluster"].isin(filtro_cluster)
        & df["categoria_vm"].isin(filtro_categoria)
        & df["prioridade_final"].isin(filtro_prioridade)
        & df["acao_final"].isin(filtro_acao)
        & df["criticidade_futura"].isin(filtro_risco)
    ].copy()

    if busca_vm.strip():
        termo = busca_vm.strip().upper()
        df_filtrado = df_filtrado[
            df_filtrado["vm"].astype(str).str.upper().str.contains(termo, na=False)
        ]

    return df_filtrado, recurso_selecionado


# =============================================================================
# Gráficos
# =============================================================================
CORES_PRIORIDADE = {
    "P0_ACAO_IMEDIATA": "#D9534F",
    "P1_ALTA": "#F08A24",
    "P2_MEDIA": "#D9A441",
    "P3_BAIXA": "#46B5E5",
    "P4_MONITORAR": "#4FAE5A",
}

CORES_RISCO = {
    "RISCO_FUTURO_30D": "#D9534F",
    "RISCO_FUTURO_60D": "#F08A24",
    "RISCO_FUTURO_90D": "#D9A441",
    "SEM_RISCO_FUTURO": "#4FAE5A",
}

CORES_RECURSO = {
    "CPU": "#2E5B9A",
    "Memória": "#F08A24",
    "Disco": "#D9534F",
}


def aplicar_layout_plotly(fig, altura=420, showlegend=True):
    fig.update_layout(
        height=altura,
        paper_bgcolor="white",
        plot_bgcolor="white",
        title_font=dict(size=18, color="#1F3B73"),
        font=dict(color="#374151"),
        margin=dict(l=20, r=20, t=60, b=40),
        showlegend=showlegend,
    )
    fig.update_xaxes(showgrid=False)
    fig.update_yaxes(gridcolor="#E5EDF7")
    return fig


def criar_grafico_prioridade(df: pd.DataFrame):
    dados = df["prioridade_final"].value_counts().reset_index()
    dados.columns = ["prioridade_final", "quantidade"]

    ordem = [
        "P0_ACAO_IMEDIATA",
        "P1_ALTA",
        "P2_MEDIA",
        "P3_BAIXA",
        "P4_MONITORAR",
    ]

    dados["prioridade_final"] = pd.Categorical(
        dados["prioridade_final"],
        categories=ordem,
        ordered=True,
    )
    dados = dados.sort_values("prioridade_final")

    fig = px.bar(
        dados,
        x="prioridade_final",
        y="quantidade",
        title="Distribuição por prioridade final",
        text="quantidade",
        color="prioridade_final",
        color_discrete_map=CORES_PRIORIDADE,
    )

    fig.update_layout(
        xaxis_title="Prioridade",
        yaxis_title="Quantidade de VMs",
    )

    return aplicar_layout_plotly(fig, altura=420, showlegend=False)


def criar_grafico_risco_futuro(df: pd.DataFrame):
    dados = df["criticidade_futura"].value_counts().reset_index()
    dados.columns = ["criticidade_futura", "quantidade"]

    ordem = [
        "RISCO_FUTURO_30D",
        "RISCO_FUTURO_60D",
        "RISCO_FUTURO_90D",
        "SEM_RISCO_FUTURO",
    ]

    dados["criticidade_futura"] = pd.Categorical(
        dados["criticidade_futura"],
        categories=ordem,
        ordered=True,
    )
    dados = dados.sort_values("criticidade_futura")

    fig = px.bar(
        dados,
        x="criticidade_futura",
        y="quantidade",
        title="Risco futuro 30/60/90 dias",
        text="quantidade",
        color="criticidade_futura",
        color_discrete_map=CORES_RISCO,
    )

    fig.update_layout(
        xaxis_title="Risco futuro",
        yaxis_title="Quantidade de VMs",
    )

    return aplicar_layout_plotly(fig, altura=420, showlegend=False)


def criar_grafico_acao(df: pd.DataFrame):
    dados = df["acao_final"].value_counts().reset_index()
    dados.columns = ["acao_final", "quantidade"]

    fig = px.bar(
        dados,
        x="quantidade",
        y="acao_final",
        orientation="h",
        title="Distribuição por ação final",
        text="quantidade",
        color_discrete_sequence=["#2E5B9A"],
    )

    fig.update_layout(
        xaxis_title="Quantidade de VMs",
        yaxis_title="Ação final",
        yaxis={"categoryorder": "total ascending"},
    )

    return aplicar_layout_plotly(fig, altura=560, showlegend=False)


def criar_grafico_top_clusters(resumo_cluster: pd.DataFrame):
    dados = resumo_cluster.head(10).copy()

    fig = px.bar(
        dados,
        x="vms_prioritarias",
        y="cluster",
        orientation="h",
        title="Top clusters por VMs P0/P1",
        text="vms_prioritarias",
        color="pct_vms_prioritarias",
        color_continuous_scale=["#46B5E5", "#2E5B9A", "#1F3B73"],
        hover_data=[
            "p0_acao_imediata",
            "p1_alta",
            "pct_vms_prioritarias",
            "score_max",
        ],
    )

    fig.update_layout(
        xaxis_title="VMs prioritárias P0/P1",
        yaxis_title="Cluster",
        yaxis={"categoryorder": "total ascending"},
    )

    return aplicar_layout_plotly(fig, altura=560, showlegend=False)


def criar_grafico_recursos_cluster(resumo_cluster: pd.DataFrame, recurso_selecionado: str):
    dados = resumo_cluster.head(10).copy()

    if dados.empty:
        return None

    if recurso_selecionado == "Todos":
        colunas = {
            "cpu_p95_medio": "CPU",
            "mem_p95_medio": "Memória",
            "disk_p95_medio": "Disco",
        }

        colunas_existentes = [c for c in colunas.keys() if c in dados.columns]

        if not colunas_existentes:
            return None

        dados_long = dados[["cluster"] + colunas_existentes].melt(
            id_vars="cluster",
            value_vars=colunas_existentes,
            var_name="recurso",
            value_name="p95_medio",
        )

        dados_long["recurso"] = dados_long["recurso"].map(colunas)

        fig = px.bar(
            dados_long,
            x="cluster",
            y="p95_medio",
            color="recurso",
            barmode="group",
            title="Pico de uso típico por cluster — CPU, Memória e Disco",
            color_discrete_map=CORES_RECURSO,
        )

        fig.update_layout(
            xaxis_title="Cluster",
            yaxis_title="Pico de uso típico (%)",
        )

        return aplicar_layout_plotly(fig, altura=520, showlegend=True)

    mapa = {
        "CPU": ("cpu_p95_medio", "CPU — pico de uso típico médio por cluster", "#2E5B9A"),
        "Memória": ("mem_p95_medio", "Memória — pico de uso típico médio por cluster", "#F08A24"),
        "Disco": ("disk_p95_medio", "Disco — pico de uso típico médio por cluster", "#D9534F"),
    }

    if recurso_selecionado not in mapa:
        return None

    coluna, titulo, cor = mapa[recurso_selecionado]

    if coluna not in dados.columns:
        return None

    dados = dados.sort_values(coluna, ascending=False)

    fig = px.bar(
        dados,
        x=coluna,
        y="cluster",
        orientation="h",
        title=titulo,
        text=coluna,
        color_discrete_sequence=[cor],
        hover_data=["total_vms", "vms_prioritarias", "score_max"],
    )

    fig.update_traces(texttemplate="%{text:.1f}%", textposition="outside")

    fig.update_layout(
        xaxis_title="Pico de uso típico (%)",
        yaxis_title="Cluster",
        yaxis={"categoryorder": "total ascending"},
    )

    return aplicar_layout_plotly(fig, altura=560, showlegend=False)


def criar_grafico_top_vms(df: pd.DataFrame, recurso_selecionado: str = "Todos"):
    if recurso_selecionado == "CPU" and "cpu_p95_pct" in df.columns:
        coluna = "cpu_p95_pct"
        titulo = "Top 30 VMs por CPU — pico de uso típico"
        cor = "#2E5B9A"
        dados = df.sort_values(coluna, ascending=False).head(30).copy()

        fig = px.bar(
            dados,
            x=coluna,
            y="vm",
            orientation="h",
            title=titulo,
            text=coluna,
            color_discrete_sequence=[cor],
            hover_data=["cluster", "categoria_vm", "status_cpu", "cpu_forecast_90d", "acao_final"],
        )

        fig.update_traces(texttemplate="%{text:.1f}%", textposition="outside")
        fig.update_layout(xaxis_title="CPU — pico típico de uso (%)", yaxis_title="VM", yaxis={"categoryorder": "total ascending"})
        return aplicar_layout_plotly(fig, altura=760, showlegend=False)

    if recurso_selecionado == "Memória" and "mem_p95_pct" in df.columns:
        coluna = "mem_p95_pct"
        titulo = "Top 30 VMs por Memória — pico de uso típico"
        cor = "#F08A24"
        dados = df.sort_values(coluna, ascending=False).head(30).copy()

        fig = px.bar(
            dados,
            x=coluna,
            y="vm",
            orientation="h",
            title=titulo,
            text=coluna,
            color_discrete_sequence=[cor],
            hover_data=["cluster", "categoria_vm", "status_memoria", "mem_forecast_90d", "acao_final"],
        )

        fig.update_traces(texttemplate="%{text:.1f}%", textposition="outside")
        fig.update_layout(xaxis_title="Memória — pico típico de uso (%)", yaxis_title="VM", yaxis={"categoryorder": "total ascending"})
        return aplicar_layout_plotly(fig, altura=760, showlegend=False)

    if recurso_selecionado == "Disco" and "disk_p95_pct" in df.columns:
        coluna = "disk_p95_pct"
        titulo = "Top 30 VMs por Disco — pico de uso típico"
        cor = "#D9534F"
        dados = df.sort_values(coluna, ascending=False).head(30).copy()

        fig = px.bar(
            dados,
            x=coluna,
            y="vm",
            orientation="h",
            title=titulo,
            text=coluna,
            color_discrete_sequence=[cor],
            hover_data=["cluster", "categoria_vm", "status_disco", "disk_forecast_90d", "acao_final"],
        )

        fig.update_traces(texttemplate="%{text:.1f}%", textposition="outside")
        fig.update_layout(xaxis_title="Disco — pico típico de uso (%)", yaxis_title="VM", yaxis={"categoryorder": "total ascending"})
        return aplicar_layout_plotly(fig, altura=760, showlegend=False)

    dados = df.sort_values("score_prioridade", ascending=False).head(30).copy()

    fig = px.bar(
        dados,
        x="score_prioridade",
        y="vm",
        orientation="h",
        color="prioridade_final",
        title="Top 30 VMs por score de prioridade",
        color_discrete_map=CORES_PRIORIDADE,
        hover_data=[
            "cluster",
            "categoria_vm",
            "status_geral",
            "risco_futuro_90d",
            "acao_final",
            "cpu_p95_pct",
            "mem_p95_pct",
            "disk_p95_pct",
        ],
    )

    fig.update_layout(
        xaxis_title="Score de prioridade",
        yaxis_title="VM",
        yaxis={"categoryorder": "total ascending"},
    )

    return aplicar_layout_plotly(fig, altura=760, showlegend=True)



def obter_mapa_recursos():
    return {
        "CPU": {
            "atual": "cpu_p95_pct",
            "f30": "cpu_forecast_30d",
            "f60": "cpu_forecast_60d",
            "f90": "cpu_forecast_90d",
            "cor": "#2E5B9A",
        },
        "Memória": {
            "atual": "mem_p95_pct",
            "f30": "mem_forecast_30d",
            "f60": "mem_forecast_60d",
            "f90": "mem_forecast_90d",
            "cor": "#F08A24",
        },
        "Disco": {
            "atual": "disk_p95_pct",
            "f30": "disk_forecast_30d",
            "f60": "disk_forecast_60d",
            "f90": "disk_forecast_90d",
            "cor": "#D9534F",
        },
    }


def calcular_tabela_forecast_30_60_90(df: pd.DataFrame, recurso_selecionado: str = "Todos") -> pd.DataFrame:
    """
    Consolida o forecast 30/60/90 dias por recurso.
    Usa média do ambiente filtrado para visão executiva.
    """

    mapa = obter_mapa_recursos()
    recursos = list(mapa.keys()) if recurso_selecionado == "Todos" else [recurso_selecionado]

    linhas = []

    for recurso in recursos:
        cfg = mapa.get(recurso)
        if not cfg:
            continue

        colunas = [cfg["atual"], cfg["f30"], cfg["f60"], cfg["f90"]]
        if not all(c in df.columns for c in colunas):
            continue

        atual = pd.to_numeric(df[cfg["atual"]], errors="coerce")
        f30 = pd.to_numeric(df[cfg["f30"]], errors="coerce")
        f60 = pd.to_numeric(df[cfg["f60"]], errors="coerce")
        f90 = pd.to_numeric(df[cfg["f90"]], errors="coerce")

        linhas.append(
            {
                "recurso": recurso,
                "valor_atual": round(float(atual.mean()), 2) if atual.notna().any() else None,
                "forecast_30d": round(float(f30.mean()), 2) if f30.notna().any() else None,
                "forecast_60d": round(float(f60.mean()), 2) if f60.notna().any() else None,
                "forecast_90d": round(float(f90.mean()), 2) if f90.notna().any() else None,
                "delta_90d": round(float(f90.mean() - atual.mean()), 2) if f90.notna().any() and atual.notna().any() else None,
                "vms_acima_85_90d": int((f90 >= 85).sum()),
                "vms_acima_95_90d": int((f90 >= 95).sum()),
            }
        )

    return pd.DataFrame(linhas)


def criar_grafico_forecast_30_60_90(df: pd.DataFrame, recurso_selecionado: str = "Todos"):
    """
    Gráfico principal de capacity planning: uso atual + previsão 30/60/90 dias.
    Quando recurso = Todos, plota CPU, Memória e Disco juntos.
    Quando recurso = CPU/Memória/Disco, plota somente o recurso selecionado.
    """

    tabela = calcular_tabela_forecast_30_60_90(df, recurso_selecionado)

    if tabela.empty:
        return None

    linhas = []
    for _, row in tabela.iterrows():
        linhas.extend(
            [
                {"recurso": row["recurso"], "periodo": "Atual", "utilizacao_pct": row["valor_atual"]},
                {"recurso": row["recurso"], "periodo": "30 dias", "utilizacao_pct": row["forecast_30d"]},
                {"recurso": row["recurso"], "periodo": "60 dias", "utilizacao_pct": row["forecast_60d"]},
                {"recurso": row["recurso"], "periodo": "90 dias", "utilizacao_pct": row["forecast_90d"]},
            ]
        )

    dados = pd.DataFrame(linhas).dropna(subset=["utilizacao_pct"])
    dados["periodo"] = pd.Categorical(
        dados["periodo"],
        categories=["Atual", "30 dias", "60 dias", "90 dias"],
        ordered=True,
    )
    dados = dados.sort_values(["recurso", "periodo"])

    titulo = "Forecast 30/60/90 dias — CPU, Memória e Disco"
    if recurso_selecionado != "Todos":
        titulo = f"Forecast 30/60/90 dias — {recurso_selecionado}"

    fig = px.line(
        dados,
        x="periodo",
        y="utilizacao_pct",
        color="recurso",
        markers=True,
        title=titulo,
        color_discrete_map=CORES_RECURSO,
    )

    fig.update_traces(line=dict(width=4), marker=dict(size=10))

    fig.add_hline(
        y=85,
        line_dash="dot",
        line_color="#D9A441",
        annotation_text="Atenção 85%",
        annotation_position="top left",
    )
    fig.add_hline(
        y=95,
        line_dash="dot",
        line_color="#D9534F",
        annotation_text="Crítico 95%",
        annotation_position="top left",
    )

    fig.update_layout(
        xaxis_title="Horizonte de previsão",
        yaxis_title="Uso previsto (%)",
        yaxis=dict(range=[0, 105]),
    )

    return aplicar_layout_plotly(fig, altura=520, showlegend=True)


def obter_colunas_top_vms(recurso_selecionado: str):
    colunas_base = [
        "cluster",
        "vm",
        "categoria_vm",
        "status_geral",
        "risco_futuro_90d",
        "score_prioridade",
        "prioridade_final",
        "acao_final",
    ]

    if recurso_selecionado == "CPU":
        colunas_recurso = [
            "status_cpu",
            "cpu_p95_pct",
            "cpu_forecast_30d",
            "cpu_forecast_60d",
            "cpu_forecast_90d",
        ]
    elif recurso_selecionado == "Memória":
        colunas_recurso = [
            "status_memoria",
            "mem_p95_pct",
            "mem_forecast_30d",
            "mem_forecast_60d",
            "mem_forecast_90d",
        ]
    elif recurso_selecionado == "Disco":
        colunas_recurso = [
            "status_disco",
            "disk_p95_pct",
            "disk_forecast_30d",
            "disk_forecast_60d",
            "disk_forecast_90d",
        ]
    else:
        colunas_recurso = [
            "cpu_p95_pct",
            "mem_p95_pct",
            "disk_p95_pct",
            "cpu_forecast_90d",
            "mem_forecast_90d",
            "disk_forecast_90d",
        ]

    return colunas_base + colunas_recurso + ["recomendacao_final"]




# =============================================================================
# Histórico, escopo e drilldown operacional
# =============================================================================
def _coluna_existente(df: pd.DataFrame, candidatos):
    if df is None or df.empty:
        return None
    mapa = {str(c).strip().lower(): c for c in df.columns}
    for c in candidatos:
        if c in df.columns:
            return c
        chave = str(c).strip().lower()
        if chave in mapa:
            return mapa[chave]
    return None


def obter_coluna_host(df: pd.DataFrame):
    return _coluna_existente(
        df,
        [
            "host",
            "Host",
            "host_name",
            "Host Name",
            "esxi",
            "ESXi",
            "mapping_parent_name",
            "parent_host",
            "nome_host",
        ],
    )


def preparar_coluna_host(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    coluna_host = obter_coluna_host(df)
    if coluna_host:
        df["host_visual"] = df[coluna_host].fillna("Host não informado").astype(str)
    else:
        df["host_visual"] = "Host não informado"
    return df


def normalizar_recurso(valor):
    texto = str(valor).strip().upper()
    if texto in ["CPU", "PROCESSADOR"] or "CPU" in texto:
        return "CPU"
    if texto in ["MEM", "MEMORIA", "MEMÓRIA", "MEMORY", "RAM"] or "MEM" in texto or "RAM" in texto:
        return "Memória"
    if texto in ["DISK", "DISCO", "STORAGE"] or "DISK" in texto or "DISCO" in texto or "STORAGE" in texto:
        return "Disco"
    return str(valor)


def normalizar_historico_recurso(df: pd.DataFrame, recurso: str) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame()

    df = df.copy()

    col_data = _coluna_existente(df, ["date", "data", "timestamp", "time", "datetime"])
    col_used = _coluna_existente(df, ["used_pct", "usage_pct", "valor", "value", "utilizacao_pct"])
    col_cluster = _coluna_existente(df, ["cluster", "Cluster", "cluster_name"])
    col_vm = _coluna_existente(df, ["vm", "VM", "resource", "name", "Nome VM"])
    col_id = _coluna_existente(df, ["vm_resource_id", "resource_id", "id"])
    col_host = obter_coluna_host(df)

    if col_data is None or col_used is None:
        return pd.DataFrame()

    out = pd.DataFrame()
    out["date"] = pd.to_datetime(df[col_data], errors="coerce")
    out["used_pct"] = pd.to_numeric(df[col_used], errors="coerce")
    out["recurso"] = recurso
    out["cluster"] = df[col_cluster].astype(str) if col_cluster else "Cluster não informado"
    out["vm"] = df[col_vm].astype(str) if col_vm else "VM não informada"
    out["vm_resource_id"] = df[col_id].astype(str) if col_id else ""
    out["host_visual"] = df[col_host].astype(str) if col_host else "Host não informado"

    out = out.dropna(subset=["date", "used_pct"])

    return out


@st.cache_data(show_spinner=False)
def carregar_historico_excel(caminho_arquivo: str) -> pd.DataFrame:
    arquivo = Path(caminho_arquivo)
    historicos = []

    try:
        df_cpu = carregar_aba(arquivo, "HIST_CPU")
        historicos.append(normalizar_historico_recurso(df_cpu, "CPU"))
    except Exception:
        pass

    try:
        df_mem = carregar_aba(arquivo, "HIST_MEM")
        historicos.append(normalizar_historico_recurso(df_mem, "Memória"))
    except Exception:
        pass

    try:
        df_disk = carregar_aba(arquivo, "HIST_DISK")
        hist_disk = normalizar_historico_recurso(df_disk, "Disco")
        if not hist_disk.empty:
            # Para disco, uma VM pode ter várias partições. Para capacity, usamos o maior uso do dia.
            chaves = ["date", "cluster", "host_visual", "vm", "vm_resource_id", "recurso"]
            hist_disk = hist_disk.groupby(chaves, dropna=False, as_index=False)["used_pct"].max()
        historicos.append(hist_disk)
    except Exception:
        pass

    historicos = [h for h in historicos if h is not None and not h.empty]
    if not historicos:
        return pd.DataFrame(columns=["date", "used_pct", "recurso", "cluster", "host_visual", "vm", "vm_resource_id"])

    return pd.concat(historicos, ignore_index=True)


def carregar_historico_banco(caminho_banco, execution_id=None) -> pd.DataFrame:
    if duckdb is None or not Path(caminho_banco).exists():
        return pd.DataFrame()

    try:
        con = duckdb.connect(str(caminho_banco), read_only=True)
        tabelas = con.execute("SHOW TABLES").fetchdf()
        nomes_tabelas = tabelas.iloc[:, 0].astype(str).tolist()
        tabela = None
        for nome in nomes_tabelas:
            if "historico" in nome.lower() or "hist" in nome.lower():
                tabela = nome
                break
        if tabela is None:
            con.close()
            return pd.DataFrame()

        cols = con.execute(f"DESCRIBE {tabela}").fetchdf()
        col_names = cols.iloc[:, 0].astype(str).tolist()

        if execution_id and "execution_id" in col_names:
            raw = con.execute(f"SELECT * FROM {tabela} WHERE execution_id = ?", [str(execution_id)]).fetchdf()
        else:
            raw = con.execute(f"SELECT * FROM {tabela}").fetchdf()
        con.close()

        if raw.empty:
            return pd.DataFrame()

        col_recurso = _coluna_existente(raw, ["recurso", "resource", "metric", "tipo_recurso"])
        if col_recurso:
            raw["__recurso_norm"] = raw[col_recurso].map(normalizar_recurso)
            partes = []
            for recurso in ["CPU", "Memória", "Disco"]:
                partes.append(normalizar_historico_recurso(raw[raw["__recurso_norm"] == recurso], recurso))
            partes = [p for p in partes if p is not None and not p.empty]
            return pd.concat(partes, ignore_index=True) if partes else pd.DataFrame()

        return pd.DataFrame()
    except Exception:
        return pd.DataFrame()


def criar_resumo_host(df: pd.DataFrame) -> pd.DataFrame:
    df = preparar_coluna_host(df)
    if df.empty or "host_visual" not in df.columns:
        return pd.DataFrame()

    resumo = (
        df.groupby("host_visual", dropna=False)
        .agg(
            total_vms=("vm", "count"),
            p0_acao_imediata=("prioridade_final", lambda x: (x == "P0_ACAO_IMEDIATA").sum()),
            p1_alta=("prioridade_final", lambda x: (x == "P1_ALTA").sum()),
            score_max=("score_prioridade", "max"),
            score_medio=("score_prioridade", "mean"),
            cpu_pico_uso=("cpu_p95_pct", "mean"),
            memoria_pico_uso=("mem_p95_pct", "mean"),
            disco_pico_uso=("disk_p95_pct", "mean"),
            cpu_previsao_90d=("cpu_forecast_90d", "mean"),
            memoria_previsao_90d=("mem_forecast_90d", "mean"),
            disco_previsao_90d=("disk_forecast_90d", "mean"),
        )
        .reset_index()
        .rename(columns={"host_visual": "host"})
    )

    resumo["vms_prioritarias"] = resumo["p0_acao_imediata"] + resumo["p1_alta"]
    resumo = resumo.sort_values(["vms_prioritarias", "score_max"], ascending=False)
    return resumo


def obter_linha_grid(evento):
    try:
        linhas = evento.selection.rows
        return linhas[0] if linhas else None
    except Exception:
        pass

    try:
        linhas = evento.get("selection", {}).get("rows", [])
        return linhas[0] if linhas else None
    except Exception:
        return None


def filtrar_historico_por_escopo(df_hist: pd.DataFrame, df_scope: pd.DataFrame, tipo_escopo: str, nome_escopo: str) -> pd.DataFrame:
    if df_hist is None or df_hist.empty:
        return pd.DataFrame()

    hist = df_hist.copy()
    df_scope = preparar_coluna_host(df_scope)

    if tipo_escopo == "VM":
        ids = []
        if "vm_resource_id" in df_scope.columns:
            ids = df_scope["vm_resource_id"].dropna().astype(str).unique().tolist()
        if ids and "vm_resource_id" in hist.columns:
            return hist[hist["vm_resource_id"].astype(str).isin(ids)].copy()
        vms = df_scope["vm"].dropna().astype(str).unique().tolist()
        return hist[hist["vm"].astype(str).isin(vms)].copy()

    if tipo_escopo == "Host":
        if "host_visual" in hist.columns:
            por_host = hist[hist["host_visual"].astype(str) == str(nome_escopo)].copy()
            if not por_host.empty:
                return por_host
        vms = df_scope["vm"].dropna().astype(str).unique().tolist()
        return hist[hist["vm"].astype(str).isin(vms)].copy()

    if tipo_escopo == "Cluster":
        return hist[hist["cluster"].astype(str) == str(nome_escopo)].copy()

    vms = df_scope["vm"].dropna().astype(str).unique().tolist() if "vm" in df_scope.columns else []
    return hist[hist["vm"].astype(str).isin(vms)].copy() if vms else hist


def criar_grafico_historico_forecast_escopo(
    df_scope: pd.DataFrame,
    df_hist: pd.DataFrame,
    tipo_escopo: str,
    nome_escopo: str,
    recurso_selecionado: str = "Todos",
):
    mapa = obter_mapa_recursos()
    recursos = list(mapa.keys()) if recurso_selecionado == "Todos" else [recurso_selecionado]
    df_scope = preparar_coluna_host(df_scope)
    hist_scope = filtrar_historico_por_escopo(df_hist, df_scope, tipo_escopo, nome_escopo)

    fig = go.Figure()
    tem_dados = False

    for recurso in recursos:
        cfg = mapa.get(recurso)
        if not cfg:
            continue

        cor = cfg["cor"]
        hist_recurso = pd.DataFrame()
        if hist_scope is not None and not hist_scope.empty:
            hist_recurso = hist_scope[hist_scope["recurso"] == recurso].copy()

        ultima_data = pd.Timestamp.today().normalize()
        ultimo_valor = None

        if not hist_recurso.empty:
            serie = (
                hist_recurso.groupby("date", as_index=False)["used_pct"]
                .mean()
                .sort_values("date")
            )
            if not serie.empty:
                tem_dados = True
                ultima_data = serie["date"].max()
                ultimo_valor = float(serie.loc[serie["date"].idxmax(), "used_pct"])
                fig.add_trace(
                    go.Scatter(
                        x=serie["date"],
                        y=serie["used_pct"],
                        mode="lines",
                        name=f"{recurso} histórico",
                        line=dict(color=cor, width=2),
                    )
                )

        colunas_forecast = [cfg["atual"], cfg["f30"], cfg["f60"], cfg["f90"]]
        if all(c in df_scope.columns for c in colunas_forecast):
            atual = pd.to_numeric(df_scope[cfg["atual"]], errors="coerce").mean()
            f30 = pd.to_numeric(df_scope[cfg["f30"]], errors="coerce").mean()
            f60 = pd.to_numeric(df_scope[cfg["f60"]], errors="coerce").mean()
            f90 = pd.to_numeric(df_scope[cfg["f90"]], errors="coerce").mean()

            base = ultimo_valor if ultimo_valor is not None else atual
            pontos = [
                (ultima_data, base),
                (ultima_data + pd.Timedelta(days=30), f30),
                (ultima_data + pd.Timedelta(days=60), f60),
                (ultima_data + pd.Timedelta(days=90), f90),
            ]
            pontos = [(x, y) for x, y in pontos if pd.notna(y)]

            if len(pontos) >= 2:
                tem_dados = True
                fig.add_trace(
                    go.Scatter(
                        x=[p[0] for p in pontos],
                        y=[float(p[1]) for p in pontos],
                        mode="lines+markers",
                        name=f"{recurso} previsão 30/60/90",
                        line=dict(color=cor, width=4, dash="dash"),
                        marker=dict(size=9),
                    )
                )

    if not tem_dados:
        return None

    fig.add_hline(y=85, line_dash="dot", line_color="#D9A441", annotation_text="Atenção 85%")
    fig.add_hline(y=95, line_dash="dot", line_color="#D9534F", annotation_text="Crítico 95%")

    fig.update_layout(
        title=f"Histórico real + previsão 30/60/90 dias — {tipo_escopo}: {nome_escopo}",
        xaxis_title="Data",
        yaxis_title="Uso (%)",
        yaxis=dict(range=[0, 105]),
        hovermode="x unified",
    )

    return aplicar_layout_plotly(fig, altura=620, showlegend=True)


def criar_tabela_forecast_escopo(df_scope: pd.DataFrame, recurso_selecionado: str):
    tabela = calcular_tabela_forecast_30_60_90(df_scope, recurso_selecionado)
    return renomear_colunas_visual(tabela) if not tabela.empty else tabela




# =============================================================================
# Recursos provisionados, Top 30 clicável e Data Stores
# =============================================================================
COLUNAS_VISUAIS.update(
    {
        "host_visual": "Host",
        "vcpus_total": "vCPU total",
        "memoria_total_gb": "Memória total (GB)",
        "disco_total_gb": "Disco total (GB)",
        "cpu_alocada": "vCPU alocada",
        "memoria_alocada_gb": "Memória alocada (GB)",
        "disco_alocado_gb": "Disco alocado (GB)",
        "cpu_ocupacao_pct": "CPU — ocupação / pico típico (%)",
        "memoria_ocupacao_pct": "Memória — ocupação / pico típico (%)",
        "disco_ocupacao_pct": "Disco — ocupação / pico típico (%)",
        "datastore": "Data store",
        "datastore_total_gb": "Data store — capacidade total (GB)",
        "datastore_usado_gb": "Data store — usado (GB)",
        "datastore_livre_gb": "Data store — livre (GB)",
        "datastore_ocupacao_pct": "Data store — ocupação (%)",
    }
)


def _serie_numerica(df: pd.DataFrame, coluna):
    if coluna is None or coluna not in df.columns:
        return pd.Series(dtype="float64")
    return pd.to_numeric(df[coluna], errors="coerce")


def _fmt_valor(valor, unidade="", vazio="Não coletado"):
    if valor is None or pd.isna(valor):
        return vazio
    try:
        n = float(valor)
        if not math.isfinite(n):
            return vazio
        if abs(n - round(n)) < 0.005 and unidade == "":
            txt = f"{int(round(n)):,}".replace(",", ".")
        else:
            txt = f"{n:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        return txt + unidade
    except Exception:
        return str(valor) + unidade

def _fmt_ratio(valor, vazio="Não coletado"):
    if valor is None or pd.isna(valor):
        return vazio
    try:
        n = float(valor)
        if not math.isfinite(n):
            return vazio
        return f"{n:.2f}:1".replace(".", ",")
    except Exception:
        return str(valor)


def _fmt_card_count(valor, vazio="N/C"):
    if valor is None or pd.isna(valor):
        return vazio
    try:
        n = float(valor)
        if not math.isfinite(n):
            return vazio
        return f"{int(round(n)):,}".replace(",", ".")
    except Exception:
        return str(valor)


def _fmt_card_pct(valor, vazio="N/C"):
    if valor is None or pd.isna(valor):
        return vazio
    try:
        n = float(valor)
        if not math.isfinite(n):
            return vazio
        return f"{n:.2f}%".replace(".", ",")
    except Exception:
        return str(valor)


def _fmt_card_ratio(valor, vazio="N/C"):
    return _fmt_ratio(valor, vazio=vazio)


def _fmt_card_storage_gb(valor, vazio="N/C"):
    if valor is None or pd.isna(valor):
        return vazio
    try:
        n = float(valor)
        if not math.isfinite(n):
            return vazio
        abs_n = abs(n)
        if abs_n >= 1000:
            return f"{n/1000:.1f} TB".replace(".", ",")
        if abs_n >= 100:
            return f"{n:,.0f} GB".replace(",", "X").replace(".", ",").replace("X", ".")
        return f"{n:,.2f} GB".replace(",", "X").replace(".", ",").replace("X", ".")
    except Exception:
        return str(valor)

def obter_coluna_host_explicita(df: pd.DataFrame):
    """Host só é considerado disponível quando a coleta trouxe coluna explícita de host/ESXi."""
    return _coluna_existente(
        df,
        [
            "host",
            "Host",
            "host_name",
            "Host Name",
            "esxi",
            "ESXi",
            "parent_host",
            "nome_host",
        ],
    )


def host_esta_disponivel(df: pd.DataFrame) -> bool:
    col = obter_coluna_host_explicita(df)
    if col is None or col not in df.columns:
        return False
    valores = df[col].dropna().astype(str).str.strip()
    valores = valores[valores != ""]
    if valores.empty:
        return False
    ruins = {"host não informado", "na", "none", "null", "nan", "não informado"}
    valores_validos = [v for v in valores.unique().tolist() if v.strip().lower() not in ruins]
    return len(valores_validos) > 0


def obter_colunas_capacidade(df: pd.DataFrame):
    return {
        "cpu": _coluna_existente(
            df,
            [
                "cpu_alloc",
                "cpu_alloc_vcpu",
                "vcpus",
                "vcpu",
                "num_cpu",
                "cpu_count",
                "cpus",
                "vCPU",
                "CPU Allocated",
            ],
        ),
        "mem": _coluna_existente(
            df,
            [
                "mem_alloc_gb",
                "memory_gb",
                "mem_gb",
                "ram_gb",
                "configured_memory_gb",
                "memory_allocated_gb",
                "Memória GB",
                "Memory GB",
            ],
        ),
        "disk": _coluna_existente(
            df,
            [
                "disk_alloc_gb",
                "disk_capacity_gb",
                "capacity_gb",
                "storage_gb",
                "provisioned_gb",
                "Disco GB",
                "Disk GB",
            ],
        ),
    }


def calcular_resumo_recursos(
    df_scope: pd.DataFrame,
    caminho_banco=None,
    tipo_escopo=None,
    nome_escopo=None,
    run_id=None,
) -> dict:
    """
    HOTFIX 15F.10.9.
    Não deixa cards ficarem "Não coletado" quando o DuckDB retornou resumo parcial.
    Faz merge:
    1) DuckDB / tabelas novas;
    2) colunas disponíveis no df_scope;
    3) faltantes reais.
    """
    if df_scope is None:
        df_scope = pd.DataFrame()

    cols = obter_colunas_capacidade(df_scope) if not df_scope.empty else {"cpu": None, "mem": None, "disk": None}

    def _mean_any(candidates):
        if df_scope.empty:
            return None
        for c in candidates:
            if c in df_scope.columns:
                s = _serie_numerica(df_scope, c)
                if s.notna().any():
                    return float(s.mean())
        return None

    local = {
        "vms": int(len(df_scope)) if not df_scope.empty else 0,
        "vcpus": _serie_numerica(df_scope, cols.get("cpu")).sum(min_count=1) if cols.get("cpu") else None,
        "memoria_gb": _serie_numerica(df_scope, cols.get("mem")).sum(min_count=1) if cols.get("mem") else None,
        "disco_gb": _serie_numerica(df_scope, cols.get("disk")).sum(min_count=1) if cols.get("disk") else None,
        "cpu_pct": _mean_any(["cpu_p95_pct", "cpu_usage_pct", "cpu_used_pct", "CPU — pico médio por cluster (%)", "CPU — pico típico de uso (%)"]),
        "mem_pct": _mean_any(["mem_p95_pct", "mem_usage_pct", "mem_used_pct", "Memória — pico médio por cluster (%)", "Memória — pico típico de uso (%)"]),
        "disk_pct": _mean_any(["disk_p95_pct", "disk_usage_pct", "disk_used_pct", "Disco — pico médio por cluster (%)", "Disco — pico típico de uso (%)"]),
        "fonte_alocacao": "df_scope",
    }

    resumo = {}
    try:
        if caminho_banco is None:
            caminho_banco = CAMINHO_BANCO
        if run_id is None:
            run_id = st.session_state.get("resource_run_id_ativo") or st.session_state.get("execution_id_ativo")
        db_resumo = calcular_alocacoes_cards(
            df_scope=df_scope,
            caminho_banco=caminho_banco,
            tipo_escopo=tipo_escopo,
            nome_escopo=nome_escopo,
            run_id=run_id,
        )
        if isinstance(db_resumo, dict):
            resumo.update(db_resumo)
    except Exception:
        resumo = {}

    # Preenche o que veio faltando do DuckDB usando o df_scope.
    for key, value in local.items():
        if key not in resumo or resumo.get(key) is None or (isinstance(resumo.get(key), float) and not math.isfinite(resumo.get(key))):
            if value is not None and not (isinstance(value, float) and not math.isfinite(value)):
                resumo[key] = value

    # Recalcula ratios após o merge.
    def _ratio_calc(a, b):
        try:
            if a is None or b is None:
                return None
            a = float(a)
            b = float(b)
            if not math.isfinite(a) or not math.isfinite(b) or b <= 0:
                return None
            return round(a / b, 2)
        except Exception:
            return None

    resumo["cpu_over_ratio"] = _ratio_calc(resumo.get("vcpus"), resumo.get("pcpus"))
    resumo["memoria_over_ratio"] = _ratio_calc(resumo.get("memoria_gb"), resumo.get("host_memoria_gb"))
    resumo["disk_over_ratio"] = _ratio_calc(resumo.get("disco_gb"), resumo.get("datastore_total_gb"))

    faltantes = []
    for key, label in [
        ("vcpus", "vCPU alocada"),
        ("pcpus", "pCPU/cores físicos"),
        ("cpu_over_ratio", "over ratio CPU"),
        ("memoria_gb", "memória alocada"),
        ("host_memoria_gb", "memória física"),
        ("memoria_over_ratio", "over ratio memória"),
        ("disco_gb", "disco provisionado"),
        ("datastore_total_gb", "datastore total"),
        ("disk_over_ratio", "over ratio disco"),
    ]:
        val = resumo.get(key)
        if val is None or (isinstance(val, float) and not math.isfinite(val)):
            faltantes.append(label)

    resumo["campos_faltantes"] = faltantes
    if "fonte_alocacao" not in resumo:
        resumo["fonte_alocacao"] = "df_scope"

    return resumo


def render_resumo_recursos(
    df_scope: pd.DataFrame,
    titulo: str,
    caminho_banco=None,
    tipo_escopo=None,
    nome_escopo=None,
    run_id=None,
):
    resumo = calcular_resumo_recursos(
        df_scope=df_scope,
        caminho_banco=caminho_banco,
        tipo_escopo=tipo_escopo,
        nome_escopo=nome_escopo,
        run_id=run_id,
    )

    st.markdown(f'<div class="section-title">{titulo}</div>', unsafe_allow_html=True)

    cards = [
        ("VMs", _fmt_card_count(resumo.get("vms", 0)), "kpi-blue", "Qtd."),
        ("Hosts", _fmt_card_count(resumo.get("host_count")), "kpi-green", "Escopo"),
        ("Datastores", _fmt_card_count(resumo.get("datastore_count")), "kpi-blue", "Qtd."),
        ("vCPU", _fmt_card_count(resumo.get("vcpus")), "kpi-cyan", "Aloc."),
        ("pCPU", _fmt_card_count(resumo.get("pcpus")), "kpi-green", "Físico"),
        ("CPU Over", _fmt_card_ratio(resumo.get("cpu_over_ratio")), "kpi-yellow", "vCPU/pCPU"),
        ("Mem Over", _fmt_card_ratio(resumo.get("memoria_over_ratio")), "kpi-orange", "Aloc/Fís."),
        ("Disk Over", _fmt_card_ratio(resumo.get("disk_over_ratio")), "kpi-yellow", "Prov/Cap."),
        ("Mem Aloc.", _fmt_card_storage_gb(resumo.get("memoria_gb")), "kpi-orange", "Memória"),
        ("Mem Física", _fmt_card_storage_gb(resumo.get("host_memoria_gb")), "kpi-green", "Hosts"),
        ("Disco Prov.", _fmt_card_storage_gb(resumo.get("disco_gb")), "kpi-yellow", "Provisionado"),
        ("DS Cap.", _fmt_card_storage_gb(resumo.get("datastore_total_gb")), "kpi-cyan", "Capacidade"),
        ("CPU", _fmt_card_pct(resumo.get("cpu_pct")), "kpi-blue", "P95"),
        ("Memória", _fmt_card_pct(resumo.get("mem_pct")), "kpi-orange", "P95"),
        ("Disco", _fmt_card_pct(resumo.get("disk_pct")), "kpi-red", "P95"),
    ]

    for i in range(0, len(cards), 5):
        cols = st.columns(5)
        for col, (label, value, css_class, foot) in zip(cols, cards[i:i+5]):
            with col:
                render_kpi_card(label, value, css_class, foot)

    fonte = resumo.get("fonte_alocacao", "não informada")
    faltantes = resumo.get("campos_faltantes", []) or []

    detalhes = []
    for key, label in [
        ("cpu_over_ratio", "cpu_over"),
        ("memoria_over_ratio", "mem_over"),
        ("disk_over_ratio", "disk_over"),
    ]:
        if resumo.get(key) is not None:
            detalhes.append(label + "=" + _fmt_ratio(resumo.get(key), vazio="N/C"))
    if resumo.get("datastore_total_gb") is not None:
        detalhes.append("datastore_total=" + _fmt_card_storage_gb(resumo.get("datastore_total_gb"), vazio="N/C"))

    if faltantes:
        st.caption(
            "Alocação: fonte="
            + str(fonte)
            + (" | " + " | ".join(detalhes) if detalhes else "")
            + " | Campos ainda não coletados: "
            + ", ".join([str(x) for x in faltantes])
            + "."
        )
    else:
        st.caption(
            "Alocação: fonte="
            + str(fonte)
            + (" | " + " | ".join(detalhes) if detalhes else "")
            + " | cards completos."
        )


def obter_colunas_vm_recursos(df: pd.DataFrame, recurso_selecionado: str = "Todos"):
    cols_cap = obter_colunas_capacidade(df)
    colunas = ["cluster"]
    if host_esta_disponivel(df):
        colunas.append(obter_coluna_host_explicita(df))
    colunas += ["vm", "categoria_vm", "prioridade_final", "acao_final", "score_prioridade"]

    for col in [cols_cap["cpu"], cols_cap["mem"], cols_cap["disk"]]:
        if col and col not in colunas:
            colunas.append(col)

    if recurso_selecionado == "CPU":
        colunas += ["status_cpu", "cpu_p95_pct", "cpu_forecast_30d", "cpu_forecast_60d", "cpu_forecast_90d"]
    elif recurso_selecionado == "Memória":
        colunas += ["status_memoria", "mem_p95_pct", "mem_forecast_30d", "mem_forecast_60d", "mem_forecast_90d"]
    elif recurso_selecionado == "Disco":
        colunas += ["status_disco", "disk_p95_pct", "disk_forecast_30d", "disk_forecast_60d", "disk_forecast_90d"]
    else:
        colunas += [
            "status_cpu",
            "cpu_p95_pct",
            "cpu_forecast_90d",
            "status_memoria",
            "mem_p95_pct",
            "mem_forecast_90d",
            "status_disco",
            "disk_p95_pct",
            "disk_forecast_90d",
        ]

    colunas += ["risco_futuro_90d", "recomendacao_final"]
    colunas = [c for c in colunas if c and c in df.columns]
    return list(dict.fromkeys(colunas))


def ordenar_vms_por_recurso(df: pd.DataFrame, recurso_selecionado: str):
    if recurso_selecionado == "CPU" and "cpu_p95_pct" in df.columns:
        return df.sort_values(["cpu_p95_pct", "score_prioridade"], ascending=False)
    if recurso_selecionado == "Memória" and "mem_p95_pct" in df.columns:
        return df.sort_values(["mem_p95_pct", "score_prioridade"], ascending=False)
    if recurso_selecionado == "Disco" and "disk_p95_pct" in df.columns:
        return df.sort_values(["disk_p95_pct", "score_prioridade"], ascending=False)
    return df.sort_values("score_prioridade", ascending=False)


def render_detalhe_vm_recursos(vm_df: pd.DataFrame):
    if vm_df is None or vm_df.empty:
        return

    vm = vm_df.iloc[0]
    cols_cap = obter_colunas_capacidade(vm_df)

    def valor_col(col):
        if col and col in vm_df.columns:
            return vm.get(col)
        return None

    linhas = [
        {
            "Recurso": "CPU",
            "Quantidade provisionada": _fmt_valor(valor_col(cols_cap["cpu"]), " vCPU"),
            "Uso atual — pico típico (%)": _fmt_valor(vm.get("cpu_p95_pct"), "%"),
            "Previsão 30 dias (%)": _fmt_valor(vm.get("cpu_forecast_30d"), "%"),
            "Previsão 60 dias (%)": _fmt_valor(vm.get("cpu_forecast_60d"), "%"),
            "Previsão 90 dias (%)": _fmt_valor(vm.get("cpu_forecast_90d"), "%"),
            "Status": vm.get("status_cpu", ""),
        },
        {
            "Recurso": "Memória",
            "Quantidade provisionada": _fmt_valor(valor_col(cols_cap["mem"]), " GB"),
            "Uso atual — pico típico (%)": _fmt_valor(vm.get("mem_p95_pct"), "%"),
            "Previsão 30 dias (%)": _fmt_valor(vm.get("mem_forecast_30d"), "%"),
            "Previsão 60 dias (%)": _fmt_valor(vm.get("mem_forecast_60d"), "%"),
            "Previsão 90 dias (%)": _fmt_valor(vm.get("mem_forecast_90d"), "%"),
            "Status": vm.get("status_memoria", ""),
        },
        {
            "Recurso": "Disco",
            "Quantidade provisionada": _fmt_valor(valor_col(cols_cap["disk"]), " GB"),
            "Uso atual — pico típico (%)": _fmt_valor(vm.get("disk_p95_pct"), "%"),
            "Previsão 30 dias (%)": _fmt_valor(vm.get("disk_forecast_30d"), "%"),
            "Previsão 60 dias (%)": _fmt_valor(vm.get("disk_forecast_60d"), "%"),
            "Previsão 90 dias (%)": _fmt_valor(vm.get("disk_forecast_90d"), "%"),
            "Status": vm.get("status_disco", ""),
        },
    ]

    st.markdown(f"#### Recursos da VM selecionada — {vm.get('vm', '')}")
    st.dataframe(formatar_dataframe_visual_2_casas(pd.DataFrame(linhas)), width="stretch", hide_index=True)

    if vm.get("recomendacao_final"):
        st.info(str(vm.get("recomendacao_final")))


def render_top30_vms_clicavel(df_scope: pd.DataFrame, recurso_selecionado: str, key: str):
    if df_scope is None or df_scope.empty:
        st.info("Nenhuma VM disponível para o escopo selecionado.")
        return None, pd.DataFrame()

    df_top = ordenar_vms_por_recurso(df_scope.copy(), recurso_selecionado).head(30).copy()
    colunas = obter_colunas_vm_recursos(df_top, recurso_selecionado)
    df_grid = df_top[colunas].copy()

    evento = st.dataframe(
        formatar_dataframe_visual_2_casas(renomear_colunas_visual(df_grid)),
        width="stretch",
        hide_index=True,
        on_select="rerun",
        selection_mode="single-row",
        key=key,
    )

    idx = obter_linha_grid(evento)
    if idx is None or idx >= len(df_top):
        return None, df_top

    vm_linha = df_top.iloc[[idx]].copy()
    return vm_linha, df_top


def obter_colunas_datastore(df: pd.DataFrame):
    return {
        "nome": _coluna_existente(
            df,
            [
                "datastore",
                "datastore_name",
                "data_store",
                "ds_name",
                "nome_datastore",
                "Datastore",
                "Data Store",
            ],
        ),
        "capacidade": _coluna_existente(
            df,
            [
                "datastore_capacity_gb",
                "datastore_total_gb",
                "ds_capacity_gb",
                "capacity_datastore_gb",
            ],
        ),
        "usado": _coluna_existente(
            df,
            [
                "datastore_used_gb",
                "ds_used_gb",
                "used_datastore_gb",
                "datastore_usado_gb",
            ],
        ),
        "livre": _coluna_existente(
            df,
            [
                "datastore_free_gb",
                "ds_free_gb",
                "free_datastore_gb",
                "datastore_livre_gb",
            ],
        ),
        "ocupacao": _coluna_existente(
            df,
            [
                "datastore_used_pct",
                "datastore_usage_pct",
                "ds_used_pct",
                "datastore_ocupacao_pct",
            ],
        ),
    }


COLUNAS_VISUAIS.update({
    "datastore_provisioned_gb": "Datastore — provisionado (GB)",
    "datastore_provisioned_pct": "Datastore — provisionamento (%)",
    "datastore_used_pct": "Datastore — ocupação real (%)",
})


def render_datastores(
    df_scope: pd.DataFrame,
    caminho_banco=None,
    tipo_escopo=None,
    nome_escopo=None,
    run_id=None,
):
    st.markdown('<div class="section-title">Data stores</div>', unsafe_allow_html=True)

    cols = obter_colunas_datastore(df_scope)

    # Primeiro tenta a coleta nova persistida no DuckDB.
    df_ds_db = pd.DataFrame()
    try:
        df_ds_db = carregar_datastores_escopo(
            df_scope=df_scope,
            caminho_banco=caminho_banco or CAMINHO_BANCO,
            tipo_escopo=tipo_escopo,
            nome_escopo=nome_escopo,
            run_id=run_id or st.session_state.get("resource_run_id_ativo"),
        )
    except Exception:
        df_ds_db = pd.DataFrame()

    if df_ds_db is not None and not df_ds_db.empty:
        st.dataframe(formatar_dataframe_visual_2_casas(renomear_colunas_visual(df_ds_db)), width="stretch", hide_index=True)
        return

    if cols["nome"] is None:
        st.info(
            "Data stores não aparecem porque a coleta atual ainda não trouxe datastore. "
            "Rode nova coleta manual após o hotfix 15F.10.2 para popular datastore_capacity_snapshots."
        )
        return

    df = df_scope.copy()
    df["datastore"] = df[cols["nome"]].fillna("Datastore não informado").astype(str)

    agg = {"vms": ("vm", "count")}
    if cols["capacidade"]:
        df["datastore_total_gb"] = _serie_numerica(df, cols["capacidade"])
        agg["datastore_total_gb"] = ("datastore_total_gb", "sum")
    if cols["usado"]:
        df["datastore_used_gb"] = _serie_numerica(df, cols["usado"])
        agg["datastore_used_gb"] = ("datastore_used_gb", "sum")
    if cols["livre"]:
        df["datastore_free_gb"] = _serie_numerica(df, cols["livre"])
        agg["datastore_free_gb"] = ("datastore_free_gb", "sum")
    if cols["ocupacao"]:
        df["datastore_used_pct"] = _serie_numerica(df, cols["ocupacao"])
        agg["datastore_used_pct"] = ("datastore_used_pct", "mean")

    resumo = df.groupby("datastore", dropna=False).agg(**agg).reset_index()
    st.dataframe(formatar_dataframe_visual_2_casas(renomear_colunas_visual(resumo)), width="stretch", hide_index=True)


# =============================================================================
# Página: Comparação entre execuções
# =============================================================================
def render_pagina_comparacao():
    st.markdown('<div class="section-title">Comparação entre execuções</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="section-subtitle">Compare duas coletas salvas no DuckDB para identificar VMs novas, removidas, pioras, melhorias e variações por cluster.</div>',
        unsafe_allow_html=True,
    )

    execucoes = carregar_execucoes_com_label(CAMINHO_BANCO)

    if execucoes.empty or len(execucoes) < 2:
        st.warning("É necessário ter pelo menos duas execuções salvas no DuckDB para comparar.")
        st.stop()

    label_para_id = dict(zip(execucoes["label"], execucoes["execution_id"]))

    col_a, col_b = st.columns(2)

    with col_a:
        label_anterior = st.selectbox(
            "Execução anterior",
            options=execucoes["label"].tolist(),
            index=1 if len(execucoes) > 1 else 0,
        )

    with col_b:
        label_atual = st.selectbox(
            "Execução atual",
            options=execucoes["label"].tolist(),
            index=0,
        )

    execution_id_anterior = label_para_id[label_anterior]
    execution_id_atual = label_para_id[label_atual]

    if execution_id_anterior == execution_id_atual:
        st.warning("Selecione duas execuções diferentes.")
        st.stop()

    with st.spinner("Comparando execuções..."):
        df_comp_vms = comparar_execucoes_vms(
            caminho_banco=CAMINHO_BANCO,
            execution_id_anterior=execution_id_anterior,
            execution_id_atual=execution_id_atual,
        )

        df_comp_cluster = comparar_resumo_cluster(
            caminho_banco=CAMINHO_BANCO,
            execution_id_anterior=execution_id_anterior,
            execution_id_atual=execution_id_atual,
        )

    resumo_comp = resumir_comparacao_execucoes(df_comp_vms)

    st.markdown('<div class="section-title">Resumo da comparação</div>', unsafe_allow_html=True)

    c1, c2, c3, c4, c5, c6 = st.columns(6)

    with c1:
        render_kpi_card("VMs comparadas", formatar_numero(resumo_comp["total_vms_comparadas"]), "kpi-blue", "Total")
    with c2:
        render_kpi_card("Novas", formatar_numero(resumo_comp["vms_novas"]), "kpi-cyan", "Entraram")
    with c3:
        render_kpi_card("Removidas", formatar_numero(resumo_comp["vms_removidas"]), "kpi-yellow", "Saíram")
    with c4:
        render_kpi_card("Pioraram", formatar_numero(resumo_comp["pioraram"] + resumo_comp["pioraram_muito"]), "kpi-red", "Atenção")
    with c5:
        render_kpi_card("Melhoraram", formatar_numero(resumo_comp["melhoraram"] + resumo_comp["melhoraram_muito"]), "kpi-green", "Redução")
    with c6:
        render_kpi_card("Sem mudança", formatar_numero(resumo_comp["sem_mudanca"]), "kpi-blue", "Estáveis")

    st.markdown('<div class="section-title">Mudança de prioridade</div>', unsafe_allow_html=True)

    dados_mudanca = df_comp_vms["mudanca_prioridade"].value_counts().reset_index()
    dados_mudanca.columns = ["mudanca_prioridade", "quantidade"]

    fig_mudanca = px.bar(
        dados_mudanca,
        x="mudanca_prioridade",
        y="quantidade",
        text="quantidade",
        title="Distribuição de mudança de prioridade por VM",
        color_discrete_sequence=["#2E5B9A"],
    )

    fig_mudanca.update_layout(
        xaxis_title="Mudança",
        yaxis_title="Quantidade de VMs",
    )

    st.plotly_chart(aplicar_layout_plotly(fig_mudanca, altura=420, showlegend=False), width="stretch")

    st.markdown('<div class="section-title">Clusters com maior aumento de VMs prioritárias</div>', unsafe_allow_html=True)

    if not df_comp_cluster.empty:
        top_delta_clusters = df_comp_cluster.sort_values(
            ["delta_vms_prioritarias", "delta_p0_acao_imediata"],
            ascending=False,
        ).head(10)

        fig_cluster_delta = px.bar(
            top_delta_clusters,
            x="delta_vms_prioritarias",
            y="cluster",
            orientation="h",
            text="delta_vms_prioritarias",
            title="Top clusters por aumento de VMs P0/P1",
            color_discrete_sequence=["#1F3B73"],
            hover_data=[
                "vms_prioritarias_anterior",
                "vms_prioritarias_atual",
                "delta_p0_acao_imediata",
                "delta_p1_alta",
                "delta_score_max",
            ],
        )

        fig_cluster_delta.update_layout(
            xaxis_title="Delta VMs prioritárias",
            yaxis_title="Cluster",
            yaxis={"categoryorder": "total ascending"},
        )

        st.plotly_chart(aplicar_layout_plotly(fig_cluster_delta, altura=540, showlegend=False), width="stretch")

    st.markdown('<div class="section-title">VMs que pioraram</div>', unsafe_allow_html=True)

    df_pioraram = df_comp_vms[
        df_comp_vms["mudanca_prioridade"].isin(["PIOROU", "PIOROU_MUITO"])
    ].copy()

    colunas_pioraram = [
        "cluster",
        "vm",
        "situacao_vm",
        "prioridade_final_anterior",
        "prioridade_final_atual",
        "mudanca_prioridade",
        "alerta_comparativo",
        "score_prioridade_anterior",
        "score_prioridade_atual",
        "delta_score_prioridade",
        "disk_p95_pct_anterior",
        "disk_p95_pct_atual",
        "delta_disk_p95_pct",
        "mem_p95_pct_anterior",
        "mem_p95_pct_atual",
        "delta_mem_p95_pct",
        "cpu_p95_pct_anterior",
        "cpu_p95_pct_atual",
        "delta_cpu_p95_pct",
    ]

    colunas_pioraram = [c for c in colunas_pioraram if c in df_pioraram.columns]

    if df_pioraram.empty:
        st.info("Nenhuma VM piorou entre as execuções selecionadas.")
    else:
        st.dataframe(
            formatar_dataframe_visual_2_casas(df_pioraram[colunas_pioraram].sort_values(
                ["delta_score_prioridade", "delta_disk_p95_pct", "delta_mem_p95_pct"],
                ascending=False,
            )),
            width="stretch",
            hide_index=True,
        )

    st.markdown('<div class="section-title">VMs novas</div>', unsafe_allow_html=True)

    df_novas = df_comp_vms[df_comp_vms["situacao_vm"] == "NOVA"].copy()

    colunas_novas = [
        "cluster",
        "vm",
        "prioridade_final_atual",
        "status_geral_atual",
        "risco_futuro_90d_atual",
        "score_prioridade_atual",
        "acao_final_atual",
        "disk_p95_pct_atual",
        "mem_p95_pct_atual",
        "cpu_p95_pct_atual",
    ]

    colunas_novas = [c for c in colunas_novas if c in df_novas.columns]

    if df_novas.empty:
        st.info("Nenhuma VM nova entre as execuções selecionadas.")
    else:
        st.dataframe(
            formatar_dataframe_visual_2_casas(df_novas[colunas_novas].sort_values(
                "score_prioridade_atual",
                ascending=False,
            )),
            width="stretch",
            hide_index=True,
        )

    st.markdown('<div class="section-title">VMs removidas</div>', unsafe_allow_html=True)

    df_removidas = df_comp_vms[df_comp_vms["situacao_vm"] == "REMOVIDA"].copy()

    colunas_removidas = [
        "cluster",
        "vm",
        "prioridade_final_anterior",
        "status_geral_anterior",
        "risco_futuro_90d_anterior",
        "score_prioridade_anterior",
        "acao_final_anterior",
    ]

    colunas_removidas = [c for c in colunas_removidas if c in df_removidas.columns]

    if df_removidas.empty:
        st.info("Nenhuma VM removida entre as execuções selecionadas.")
    else:
        st.dataframe(
            formatar_dataframe_visual_2_casas(df_removidas[colunas_removidas].sort_values(
                "score_prioridade_anterior",
                ascending=False,
            )),
            width="stretch",
            hide_index=True,
        )

    st.markdown('<div class="section-title">Comparação por cluster</div>', unsafe_allow_html=True)

    st.dataframe(
        formatar_dataframe_visual_2_casas(df_comp_cluster),
        width="stretch",
        hide_index=True,
    )

    csv_comp = df_comp_vms.to_csv(index=False).encode("utf-8")

    st.download_button(
        label="Baixar comparação de VMs em CSV",
        data=csv_comp,
        file_name="rmc_copilot_comparacao_vms.csv",
        mime="text/csv",
    )


# =============================================================================
# Fonte de dados
# =============================================================================
def carregar_dados_dashboard():
    st.sidebar.markdown("### Fonte de dados")

    st.session_state.setdefault("df_historico_metricas", pd.DataFrame())
    st.session_state.setdefault("execution_id_ativo", None)
    st.session_state.setdefault("resource_run_id_ativo", None)
    st.session_state.setdefault("resource_run_label_ativo", None)
    st.session_state.setdefault("historico_dias_ativo", 90)
    st.session_state.setdefault("forecast_horizons_ativo", [30, 60, 90])
    st.session_state.setdefault("coleta_manual_log_path", None)

    modo = st.sidebar.radio(
        "Como carregar os dados?",
        [
            "Ler última execução do banco",
            "Selecionar execução do banco",
            "Usar arquivo padrão",
            "Enviar arquivo Excel",
        ],
    )

    arquivo_processar = None
    df_analise_v4 = None


    with st.sidebar.expander("Coleta vROps / histórico DuckDB", expanded=True):
        df_resource_runs = listar_resource_collection_runs_dashboard(CAMINHO_BANCO)

        if df_resource_runs is None or df_resource_runs.empty:
            st.caption("Nenhuma coleta encontrada em resource_collection_runs.")
            st.session_state["resource_run_id_ativo"] = None
            st.session_state["resource_run_label_ativo"] = None
        else:
            opcoes_runs = df_resource_runs["label"].tolist()

            run_atual = st.session_state.get("resource_run_id_ativo")
            idx_run = 0
            if run_atual:
                matches = df_resource_runs.index[df_resource_runs["run_id"].astype(str) == str(run_atual)].tolist()
                if matches:
                    idx_run = int(matches[0])

            label_run = st.selectbox(
                "Coleta para histórico/partições/alocação",
                options=opcoes_runs,
                index=idx_run,
                key="select_resource_run_ativo_15f9",
            )

            linha_run = df_resource_runs[df_resource_runs["label"] == label_run].iloc[0]
            st.session_state["resource_run_id_ativo"] = str(linha_run["run_id"])
            st.session_state["resource_run_label_ativo"] = str(label_run)

            st.caption(
                "Run ativo: "
                + str(linha_run["run_id"])
                + " | período="
                + str(linha_run.get("primeira_data", ""))
                + " a "
                + str(linha_run.get("ultima_data", ""))
                + " | linhas="
                + str(linha_run.get("linhas", ""))
            )

        st.markdown("---")
        st.markdown("**Coleta manual vROps**")

        historico_dias_manual = st.number_input(
            "Prazo de histórico para coletar/analisar (dias)",
            min_value=1,
            max_value=730,
            value=int(st.session_state.get("historico_dias_ativo", 90) or 90),
            step=1,
            key="manual_history_days_15f10",
        )
        st.session_state["historico_dias_ativo"] = int(historico_dias_manual)

        forecast_txt = st.text_input(
            "Prazos de forecast (dias, separados por vírgula)",
            value=",".join([str(x) for x in st.session_state.get("forecast_horizons_ativo", [30, 60, 90])]),
            key="manual_forecast_horizons_15f10",
        )
        try:
            forecast_horizons = sorted(
                {
                    int(x.strip())
                    for x in str(forecast_txt).split(",")
                    if x.strip() and int(x.strip()) > 0
                }
            )
        except Exception:
            forecast_horizons = [30, 60, 90]
            st.caption("Forecast inválido; usando 30,60,90.")
        if not forecast_horizons:
            forecast_horizons = [30, 60, 90]
        st.session_state["forecast_horizons_ativo"] = forecast_horizons

        tipo_coleta_manual = st.selectbox(
            "Escopo da coleta manual",
            options=["Cluster", "Host", "VM", "Todos os clusters esperados"],
            index=0,
            key="manual_scope_type_15f10",
        )

        nome_coleta_manual = st.text_input(
            "Nome do cluster/host/VM",
            value="",
            key="manual_scope_name_15f10",
            help="Obrigatório para Cluster, Host e VM. Em 'Todos os clusters esperados', pode ficar vazio.",
        )

        max_vms_manual = st.number_input(
            "Limite de VMs nesta coleta",
            min_value=1,
            max_value=10000,
            value=500,
            step=50,
            key="manual_max_vms_15f10",
        )

        with st.expander("Credenciais vROps para coleta manual", expanded=False):
            vrops_host_manual = st.text_input(
                "vROps host",
                value=os.environ.get("RMC_VROPS_HOST", "mor-vropsprd01.bvnet.bv"),
                key="manual_vrops_host_15f10",
            )
            auth_source_manual = st.text_input(
                "AuthSource",
                value=os.environ.get("RMC_VROPS_AUTH_SOURCE", "bvnet.bv"),
                key="manual_vrops_auth_source_15f10",
            )
            usuario_manual = st.text_input(
                "Usuário vROps",
                value=os.environ.get("RMC_VROPS_USERNAME", ""),
                key="manual_vrops_user_15f10",
            )
            senha_manual = st.text_input(
                "Senha vROps (não é gravada)",
                type="password",
                key="manual_vrops_password_15f10",
            )

        iniciar_coleta = st.button(
            "Iniciar coleta manual agora",
            type="primary",
            key="btn_iniciar_coleta_manual_15f10",
        )

        if iniciar_coleta:
            if tipo_coleta_manual != "Todos os clusters esperados" and not str(nome_coleta_manual).strip():
                st.error("Informe o nome do cluster, host ou VM para iniciar a coleta.")
            elif not str(usuario_manual).strip():
                st.error("Informe o usuário vROps.")
            elif not str(senha_manual).strip():
                st.error("Informe a senha vROps. Ela será passada somente por variável de ambiente para o processo.")
            else:
                scope_map = {
                    "Cluster": "cluster",
                    "Host": "host",
                    "VM": "vm",
                    "Todos os clusters esperados": "all",
                }

                log_dir = PROJECT_ROOT / "data/logs/coletas"
                log_dir.mkdir(parents=True, exist_ok=True)
                log_path = log_dir / ("coleta_manual_streamlit_" + pd.Timestamp.now().strftime("%Y%m%d_%H%M%S") + ".log")

                cmd = [
                    sys.executable,
                    str(PROJECT_ROOT / "scripts" / "100_coleta_manual_vrops.py"),
                    "--scope-type",
                    scope_map[tipo_coleta_manual],
                    "--scope-name",
                    str(nome_coleta_manual).strip(),
                    "--history-days",
                    str(int(historico_dias_manual)),
                    "--forecast-days",
                    ",".join([str(x) for x in forecast_horizons]),
                    "--max-vms",
                    str(int(max_vms_manual)),
                    "--db",
                    str(CAMINHO_BANCO),
                    "--host",
                    str(vrops_host_manual).strip(),
                    "--auth-source",
                    str(auth_source_manual).strip(),
                    "--username",
                    str(usuario_manual).strip(),
                ]

                env = os.environ.copy()
                env["RMC_VROPS_PASSWORD"] = str(senha_manual)
                env["RMC_VROPS_HOST"] = str(vrops_host_manual).strip()
                env["RMC_VROPS_AUTH_SOURCE"] = str(auth_source_manual).strip()
                env["RMC_VROPS_USERNAME"] = str(usuario_manual).strip()

                with open(log_path, "w", encoding="utf-8") as fh:
                    proc = subprocess.Popen(
                        cmd,
                        cwd=str(PROJECT_ROOT),
                        stdout=fh,
                        stderr=subprocess.STDOUT,
                        env=env,
                    )

                st.session_state["coleta_manual_log_path"] = str(log_path)
                st.success(f"Coleta manual iniciada em segundo plano. PID={proc.pid}")
                st.code(str(log_path), language="text")
                st.caption("Após a coleta finalizar, use o validador 90 ou atualize o dashboard para selecionar o novo run_id.")

        ultimo_log_coleta = st.session_state.get("coleta_manual_log_path")
        if ultimo_log_coleta:
            st.caption("Último log de coleta manual: " + str(ultimo_log_coleta))
            if Path(ultimo_log_coleta).exists():
                try:
                    tail = Path(ultimo_log_coleta).read_text(encoding="utf-8", errors="ignore").splitlines()[-20:]
                    with st.expander("Últimas linhas do log de coleta", expanded=False):
                        st.code("\n".join(tail), language="text")
                except Exception:
                    pass


    if modo == "Ler última execução do banco":
        ultima_execucao = obter_ultima_execucao(CAMINHO_BANCO)

        if ultima_execucao is None:
            st.sidebar.warning("Nenhuma execução encontrada no banco. Use Excel para processar a primeira carga.")
        else:
            st.sidebar.success(f"Última execução carregada:\n\n{ultima_execucao}")
            df_analise_v4, _ = carregar_execucao_banco(str(CAMINHO_BANCO), ultima_execucao)
            st.session_state["execution_id_ativo"] = ultima_execucao
            st.session_state["df_historico_metricas"] = carregar_historico_banco(str(CAMINHO_BANCO), ultima_execucao)

    elif modo == "Selecionar execução do banco":
        execucoes = listar_execucoes(CAMINHO_BANCO)

        if execucoes.empty:
            st.sidebar.warning("Nenhuma execução encontrada no banco.")
        else:
            opcoes = execucoes["execution_id"].tolist()

            execution_id_selecionado = st.sidebar.selectbox(
                "Execução",
                options=opcoes,
            )

            df_analise_v4, _ = carregar_execucao_banco(
                str(CAMINHO_BANCO),
                execution_id_selecionado,
            )
            st.session_state["execution_id_ativo"] = execution_id_selecionado
            st.session_state["df_historico_metricas"] = carregar_historico_banco(str(CAMINHO_BANCO), execution_id_selecionado)

    elif modo == "Usar arquivo padrão":
        if ARQUIVO_PADRAO.exists():
            arquivo_processar = ARQUIVO_PADRAO
            st.sidebar.success("Arquivo padrão encontrado.")
        else:
            st.sidebar.error(f"Arquivo padrão não encontrado: {ARQUIVO_PADRAO}")

    elif modo == "Enviar arquivo Excel":
        upload = st.sidebar.file_uploader(
            "Envie a saída Excel do RMC/vROps",
            type=["xlsx"],
        )

        if upload is not None:
            temp_dir = PROJECT_ROOT / "data/raw/uploads"
            temp_dir.mkdir(parents=True, exist_ok=True)

            arquivo_temp = temp_dir / upload.name

            with open(arquivo_temp, "wb") as f:
                f.write(upload.getbuffer())

            arquivo_processar = arquivo_temp
            st.sidebar.success("Arquivo enviado com sucesso.")

    if df_analise_v4 is None:
        if arquivo_processar is None:
            st.info("Selecione uma execução do banco ou envie um arquivo Excel para iniciar.")
            st.stop()

        with st.spinner("Processando dados do RMC Copilot..."):
            df_analise_v4 = processar_excel(str(arquivo_processar))
            st.session_state["df_historico_metricas"] = carregar_historico_excel(str(arquivo_processar))
            st.session_state["execution_id_ativo"] = None

        resumo_para_salvar = criar_resumo_cluster(df_analise_v4)

        salvar_no_banco = st.sidebar.checkbox(
            "Salvar execução no DuckDB",
            value=True,
        )

        observacao_execucao = st.sidebar.text_input(
            "Observação da execução",
            value="Carga via dashboard",
        )

        if salvar_no_banco:
            execution_id_salvo = salvar_execucao(
                caminho_banco=CAMINHO_BANCO,
                df_analise_v4=df_analise_v4,
                resumo_cluster_v04=resumo_para_salvar,
                nome_arquivo=str(arquivo_processar.name),
                observacao=observacao_execucao,
            )
            st.session_state["execution_id_ativo"] = execution_id_salvo
            st.sidebar.success(f"Execução salva:\n\n{execution_id_salvo}")

    return preparar_coluna_host(df_analise_v4)


# =============================================================================
# Página dashboard operacional
# =============================================================================
def render_dashboard_operacional():
    df_analise_v4 = carregar_dados_dashboard()

    df_filtrado, recurso_selecionado = filtrar_dataframe(df_analise_v4)
    df_filtrado = preparar_coluna_host(df_filtrado)

    if df_filtrado.empty:
        st.warning("Nenhum dado encontrado com os filtros selecionados.")
        st.stop()

    resumo_cluster = criar_resumo_cluster(df_filtrado)

    total = len(df_filtrado)
    prioridade_counts = df_filtrado["prioridade_final"].value_counts()
    criticidade_counts = df_filtrado["criticidade_futura"].value_counts()

    p0 = int(prioridade_counts.get("P0_ACAO_IMEDIATA", 0))
    p1 = int(prioridade_counts.get("P1_ALTA", 0))
    p2 = int(prioridade_counts.get("P2_MEDIA", 0))
    p3 = int(prioridade_counts.get("P3_BAIXA", 0))
    p4 = int(prioridade_counts.get("P4_MONITORAR", 0))

    risco_30 = int(criticidade_counts.get("RISCO_FUTURO_30D", 0))
    risco_60 = int(criticidade_counts.get("RISCO_FUTURO_60D", 0))
    risco_90 = int(criticidade_counts.get("RISCO_FUTURO_90D", 0))

    st.markdown('<div class="section-title">Visão geral</div>', unsafe_allow_html=True)

    c1, c2, c3, c4, c5, c6 = st.columns(6)

    with c1:
        render_kpi_card("Total de VMs", formatar_numero(total), "kpi-blue", "Ambiente analisado")
    with c2:
        render_kpi_card("P0 Imediata", formatar_numero(p0), "kpi-red", "Ação imediata")
    with c3:
        render_kpi_card("P1 Alta", formatar_numero(p1), "kpi-orange", "Curto prazo")
    with c4:
        render_kpi_card("P2 Média", formatar_numero(p2), "kpi-yellow", "Acompanhar")
    with c5:
        render_kpi_card("P3 Baixa", formatar_numero(p3), "kpi-cyan", "Otimização")
    with c6:
        render_kpi_card("P4 Monitorar", formatar_numero(p4), "kpi-green", "Baixa criticidade")

    st.markdown('<div class="section-title">Forecast de risco</div>', unsafe_allow_html=True)

    f1, f2, f3 = st.columns(3)

    with f1:
        render_kpi_card("Risco 30 dias", formatar_numero(risco_30), "kpi-red", "Atenção imediata")
    with f2:
        render_kpi_card("Risco 60 dias", formatar_numero(risco_60), "kpi-orange", "Planejamento")
    with f3:
        render_kpi_card("Risco 90 dias", formatar_numero(risco_90), "kpi-yellow", "Radar futuro")

    st.markdown('<div class="section-title">Seleção de escopo para análise</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="section-subtitle">Escolha se quer analisar por cluster, host ou VM. Host só aparece quando a coleta trouxer coluna explícita de host/ESXi. O gráfico abaixo mostra o escopo escolhido com histórico real e previsão de 30/60/90 dias.</div>',
        unsafe_allow_html=True,
    )

    host_disponivel = host_esta_disponivel(df_filtrado)
    opcoes_escopo = ["Cluster"]
    if host_disponivel:
        opcoes_escopo.append("Host")
    opcoes_escopo.append("VM")

    tipo_drilldown = st.radio(
        "Nível de análise",
        options=opcoes_escopo,
        horizontal=True,
        index=0,
    )

    if not host_disponivel:
        st.caption("Host não será exibido porque a coleta atual não trouxe uma coluna explícita de host/ESXi. Isso evita usar mapeamento incorreto como se fosse host.")

    cluster_selecionado = None
    host_selecionado = None
    vm_selecionada = None
    df_scope = df_filtrado.copy()
    tipo_escopo = "Filtro atual"
    nome_escopo = "Todos os itens filtrados"

    if tipo_drilldown == "Cluster":
        st.markdown("#### Escolha o cluster no grid")
        col_cluster = [
            "cluster",
            "total_vms",
            "p0_acao_imediata",
            "p1_alta",
            "vms_prioritarias",
            "pct_vms_prioritarias",
            "score_max",
            "cpu_p95_medio",
            "mem_p95_medio",
            "disk_p95_medio",
        ]
        col_cluster = [c for c in col_cluster if c in resumo_cluster.columns]
        resumo_cluster_grid = resumo_cluster[col_cluster].copy()
        evento_cluster = st.dataframe(
            formatar_dataframe_visual_2_casas(renomear_colunas_visual(resumo_cluster_grid)),
            width="stretch",
            hide_index=True,
            on_select="rerun",
            selection_mode="single-row",
            key="grid_cluster_escopo_v2",
        )
        idx_cluster = obter_linha_grid(evento_cluster)
        if idx_cluster is not None and idx_cluster < len(resumo_cluster_grid):
            cluster_selecionado = resumo_cluster_grid.iloc[idx_cluster]["cluster"]
            df_scope = df_filtrado[df_filtrado["cluster"].astype(str) == str(cluster_selecionado)].copy()
            tipo_escopo = "Cluster"
            nome_escopo = str(cluster_selecionado)
        else:
            st.info("Clique em um cluster para ver recursos totais, Top 30 VMs, histórico e forecast desse cluster.")

    elif tipo_drilldown == "Host":
        st.markdown("#### Escolha o host no grid")
        col_host_original = obter_coluna_host_explicita(df_filtrado)
        df_host_base = df_filtrado.copy()
        df_host_base["host_visual"] = df_host_base[col_host_original].fillna("Host não informado").astype(str)
        resumo_host = criar_resumo_host(df_host_base)
        col_host = [
            "host",
            "total_vms",
            "p0_acao_imediata",
            "p1_alta",
            "vms_prioritarias",
            "score_max",
            "cpu_pico_uso",
            "memoria_pico_uso",
            "disco_pico_uso",
        ]
        col_host = [c for c in col_host if c in resumo_host.columns]
        resumo_host_grid = resumo_host[col_host].copy()
        evento_host = st.dataframe(
            formatar_dataframe_visual_2_casas(renomear_colunas_visual(resumo_host_grid)),
            width="stretch",
            hide_index=True,
            on_select="rerun",
            selection_mode="single-row",
            key="grid_host_escopo_v2",
        )
        idx_host = obter_linha_grid(evento_host)
        if idx_host is not None and idx_host < len(resumo_host_grid):
            host_selecionado = resumo_host_grid.iloc[idx_host]["host"]
            df_scope = df_host_base[df_host_base["host_visual"].astype(str) == str(host_selecionado)].copy()
            tipo_escopo = "Host"
            nome_escopo = str(host_selecionado)
        else:
            st.info("Clique em um host para ver recursos totais, Top 30 VMs, histórico e forecast desse host.")

    elif tipo_drilldown == "VM":
        st.markdown("#### Top 30 VMs — clique em uma VM")
        vm_selecionada, df_top30_vm_global = render_top30_vms_clicavel(
            df_scope=df_filtrado,
            recurso_selecionado=recurso_selecionado,
            key="grid_top30_vm_global_v2",
        )
        render_resumo_recursos(df_top30_vm_global, "Recursos das Top 30 VMs listadas", caminho_banco=CAMINHO_BANCO, tipo_escopo="VM", nome_escopo="Top 30", run_id=st.session_state.get("resource_run_id_ativo"))
        if vm_selecionada is not None and not vm_selecionada.empty:
            df_scope = vm_selecionada.copy()
            tipo_escopo = "VM"
            nome_escopo = str(vm_selecionada.iloc[0]["vm"])
            render_detalhe_vm_recursos(vm_selecionada)
        else:
            st.info("Clique em uma VM da Top 30 para abrir os dados de recursos, ocupação, histórico e forecast.")

    if tipo_drilldown in ["Cluster", "Host"] and not df_scope.empty and tipo_escopo in ["Cluster", "Host"]:
        render_resumo_recursos(df_scope, f"Recursos totais — {tipo_escopo}: {nome_escopo}", caminho_banco=CAMINHO_BANCO, tipo_escopo=tipo_escopo, nome_escopo=nome_escopo, run_id=st.session_state.get("resource_run_id_ativo"))
        render_datastores(df_scope, caminho_banco=CAMINHO_BANCO, tipo_escopo=tipo_escopo, nome_escopo=nome_escopo, run_id=st.session_state.get("resource_run_id_ativo"))

        st.markdown(f"#### Top 30 VMs dentro do escopo selecionado — {tipo_escopo}: {nome_escopo}")
        vm_top30, df_top30 = render_top30_vms_clicavel(
            df_scope=df_scope,
            recurso_selecionado=recurso_selecionado,
            key=f"grid_top30_{tipo_escopo.lower()}_v2",
        )
        render_resumo_recursos(df_top30, "Recursos das Top 30 VMs listadas", caminho_banco=CAMINHO_BANCO, tipo_escopo="VM", nome_escopo="Top 30", run_id=st.session_state.get("resource_run_id_ativo"))

        if vm_top30 is not None and not vm_top30.empty:
            st.success(f"VM selecionada na Top 30: {vm_top30.iloc[0]['vm']}")
            render_detalhe_vm_recursos(vm_top30)
            df_scope_grafico = vm_top30.copy()
            tipo_grafico = "VM"
            nome_grafico = str(vm_top30.iloc[0]["vm"])
        else:
            df_scope_grafico = df_scope.copy()
            tipo_grafico = tipo_escopo
            nome_grafico = nome_escopo
    else:
        df_scope_grafico = df_scope.copy()
        tipo_grafico = tipo_escopo
        nome_grafico = nome_escopo
        if tipo_escopo == "VM":
            render_datastores(df_scope_grafico, caminho_banco=CAMINHO_BANCO, tipo_escopo=tipo_grafico, nome_escopo=nome_grafico, run_id=execution_id_grafico)

    st.markdown(f'<div class="section-title">Histórico real + previsão 30/60/90 dias — {tipo_grafico}: {nome_grafico}</div>', unsafe_allow_html=True)

    # HOTFIX 15F.6.7: PASSO 1 - histórico real direto do DuckDB com fallback por VM/run.
    # Não usa df_historico_metricas para decidir se há histórico; consulta vm_resource_timeseries.
    execution_id_grafico = st.session_state.get("resource_run_id_ativo") or st.session_state.get("execution_id_ativo")
    fig_escopo = criar_grafico_historico_forecast_escopo_duckdb(
        df_scope=df_scope_grafico,
        df_hist=None,
        tipo_escopo=tipo_grafico,
        nome_escopo=nome_grafico,
        recurso_selecionado=recurso_selecionado,
        caminho_banco=CAMINHO_BANCO,
        run_id=execution_id_grafico,
        dias=int(st.session_state.get("historico_dias_ativo", 90) or 90),
    )

    if fig_escopo is not None:
        st.plotly_chart(fig_escopo, width="stretch")
        st.caption(f"Histórico real DuckDB | run_id_coleta={execution_id_grafico or 'última coleta OK'} | banco={CAMINHO_BANCO}")
        tabela_escopo = criar_tabela_forecast_escopo(df_scope_grafico, recurso_selecionado)
        if not tabela_escopo.empty:
            st.dataframe(formatar_dataframe_visual_2_casas(tabela_escopo), width="stretch", hide_index=True)
    else:
        st.warning("Não há histórico real/forecast suficiente para montar o gráfico deste escopo.")


    st.markdown(f'<div class="section-title">Histórico e forecast das partições de disco — {tipo_grafico}: {nome_grafico}</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="section-subtitle">A visão abaixo analisa DISCO por partição/volume, com forecast 30/60/90 dias para gerar FUP. Para VM, mostra as partições da VM selecionada. Para Cluster/Host, mostra as partições mais críticas do escopo.</div>',
        unsafe_allow_html=True,
    )

    col_part_1, col_part_2 = st.columns(2)
    with col_part_1:
        dias_particoes = st.slider(
            "Período do histórico de partições (dias)",
            min_value=7,
            max_value=365,
            value=int(st.session_state.get("historico_dias_ativo", 90) or 90),
            step=7,
            key=f"slider_dias_particoes_{tipo_grafico}_{nome_grafico}",
        )
    with col_part_2:
        top_particoes = st.slider(
            "Quantidade máxima de partições no gráfico",
            min_value=3,
            max_value=30,
            value=10,
            step=1,
            key=f"slider_top_particoes_{tipo_grafico}_{nome_grafico}",
        )

    fig_particoes, df_particoes, msg_particoes = criar_grafico_historico_particoes_duckdb(
        df_scope=df_scope_grafico,
        tipo_escopo=tipo_grafico,
        nome_escopo=nome_grafico,
        caminho_banco=CAMINHO_BANCO,
        run_id=execution_id_grafico,
        dias=dias_particoes,
        top_n=top_particoes,
        forecast_horizons=st.session_state.get("forecast_horizons_ativo", [30, 60, 90]),
    )

    if msg_particoes:
        st.caption(msg_particoes)

    if fig_particoes is not None:
        st.plotly_chart(fig_particoes, width="stretch")
    else:
        st.info("Não há histórico de partições de disco para montar o gráfico deste escopo.")

    if df_particoes is not None and not df_particoes.empty:
        st.dataframe(
            formatar_dataframe_visual_2_casas(renomear_colunas_visual(df_particoes)),
            width="stretch",
            hide_index=True,
        )


    st.markdown('<div class="section-title">Resumo executivo automático</div>', unsafe_allow_html=True)

    resumo_textual = gerar_resumo_textual(df_filtrado, resumo_cluster)

    st.markdown(
        f"""
        <div class="summary-box">
        {resumo_textual}
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown('<div class="section-title">Visão gráfica</div>', unsafe_allow_html=True)

    g1, g2 = st.columns(2)

    with g1:
        st.plotly_chart(criar_grafico_prioridade(df_filtrado), width="stretch")

    with g2:
        st.plotly_chart(criar_grafico_risco_futuro(df_filtrado), width="stretch")

    g3, g4 = st.columns(2)

    with g3:
        st.plotly_chart(criar_grafico_top_clusters(resumo_cluster), width="stretch")

    with g4:
        st.plotly_chart(criar_grafico_acao(df_filtrado), width="stretch")

    fig_recursos = criar_grafico_recursos_cluster(resumo_cluster, recurso_selecionado)
    if fig_recursos is not None:
        st.plotly_chart(fig_recursos, width="stretch")

    st.plotly_chart(criar_grafico_top_vms(df_filtrado, recurso_selecionado), width="stretch")

    st.markdown('<div class="section-title">Exportação opcional</div>', unsafe_allow_html=True)

    csv = df_filtrado.to_csv(index=False).encode("utf-8")

    st.download_button(
        label="Baixar dados filtrados em CSV",
        data=csv,
        file_name="rmc_copilot_dados_filtrados.csv",
        mime="text/csv",
    )


# =============================================================================
# Main
# =============================================================================
def render_pagina_particoes_criticas():
    df_analise_v4 = carregar_dados_dashboard()
    df_filtrado, _ = filtrar_dataframe(df_analise_v4)
    df_filtrado = preparar_coluna_host(df_filtrado)

    if df_filtrado.empty:
        st.warning("Nenhum dado encontrado com os filtros selecionados.")
        st.stop()

    st.markdown('<div class="section-title">Partições críticas acima de 80%</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="section-subtitle">Atalho de FUP: lista todas as partições com uso atual ou forecast 30/60/90 dias acima do limite informado.</div>',
        unsafe_allow_html=True,
    )

    c1, c2, c3 = st.columns([1, 1, 2])
    with c1:
        limite = st.slider("Limite crítico (%)", min_value=70, max_value=99, value=80, step=1, key="particoes_limite_15f10_17")
    with c2:
        top_n = st.slider("Máx. partições consultadas", min_value=50, max_value=5000, value=1000, step=50, key="particoes_topn_15f10_17")
    with c3:
        st.caption("Dica: use os filtros da barra lateral para restringir cluster, categoria, prioridade, ação ou VM antes da consulta.")

    fig_particoes, df_particoes, msg_particoes = criar_grafico_historico_particoes_duckdb(
        df_scope=df_filtrado,
        tipo_escopo="Global",
        nome_escopo="Filtro atual",
        caminho_banco=CAMINHO_BANCO,
        run_id=st.session_state.get("resource_run_id_ativo"),
        dias=int(st.session_state.get("historico_dias_ativo", 90) or 90),
        top_n=int(top_n),
        forecast_horizons=st.session_state.get("forecast_horizons_ativo", [30, 60, 90]),
    )

    if msg_particoes:
        st.caption(msg_particoes)

    if df_particoes is None or df_particoes.empty:
        st.info("Nenhuma partição encontrada para o escopo filtrado.")
        return

    df_particoes = df_particoes.copy()
    for col in ["last_used_pct", "max_used_pct", "avg_used_pct", "forecast_30d", "forecast_60d", "forecast_90d", "slope_pp_day"]:
        if col in df_particoes.columns:
            df_particoes[col] = pd.to_numeric(df_particoes[col], errors="coerce")

    colunas_alerta = [c for c in ["last_used_pct", "forecast_30d", "forecast_60d", "forecast_90d"] if c in df_particoes.columns]
    if not colunas_alerta:
        st.warning("Resumo de partições retornou sem colunas de uso/forecast.")
        return

    mascara = pd.Series(False, index=df_particoes.index)
    for col in colunas_alerta:
        mascara = mascara | (df_particoes[col].fillna(0) >= float(limite))

    criticas = df_particoes[mascara].copy()
    if criticas.empty:
        st.success(f"Nenhuma partição acima de {limite}% no uso atual ou nos forecasts 30/60/90 dias.")
        return

    criticas["maior_pct"] = criticas[colunas_alerta].max(axis=1, skipna=True)
    criticas = criticas.sort_values(["maior_pct", "last_used_pct"], ascending=False, na_position="last")

    k1, k2, k3, k4 = st.columns(4)
    with k1:
        render_kpi_card("Partições críticas", _fmt_card_count(len(criticas)), "kpi-red", f">= {limite}%")
    with k2:
        render_kpi_card("VMs afetadas", _fmt_card_count(criticas["vm"].nunique() if "vm" in criticas.columns else 0), "kpi-blue", "Qtd.")
    with k3:
        atuais = int((criticas["last_used_pct"].fillna(0) >= float(limite)).sum()) if "last_used_pct" in criticas.columns else 0
        render_kpi_card("Uso atual", _fmt_card_count(atuais), "kpi-orange", ">= limite")
    with k4:
        em_30 = int((criticas["forecast_30d"].fillna(0) >= float(limite)).sum()) if "forecast_30d" in criticas.columns else 0
        render_kpi_card("Forecast 30d", _fmt_card_count(em_30), "kpi-yellow", ">= limite")

    if fig_particoes is not None:
        st.plotly_chart(fig_particoes, width="stretch")

    colunas_exibir = [
        c for c in [
            "cluster",
            "host",
            "vm",
            "partition_name",
            "last_used_pct",
            "max_used_pct",
            "avg_used_pct",
            "forecast_30d",
            "forecast_60d",
            "forecast_90d",
            "dias_ate_85",
            "dias_ate_95",
            "status_fup",
            "acao_fup",
            "evidencia_fup",
        ] if c in criticas.columns
    ]

    st.dataframe(
        formatar_dataframe_visual_2_casas(renomear_colunas_visual(criticas[colunas_exibir])),
        width="stretch",
        hide_index=True,
    )


def render_placeholder_integrado(titulo: str, mensagem: str, pagina_destino: str = "dashboard"):
    st.info(mensagem)
    if pagina_destino == "particoes":
        render_pagina_particoes_criticas()
    else:
        render_dashboard_operacional()




# =============================================================================
# HOTFIX 15F.10.21 - helpers para Analise Individual e FUP refinado
# =============================================================================
def _pick_col_case_insensitive(cols, candidates):
    # HOTFIX 15F.10.28:
    # pandas.Index não pode ser usado em expressão booleana: "cols or []"
    # Isso quebrava a Análise Individual ao colar o nome do servidor.
    if cols is None:
        cols_list = []
    else:
        cols_list = list(cols)
    lower = {str(c).strip().lower(): c for c in cols_list}
    for cand in candidates:
        if cand in cols_list:
            return cand
        key = str(cand).strip().lower()
        if key in lower:
            return lower[key]
    return None


def _first_text_value(df, candidates, default="N/C"):
    if df is None or df.empty:
        return default
    col = _pick_col_case_insensitive(df.columns, candidates)
    if not col:
        return default
    vals = df[col].dropna().astype(str).str.strip()
    vals = vals[(vals != "") & (vals.str.upper() != "NAN")]
    if vals.empty:
        return default
    uniques = vals.unique().tolist()
    if len(uniques) <= 2:
        return ", ".join(uniques)
    return ", ".join(uniques[:2]) + f" +{len(uniques)-2}"


def _first_numeric_value(df, candidates):
    if df is None or df.empty:
        return None
    col = _pick_col_case_insensitive(df.columns, candidates)
    if not col:
        return None
    s = pd.to_numeric(df[col], errors="coerce").dropna()
    if s.empty:
        return None
    return float(s.iloc[0])


def _duck_table_exists(con, table):
    try:
        return bool(con.execute(
            "SELECT COUNT(*) FROM information_schema.tables WHERE lower(table_name)=lower(?)",
            [str(table)],
        ).fetchone()[0])
    except Exception:
        return False


def _duck_cols(con, table):
    try:
        return con.execute('DESCRIBE "' + str(table).replace('"', '""') + '"').fetchdf().iloc[:, 0].astype(str).tolist()
    except Exception:
        return []


def _qident(name):
    return '"' + str(name).replace('"', '""') + '"'


def _buscar_valor_vm_duckdb(nome_vm, table_candidates, value_candidates):
    if duckdb is None or not CAMINHO_BANCO.exists() or not nome_vm:
        return None
    con = None
    try:
        con = duckdb.connect(str(CAMINHO_BANCO), read_only=True)
        for table in table_candidates:
            if not _duck_table_exists(con, table):
                continue
            cols = _duck_cols(con, table)
            vm_col = _pick_col_case_insensitive(cols, ["vm", "vm_name", "name", "resource_name", "Nome VM"])
            value_col = _pick_col_case_insensitive(cols, value_candidates)
            if not vm_col or not value_col:
                continue
            sql = f"SELECT {_qident(value_col)} FROM {_qident(table)} WHERE UPPER(CAST({_qident(vm_col)} AS VARCHAR)) = ? AND {_qident(value_col)} IS NOT NULL ORDER BY 1 DESC LIMIT 1"
            row = con.execute(sql, [str(nome_vm).upper()]).fetchone()
            if row and row[0] is not None:
                con.close()
                return row[0]
        con.close()
    except Exception:
        try:
            if con is not None:
                con.close()
        except Exception:
            pass
    return None


def _buscar_datastores_vm_duckdb(nome_vm):
    if duckdb is None or not CAMINHO_BANCO.exists() or not nome_vm:
        return "N/C"
    tables = ["vm_disk_partitions", "vm_allocation_snapshots", "vm_inventory_snapshots", "vm_resource_timeseries"]
    con = None
    try:
        con = duckdb.connect(str(CAMINHO_BANCO), read_only=True)
        encontrados = []
        for table in tables:
            if not _duck_table_exists(con, table):
                continue
            cols = _duck_cols(con, table)
            vm_col = _pick_col_case_insensitive(cols, ["vm", "vm_name", "name", "resource_name", "Nome VM"])
            ds_col = _pick_col_case_insensitive(cols, ["datastore", "datastore_name", "ds_name", "data_store", "storage", "subresource"])
            if not vm_col or not ds_col:
                continue
            sql = f"SELECT DISTINCT CAST({_qident(ds_col)} AS VARCHAR) FROM {_qident(table)} WHERE UPPER(CAST({_qident(vm_col)} AS VARCHAR)) = ? AND {_qident(ds_col)} IS NOT NULL LIMIT 10"
            rows = con.execute(sql, [str(nome_vm).upper()]).fetchall()
            for r in rows:
                val = str(r[0]).strip()
                if val and val.upper() not in {"NAN", "NONE"} and val not in encontrados:
                    encontrados.append(val)
        con.close()
        if not encontrados:
            return "N/C"
        if len(encontrados) <= 2:
            return ", ".join(encontrados)
        return ", ".join(encontrados[:2]) + f" +{len(encontrados)-2}"
    except Exception:
        try:
            if con is not None:
                con.close()
        except Exception:
            pass
    return "N/C"


def _render_cards_vm_individual(vm_df: pd.DataFrame, df_contexto: pd.DataFrame, nome_vm: str):
    host = _first_text_value(vm_df, ["host", "Host", "host_name", "ESXi", "esxi", "nome_host", "mapping_parent_name"])
    datastore = _first_text_value(vm_df, ["datastore", "datastore_name", "ds_name", "storage"], default=None)
    if not datastore:
        datastore = _buscar_datastores_vm_duckdb(nome_vm)

    vcpu = _first_numeric_value(vm_df, ["vcpu_count", "vcpus", "vcpu", "num_cpu", "cpu_count", "vCPU"])
    if vcpu is None:
        vcpu = _buscar_valor_vm_duckdb(nome_vm, ["vm_allocation_snapshots", "vm_inventory_snapshots"], ["vcpu_count", "vcpus", "vcpu", "num_cpu", "cpu_count"])
        try:
            vcpu = float(vcpu) if vcpu is not None else None
        except Exception:
            vcpu = None

    mem_aloc = _first_numeric_value(vm_df, ["memory_gb", "mem_alloc_gb", "mem_gb", "memoria_gb", "ram_gb"])
    if mem_aloc is None:
        mem_aloc = _buscar_valor_vm_duckdb(nome_vm, ["vm_allocation_snapshots", "vm_inventory_snapshots"], ["memory_gb", "mem_alloc_gb", "mem_gb", "ram_gb"])
        try:
            mem_aloc = float(mem_aloc) if mem_aloc is not None else None
        except Exception:
            mem_aloc = None

    disco_prov = _first_numeric_value(vm_df, ["provisioned_gb", "disk_provisioned_gb", "disk_alloc_gb", "disco_gb", "capacity_gb"])
    if disco_prov is None:
        disco_prov = _buscar_valor_vm_duckdb(nome_vm, ["vm_allocation_snapshots", "vm_inventory_snapshots"], ["provisioned_gb", "disk_provisioned_gb", "disk_alloc_gb", "disk_gb"])
        try:
            disco_prov = float(disco_prov) if disco_prov is not None else None
        except Exception:
            disco_prov = None

    mem_pct = _first_numeric_value(vm_df, ["mem_p95_pct", "mem_usage_pct", "mem_used_pct", "Memória — pico típico de uso (%)"])
    cpu_pct = _first_numeric_value(vm_df, ["cpu_p95_pct", "cpu_usage_pct", "cpu_used_pct", "CPU — pico típico de uso (%)"])
    disk_pct = _first_numeric_value(vm_df, ["disk_p95_pct", "disk_usage_pct", "disk_used_pct", "Disco — pico típico de uso (%)"])

    mem_fisica = _first_numeric_value(vm_df, ["mem_used_gb", "memory_used_gb", "guest_memory_used_gb", "active_mem_gb", "consumed_memory_gb"])
    if mem_fisica is None and mem_aloc is not None and mem_pct is not None:
        mem_fisica = float(mem_aloc) * float(mem_pct) / 100.0

    pcpu_eq = None
    cluster_nome = _first_text_value(vm_df, ["cluster", "Cluster", "cluster_name"], default=None)
    if vcpu is not None and cluster_nome:
        try:
            df_cluster = df_contexto[df_contexto["cluster"].astype(str) == str(cluster_nome)].copy() if "cluster" in df_contexto.columns else pd.DataFrame()
            resumo_cluster = calcular_alocacoes_cards(
                df_scope=df_cluster,
                caminho_banco=CAMINHO_BANCO,
                tipo_escopo="Cluster",
                nome_escopo=cluster_nome,
                run_id=st.session_state.get("resource_run_id_ativo"),
            )
            ratio = resumo_cluster.get("cpu_over_ratio")
            if ratio:
                pcpu_eq = float(vcpu) / float(ratio)
        except Exception:
            pcpu_eq = None

    st.markdown(f'<div class="section-title">Recursos da VM — {nome_vm}</div>', unsafe_allow_html=True)
    cards = [
        ("Host", host, "kpi-green", "ESXi"),
        ("Datastore", datastore, "kpi-blue", "Uso VM"),
        ("vCPU", _fmt_card_count(vcpu), "kpi-cyan", "Alocada"),
        ("pCPU Eq.", _fmt_valor(pcpu_eq, vazio="N/C"), "kpi-green", "Aprox."),
        ("Mem Aloc.", _fmt_card_storage_gb(mem_aloc), "kpi-orange", "Alocada"),
        ("Mem Física", _fmt_card_storage_gb(mem_fisica), "kpi-green", "Uso aprox."),
        ("Disco Prov.", _fmt_card_storage_gb(disco_prov), "kpi-yellow", "Provisionado"),
        ("CPU P95", _fmt_card_pct(cpu_pct), "kpi-blue", "Uso"),
        ("Mem P95", _fmt_card_pct(mem_pct), "kpi-orange", "Uso"),
        ("Disco P95", _fmt_card_pct(disk_pct), "kpi-red", "Uso"),
    ]
    for i in range(0, len(cards), 5):
        cols = st.columns(5)
        for col, (label, value, css_class, foot) in zip(cols, cards[i:i+5]):
            with col:
                render_kpi_card(label, value, css_class, foot)


def _enriquecer_particoes_fup(fup_part: pd.DataFrame, df_contexto: pd.DataFrame) -> pd.DataFrame:
    if fup_part is None or fup_part.empty:
        return fup_part
    out = fup_part.copy()
    if "vm" in out.columns and df_contexto is not None and not df_contexto.empty and "vm" in df_contexto.columns:
        map_cols = [c for c in ["vm", "cluster", "host"] if c in df_contexto.columns]
        extra = df_contexto[map_cols].drop_duplicates("vm")
        for c in ["cluster", "host"]:
            if c in extra.columns and c not in out.columns:
                out = out.merge(extra[["vm", c]], on="vm", how="left")
    if "datastore" not in out.columns:
        out["datastore"] = "N/C"
        if "vm" in out.columns:
            cache = {}
            for vm in out["vm"].dropna().astype(str).unique().tolist():
                cache[vm] = _buscar_datastores_vm_duckdb(vm)
            out["datastore"] = out["vm"].astype(str).map(cache).fillna("N/C")
    return out

# =============================================================================
# HOTFIX 15F.10.20 — páginas do menu recompostas
# =============================================================================
def render_capacity_dashboard():
    """Página principal do ambiente de Capacity."""
    st.markdown('<div class="section-title">Capacity Dashboard — Capacity Geral</div>', unsafe_allow_html=True)
    st.caption(
        "Ambiente geral de capacidade: visão geral, forecast de risco, seleção de escopo, "
        "Top 30, recursos, histórico real, partições, resumo executivo e visão gráfica."
    )
    render_dashboard_operacional()


def render_analise_individual_recursos():
    """Análise individual por número de solicitação + nome de servidor/VM."""
    df_analise_v4 = carregar_dados_dashboard()
    df_filtrado, recurso_selecionado = filtrar_dataframe(df_analise_v4)
    df_filtrado = preparar_coluna_host(df_filtrado)

    st.markdown('<div class="section-title">Análise Individual de Recursos</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="section-subtitle">Diretiva: informar obrigatoriamente o número da solicitação e o servidor/VM. '
        'A análise deve mostrar cards da VM, histórico real, forecast, partições e relatório demonstrativo assistido por LLM.</div>',
        unsafe_allow_html=True,
    )

    col_req, col_dest, col_vm = st.columns([1, 1.3, 2])
    with col_req:
        numero_solicitacao = st.text_input(
            "Número da solicitação (obrigatório)",
            value=st.session_state.get("numero_solicitacao_analise_individual_15f28", ""),
            placeholder="Ex.: SOL1809645",
            key="numero_solicitacao_analise_individual_15f28",
        )
    with col_dest:
        solicitante_relatorio = st.text_input(
            "Solicitante / destinatário",
            value=st.session_state.get("solicitante_relatorio_analise_15f31", ""),
            placeholder="Nome de quem receberá o relatório",
            key="solicitante_relatorio_analise_15f31",
        )
    with col_vm:
        busca = st.text_input(
            "Nome do servidor / VM (obrigatório)",
            "",
            placeholder="Cole ou digite o nome da VM/servidor",
            key="busca_servidor_analise_individual_15f28",
        )

    numero_solicitacao = str(numero_solicitacao or "").strip()
    solicitante_relatorio = str(solicitante_relatorio or "").strip()
    busca = str(busca or "").strip()

    if not numero_solicitacao:
        st.warning("Informe o número da solicitação para gerar a análise e o relatório.")
        return

    if not busca:
        st.info("Digite parte do nome do servidor para consultar.")
        return

    st.session_state["solicitacao_analise_individual_ativa"] = numero_solicitacao

    termo = busca.upper()
    candidatos = df_filtrado[df_filtrado["vm"].astype(str).str.upper().str.contains(termo, na=False)].copy()

    if candidatos.empty:
        st.warning("Nenhuma VM encontrada para o termo informado.")
        return

    candidatos = ordenar_vms_por_recurso(candidatos, recurso_selecionado)

    st.markdown("#### VMs encontradas")
    colunas = obter_colunas_vm_recursos(candidatos, recurso_selecionado)
    evento = st.dataframe(
        formatar_dataframe_visual_2_casas(renomear_colunas_visual(candidatos[colunas].head(100))),
        width="stretch",
        hide_index=True,
        on_select="rerun",
        selection_mode="single-row",
        key="grid_analise_individual_servidor_15f28",
    )

    idx = obter_linha_grid(evento)
    if idx is None:
        vm_df = candidatos.iloc[[0]].copy()
        st.caption("Nenhuma linha selecionada; exibindo a primeira VM encontrada.")
    else:
        vm_df = candidatos.iloc[[idx]].copy()

    nome_vm = str(vm_df.iloc[0]["vm"])

    st.caption(f"Solicitação: {numero_solicitacao} | VM selecionada: {nome_vm}")

    _render_cards_vm_individual(vm_df, df_filtrado, nome_vm)
    render_detalhe_vm_recursos(vm_df)

    execution_id_grafico = st.session_state.get("resource_run_id_ativo") or st.session_state.get("execution_id_ativo")

    st.markdown(f'<div class="section-title">Histórico real — VM: {nome_vm}</div>', unsafe_allow_html=True)
    fig_escopo = criar_grafico_historico_forecast_escopo_duckdb(
        df_scope=vm_df,
        df_hist=None,
        tipo_escopo="VM",
        nome_escopo=nome_vm,
        recurso_selecionado=recurso_selecionado,
        caminho_banco=CAMINHO_BANCO,
        run_id=execution_id_grafico,
        dias=int(st.session_state.get("historico_dias_ativo", 90) or 90),
    )

    tabela_escopo = pd.DataFrame()
    if fig_escopo is not None:
        st.plotly_chart(fig_escopo, width="stretch")
        tabela_escopo = criar_tabela_forecast_escopo(vm_df, recurso_selecionado)
        if tabela_escopo is not None and not tabela_escopo.empty:
            try:
                tabela_escopo = _adicionar_threshold_fup(tabela_escopo, recurso_selecionado, 80.0)
            except Exception:
                pass
            st.dataframe(formatar_dataframe_visual_2_casas(renomear_colunas_visual(tabela_escopo)), width="stretch", hide_index=True)
            try:
                _render_download_csv(
                    "Baixar forecast da análise em CSV",
                    tabela_escopo,
                    f"analise_individual_forecast_{_normalizar_nome_arquivo(numero_solicitacao)}_{_normalizar_nome_arquivo(nome_vm)}.csv",
                    key=f"download_analise_individual_forecast_{_normalizar_nome_arquivo(numero_solicitacao)}_{_normalizar_nome_arquivo(nome_vm)}",
                )
            except Exception:
                pass
    else:
        st.warning("Não há histórico real/forecast suficiente para esta VM.")

    st.markdown(f'<div class="section-title">Histórico e forecast das partições de disco — VM: {nome_vm}</div>', unsafe_allow_html=True)
    fig_part, df_part, msg_part = criar_grafico_historico_particoes_duckdb(
        df_scope=vm_df,
        tipo_escopo="VM",
        nome_escopo=nome_vm,
        caminho_banco=CAMINHO_BANCO,
        run_id=execution_id_grafico,
        dias=int(st.session_state.get("historico_dias_ativo", 90) or 90),
        top_n=30,
        forecast_horizons=st.session_state.get("forecast_horizons_ativo", [30, 60, 90]),
    )
    if msg_part:
        st.caption(msg_part)
    if fig_part is not None:
        st.plotly_chart(fig_part, width="stretch")

    df_part_relatorio = pd.DataFrame()
    if df_part is not None and not df_part.empty:
        df_part = _enriquecer_particoes_fup(df_part, df_filtrado)
        try:
            df_part = _adicionar_threshold_fup(df_part, "Disco", 80.0)
        except Exception:
            pass
        df_part_relatorio = df_part.copy()
        st.dataframe(formatar_dataframe_visual_2_casas(renomear_colunas_visual(df_part)), width="stretch", hide_index=True)
        try:
            _render_download_csv(
                "Baixar partições da análise em CSV",
                df_part,
                f"analise_individual_particoes_{_normalizar_nome_arquivo(numero_solicitacao)}_{_normalizar_nome_arquivo(nome_vm)}.csv",
                key=f"download_analise_individual_particoes_{_normalizar_nome_arquivo(numero_solicitacao)}_{_normalizar_nome_arquivo(nome_vm)}",
            )
        except Exception:
            pass

    try:
        _render_analise_completa_vm(
            numero_solicitacao=numero_solicitacao,
            nome_vm=nome_vm,
            vm_df=vm_df,
            recurso_selecionado=recurso_selecionado,
            solicitante=solicitante_relatorio or "Solicitante",
        )
    except Exception as exc:
        st.warning(f"Não foi possível abrir a análise completa v1: {exc}")

    try:
        _render_relatorio_vm_fup(
            vm_nome=nome_vm,
            recurso=recurso_selecionado,
            origem=f"Análise Individual | Solicitação {numero_solicitacao}",
            limite=80.0,
            vm_df=vm_df,
            tabela_forecast=tabela_escopo,
            df_particoes=df_part_relatorio,
        )
    except Exception as exc:
        st.warning(f"Não foi possível abrir o relatório assistido por LLM nesta análise: {exc}")




# =============================================================================
# HOTFIX 15F.10.30 — IA comercial, período configurável e relatório ilustrado
# =============================================================================
def _safe_num(v, default=None):
    try:
        if v is None:
            return default
        x = pd.to_numeric(pd.Series([v]), errors="coerce").iloc[0]
        if pd.isna(x):
            return default
        return float(x)
    except Exception:
        return default


def _analise_pick_num(df: pd.DataFrame, candidates, default=None):
    if df is None or df.empty:
        return default
    col = _pick_col_case_insensitive(df.columns, candidates)
    if not col:
        return default
    try:
        vals = pd.to_numeric(df[col], errors="coerce").dropna()
        if vals.empty:
            return default
        return float(vals.iloc[0])
    except Exception:
        return default


def _inferir_capacidade_recurso_vm(vm_df: pd.DataFrame, recurso: str, particao_row: dict = None):
    recurso_l = str(recurso or "").lower()
    if "cpu" in recurso_l:
        cap = _analise_pick_num(vm_df, ["cpu_total_ghz", "cpu_capacity_ghz", "capacity_cpu_ghz", "cpu_alocado_ghz", "cpu_ghz_total"])
        if cap is not None:
            return cap, "GHz", "capacidade absoluta coletada"
        vcpu = _analise_pick_num(vm_df, ["vcpu_count", "vcpus", "vcpu", "num_cpu", "cpu_count", "vCPU"])
        if vcpu is not None:
            return float(vcpu), "vCPU eq.", "estimado por vCPU atual"
        return 100.0, "%", "percentual puro"
    if "mem" in recurso_l or "memória" in recurso_l or "memoria" in recurso_l:
        cap = _analise_pick_num(vm_df, ["memory_gb", "mem_alloc_gb", "mem_gb", "memoria_gb", "ram_gb", "memory_total_gb"])
        if cap is not None:
            return cap, "GB", "memória alocada atual"
        return 100.0, "%", "percentual puro"
    if "disco" in recurso_l or "disk" in recurso_l or "part" in recurso_l:
        if particao_row:
            for c in ["capacity_gb", "total_gb", "partition_capacity_gb", "particao_total_gb", "capacidade_gb"]:
                if c in particao_row and pd.notna(particao_row.get(c)):
                    val = _safe_num(particao_row.get(c))
                    if val is not None:
                        return val, "GB", "capacidade da partição"
        cap = _analise_pick_num(vm_df, ["provisioned_gb", "disk_provisioned_gb", "disk_alloc_gb", "disco_gb", "capacity_gb"])
        if cap is not None:
            return cap, "GB", "disco provisionado atual"
        return 100.0, "%", "percentual puro"
    return 100.0, "%", "percentual puro"


def _serie_sintetica_percentual_vm(vm_df: pd.DataFrame, recurso: str, capacidade: float, unidade: str) -> pd.DataFrame:
    recurso_l = str(recurso or "").lower()
    if "cpu" in recurso_l:
        atual = _analise_pick_num(vm_df, ["cpu_p95_pct", "cpu_peak_typical_pct", "cpu_atual_pct", "CPU — pico típico de uso (%)"])
        f30 = _analise_pick_num(vm_df, ["cpu_forecast_30d", "cpu_previsao_30d_pct"])
        f60 = _analise_pick_num(vm_df, ["cpu_forecast_60d", "cpu_previsao_60d_pct"])
        f90 = _analise_pick_num(vm_df, ["cpu_forecast_90d", "CPU — previsão 90 dias (%)", "cpu_previsao_90d_pct"])
    elif "mem" in recurso_l or "memória" in recurso_l or "memoria" in recurso_l:
        atual = _analise_pick_num(vm_df, ["mem_p95_pct", "memory_p95_pct", "mem_atual_pct", "Memória — pico típico de uso (%)"])
        f30 = _analise_pick_num(vm_df, ["mem_forecast_30d", "mem_previsao_30d_pct"])
        f60 = _analise_pick_num(vm_df, ["mem_forecast_60d", "mem_previsao_60d_pct"])
        f90 = _analise_pick_num(vm_df, ["mem_forecast_90d", "Memória — previsão 90 dias (%)", "mem_previsao_90d_pct"])
    else:
        atual = _analise_pick_num(vm_df, ["disk_p95_pct", "disk_atual_pct", "Disco — pico típico de uso (%)"])
        f30 = _analise_pick_num(vm_df, ["disk_forecast_30d", "forecast_30d", "disk_previsao_30d_pct"])
        f60 = _analise_pick_num(vm_df, ["disk_forecast_60d", "forecast_60d", "disk_previsao_60d_pct"])
        f90 = _analise_pick_num(vm_df, ["disk_forecast_90d", "Disco — previsão 90 dias (%)", "forecast_90d", "disk_previsao_90d_pct"])

    vals = [v for v in [atual, f30, f60, f90] if v is not None]
    if not vals:
        atual, f30, f60, f90 = 0.0, None, None, None
    elif atual is None:
        atual = vals[0]

    hoje = pd.Timestamp.today().normalize()
    pontos = [
        {"data": hoje, "momento": "Atual", "percentual": atual},
        {"data": hoje + pd.Timedelta(days=30), "momento": "30 dias", "percentual": f30},
        {"data": hoje + pd.Timedelta(days=60), "momento": "60 dias", "percentual": f60},
        {"data": hoje + pd.Timedelta(days=90), "momento": "90 dias", "percentual": f90},
    ]
    df = pd.DataFrame(pontos)
    df["percentual"] = pd.to_numeric(df["percentual"], errors="coerce")
    df = df.dropna(subset=["percentual"]).copy()
    if df.empty:
        return df
    if unidade == "%":
        df["uso"] = df["percentual"]
        df["capacidade"] = 100.0
        df["threshold_80"] = 80.0
    else:
        df["uso"] = df["percentual"] * float(capacidade) / 100.0
        df["capacidade"] = float(capacidade)
        df["threshold_80"] = 0.80 * float(capacidade)
    return df


def _buscar_timeseries_vm_duckdb(nome_vm: str, recurso: str, capacidade: float, unidade: str, run_id=None, dias=90) -> pd.DataFrame:
    if duckdb is None or not CAMINHO_BANCO.exists():
        return pd.DataFrame()
    recurso_l = str(recurso or "").lower()
    tabela_candidatas = ["vm_resource_timeseries", "resource_timeseries", "HIST_CPU", "HIST_MEM", "HIST_DISK"]
    try:
        con = duckdb.connect(str(CAMINHO_BANCO), read_only=True)
    except Exception:
        return pd.DataFrame()
    try:
        for tabela in tabela_candidatas:
            try:
                exists = con.execute("select count(*) from information_schema.tables where lower(table_name)=lower(?)", [tabela]).fetchone()[0]
                if not exists:
                    continue
                cols = [r[1] for r in con.execute(f'PRAGMA table_info("{tabela}")').fetchall()]
                col_vm = _pick_col_case_insensitive(cols, ["vm", "vm_name", "name", "resource_name"])
                col_date = _pick_col_case_insensitive(cols, ["date", "data", "timestamp", "time", "dt", "sample_time", "collection_time"])
                if not col_vm or not col_date:
                    continue
                if "cpu" in recurso_l:
                    pct_col = _pick_col_case_insensitive(cols, ["cpu_pct", "cpu_usage_pct", "usage_pct_cpu", "cpu_p95_pct", "value_pct"])
                    abs_col = _pick_col_case_insensitive(cols, ["cpu_usage_ghz", "cpu_used_ghz", "usage_ghz", "cpu_mhz", "usage_mhz"])
                    total_col = _pick_col_case_insensitive(cols, ["cpu_total_ghz", "cpu_capacity_ghz", "capacity_ghz", "cpu_total_mhz"])
                    resource_filter = "CPU"
                elif "mem" in recurso_l or "memória" in recurso_l or "memoria" in recurso_l:
                    pct_col = _pick_col_case_insensitive(cols, ["mem_pct", "memory_pct", "mem_usage_pct", "memory_usage_pct", "mem_p95_pct", "value_pct"])
                    abs_col = _pick_col_case_insensitive(cols, ["mem_used_gb", "memory_used_gb", "usage_gb", "mem_kb", "usage_kb"])
                    total_col = _pick_col_case_insensitive(cols, ["mem_total_gb", "memory_total_gb", "capacity_gb", "mem_alloc_gb"])
                    resource_filter = "MEM"
                else:
                    pct_col = _pick_col_case_insensitive(cols, ["disk_pct", "disk_usage_pct", "usage_pct_disk", "disk_p95_pct", "value_pct"])
                    abs_col = _pick_col_case_insensitive(cols, ["disk_used_gb", "used_gb", "usage_gb"])
                    total_col = _pick_col_case_insensitive(cols, ["disk_total_gb", "capacity_gb", "total_gb"])
                    resource_filter = "DISK"
                col_resource = _pick_col_case_insensitive(cols, ["resource", "recurso", "metric", "metric_name"])
                where = [f'upper("{col_vm}") = upper(?)']
                params = [nome_vm]
                if col_resource:
                    where.append(f'upper("{col_resource}") like ?')
                    params.append(f"%{resource_filter}%")
                if run_id and _pick_col_case_insensitive(cols, ["run_id", "execution_id"]):
                    rid = _pick_col_case_insensitive(cols, ["run_id", "execution_id"])
                    where.append(f'"{rid}" = ?')
                    params.append(str(run_id))
                select_parts = [f'"{col_date}" as data']
                if pct_col:
                    select_parts.append(f'"{pct_col}" as percentual')
                elif abs_col and total_col:
                    select_parts.append(f'case when "{total_col}" != 0 then ("{abs_col}" / "{total_col}") * 100 end as percentual')
                else:
                    continue
                if abs_col:
                    if "mhz" in str(abs_col).lower():
                        select_parts.append(f'("{abs_col}" / 1000.0) as uso')
                    elif "kb" in str(abs_col).lower():
                        select_parts.append(f'("{abs_col}" / 1024.0 / 1024.0) as uso')
                    else:
                        select_parts.append(f'"{abs_col}" as uso')
                else:
                    select_parts.append("NULL as uso")
                if total_col:
                    if "mhz" in str(total_col).lower():
                        select_parts.append(f'("{total_col}" / 1000.0) as capacidade')
                    elif "kb" in str(total_col).lower():
                        select_parts.append(f'("{total_col}" / 1024.0 / 1024.0) as capacidade')
                    else:
                        select_parts.append(f'"{total_col}" as capacidade')
                else:
                    select_parts.append("NULL as capacidade")
                sql = f"""
                    select {", ".join(select_parts)}
                    from "{tabela}"
                    where {" and ".join(where)}
                    order by "{col_date}"
                """
                df = con.execute(sql, params).df()
                if df is None or df.empty:
                    continue
                df["data"] = pd.to_datetime(df["data"], errors="coerce")
                df["percentual"] = pd.to_numeric(df["percentual"], errors="coerce")
                df["uso"] = pd.to_numeric(df["uso"], errors="coerce")
                df["capacidade"] = pd.to_numeric(df["capacidade"], errors="coerce")
                df = df.dropna(subset=["data", "percentual"]).copy()
                if dias:
                    corte = pd.Timestamp.today().normalize() - pd.Timedelta(days=int(dias))
                    df = df[df["data"] >= corte].copy()
                if df.empty:
                    continue
                if df["uso"].isna().all():
                    if unidade == "%":
                        df["uso"] = df["percentual"]
                        df["capacidade"] = 100.0
                    else:
                        df["uso"] = df["percentual"] * float(capacidade) / 100.0
                        df["capacidade"] = float(capacidade)
                else:
                    df["capacidade"] = df["capacidade"].fillna(float(capacidade) if unidade != "%" else 100.0)
                df["threshold_80"] = df["capacidade"] * 0.80
                return df[["data", "percentual", "uso", "capacidade", "threshold_80"]].copy()
            except Exception:
                continue
    finally:
        try:
            con.close()
        except Exception:
            pass
    return pd.DataFrame()


def _calcular_forecast_linear(df_hist: pd.DataFrame, capacidade: float, unidade: str) -> pd.DataFrame:
    if df_hist is None or df_hist.empty or len(df_hist.dropna(subset=["uso"])) < 2:
        return pd.DataFrame()
    df = df_hist.dropna(subset=["data", "uso"]).sort_values("data").copy()
    try:
        import numpy as np
        x = (df["data"] - df["data"].min()).dt.total_seconds() / 86400.0
        y = pd.to_numeric(df["uso"], errors="coerce")
        mask = y.notna()
        x = x[mask]
        y = y[mask]
        if len(y) < 2:
            return pd.DataFrame()
        coef = np.polyfit(x, y, 1)
        last_x = float(x.max())
        last_date = df["data"].max()
        pontos = []
        for h in [30, 60, 90]:
            pred = max(0.0, float(coef[0] * (last_x + h) + coef[1]))
            pct = pred if unidade == "%" else (pred / float(capacidade) * 100.0 if capacidade else None)
            pontos.append({
                "data": last_date + pd.Timedelta(days=h),
                "momento": f"{h} dias",
                "uso": pred,
                "percentual": pct,
                "capacidade": float(capacidade) if unidade != "%" else 100.0,
                "threshold_80": (float(capacidade) * 0.80) if unidade != "%" else 80.0,
            })
        return pd.DataFrame(pontos)
    except Exception:
        return pd.DataFrame()


def _describe_analise(df_hist: pd.DataFrame, unidade: str) -> pd.DataFrame:
    if df_hist is None or df_hist.empty:
        return pd.DataFrame()
    uso = pd.to_numeric(df_hist["uso"], errors="coerce").dropna()
    pct = pd.to_numeric(df_hist["percentual"], errors="coerce").dropna()
    rows = []
    if not uso.empty:
        rows.extend([
            {"Métrica": "Média", "Valor": round(float(uso.mean()), 2), "Unidade": unidade},
            {"Métrica": "Mediana", "Valor": round(float(uso.median()), 2), "Unidade": unidade},
            {"Métrica": "Q1", "Valor": round(float(uso.quantile(0.25)), 2), "Unidade": unidade},
            {"Métrica": "Q3", "Valor": round(float(uso.quantile(0.75)), 2), "Unidade": unidade},
            {"Métrica": "Máximo", "Valor": round(float(uso.max()), 2), "Unidade": unidade},
        ])
    if not pct.empty:
        rows.extend([
            {"Métrica": "Média %", "Valor": round(float(pct.mean()), 2), "Unidade": "%"},
            {"Métrica": "Q3 %", "Valor": round(float(pct.quantile(0.75)), 2), "Unidade": "%"},
            {"Métrica": "Máximo %", "Valor": round(float(pct.max()), 2), "Unidade": "%"},
        ])
    return pd.DataFrame(rows)



# Paleta executiva para relatórios: mantém HTML, Word e PDF consistentes.
RMC_RELATORIO_CORES = {
    "historico": "#1F77B4",
    "historico_fill": "rgba(31, 119, 180, 0.12)",
    "projecao": "#F28E2B",
    "capacidade": "#2CA02C",
    "threshold": "#D62728",
    "media": "#4E79A7",
    "pico": "#9467BD",
    "forecast": "#F28E2B",
    "grid": "#E6ECF5",
    "texto": "#24324A",
    "fundo_ok": "rgba(44, 160, 44, 0.08)",
    "fundo_atencao": "rgba(255, 193, 7, 0.12)",
    "fundo_risco": "rgba(214, 39, 40, 0.10)",
}


def _plotly_layout_executivo(fig, titulo: str, yaxis_title: str, height: int = 420):
    fig.update_layout(
        title=dict(text=titulo, x=0.02, font=dict(size=18, color=RMC_RELATORIO_CORES["texto"])),
        template="plotly_white",
        height=height,
        margin=dict(l=48, r=24, t=62, b=48),
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
        font=dict(color=RMC_RELATORIO_CORES["texto"]),
        plot_bgcolor="white",
        paper_bgcolor="white",
        xaxis=dict(showgrid=True, gridcolor=RMC_RELATORIO_CORES["grid"], title="Data"),
        yaxis=dict(showgrid=True, gridcolor=RMC_RELATORIO_CORES["grid"], title=yaxis_title),
    )
    return fig


def _forecast_continuo_com_historico(df_hist: pd.DataFrame, df_forecast: pd.DataFrame) -> pd.DataFrame:
    """Faz a projeção iniciar no último ponto real para evitar quebra visual entre histórico e forecast."""
    if df_forecast is None or df_forecast.empty:
        return pd.DataFrame()
    dff = df_forecast.copy()
    if df_hist is None or df_hist.empty:
        return dff.sort_values("data")
    try:
        last = df_hist.sort_values("data").iloc[-1]
        first = {c: last[c] if c in last.index else None for c in dff.columns}
        first["data"] = last["data"]
        if "momento" in dff.columns:
            first["momento"] = "Último ponto real"
        dff = pd.concat([pd.DataFrame([first]), dff], ignore_index=True)
        dff["data"] = pd.to_datetime(dff["data"], errors="coerce")
        return dff.dropna(subset=["data"]).sort_values("data")
    except Exception:
        return dff.sort_values("data")

def _fig_comparacao_previsao(df_hist: pd.DataFrame, df_forecast: pd.DataFrame, titulo: str, unidade: str):
    if (df_hist is None or df_hist.empty) and (df_forecast is None or df_forecast.empty):
        return None
    fig = go.Figure()
    if df_hist is not None and not df_hist.empty:
        dfh = df_hist.sort_values("data").copy()
        fig.add_trace(go.Scatter(
            x=dfh["data"], y=dfh["uso"], mode="lines+markers", name="Uso observado",
            line=dict(color=RMC_RELATORIO_CORES["historico"], width=3),
            marker=dict(size=6), fill="tozeroy", fillcolor=RMC_RELATORIO_CORES["historico_fill"],
        ))
        if "capacidade" in dfh.columns:
            fig.add_trace(go.Scatter(
                x=dfh["data"], y=dfh["capacidade"], mode="lines", name="Capacidade atual",
                line=dict(color=RMC_RELATORIO_CORES["capacidade"], width=2),
            ))
        if "threshold_80" in dfh.columns:
            fig.add_trace(go.Scatter(
                x=dfh["data"], y=dfh["threshold_80"], mode="lines", name="Faixa de atenção (80%)",
                line=dict(color=RMC_RELATORIO_CORES["threshold"], width=2, dash="dash"),
            ))
    dff = _forecast_continuo_com_historico(df_hist, df_forecast)
    if dff is not None and not dff.empty:
        fig.add_trace(go.Scatter(
            x=dff["data"], y=dff["uso"], mode="lines+markers", name="Projeção",
            line=dict(color=RMC_RELATORIO_CORES["projecao"], width=3, dash="dot"),
            marker=dict(size=7, color=RMC_RELATORIO_CORES["projecao"]),
        ))
        if df_hist is not None and not df_hist.empty:
            last = df_hist.sort_values("data").iloc[-1]
            fig.add_trace(go.Scatter(
                x=[last["data"]], y=[last["uso"]], mode="markers", name="Início da projeção",
                marker=dict(size=11, color=RMC_RELATORIO_CORES["projecao"], symbol="diamond"),
            ))
    fig = _plotly_layout_executivo(fig, titulo, unidade, height=430)
    return fig


def _fig_media_movel(df_hist: pd.DataFrame, titulo: str, unidade: str):
    if df_hist is None or df_hist.empty:
        return None
    df = df_hist.sort_values("data").copy()
    df["media_movel"] = pd.to_numeric(df["uso"], errors="coerce").rolling(window=min(7, max(2, len(df)//5)), min_periods=1).mean()
    fig = go.Figure()
    fig.add_trace(go.Scatter(x=df["data"], y=df["uso"], mode="lines+markers", name="Uso observado", line=dict(color=RMC_RELATORIO_CORES["historico"], width=2), marker=dict(size=5)))
    fig.add_trace(go.Scatter(x=df["data"], y=df["media_movel"], mode="lines", name="Tendência suavizada", line=dict(color=RMC_RELATORIO_CORES["projecao"], width=3)))
    fig = _plotly_layout_executivo(fig, titulo, unidade, height=360)
    return fig


def _fig_decomposicao_simples(df_hist: pd.DataFrame, titulo: str, unidade: str):
    if df_hist is None or df_hist.empty:
        return None
    df = df_hist.sort_values("data").copy()
    y = pd.to_numeric(df["uso"], errors="coerce").ffill().bfill()
    if y.dropna().empty:
        return None
    tendencia = y.rolling(window=min(14, max(2, len(df)//4)), min_periods=1).mean()
    sazonalidade = y - tendencia
    residuos = y - tendencia - sazonalidade.rolling(window=min(7, max(2, len(df)//5)), min_periods=1).mean()
    fig = go.Figure()
    fig.add_trace(go.Scatter(x=df["data"], y=y, mode="lines", name="Série original", line=dict(color=RMC_RELATORIO_CORES["historico"], width=2)))
    fig.add_trace(go.Scatter(x=df["data"], y=tendencia, mode="lines", name="Tendência", line=dict(color=RMC_RELATORIO_CORES["projecao"], width=3)))
    fig.add_trace(go.Scatter(x=df["data"], y=sazonalidade, mode="lines", name="Oscilações", line=dict(color="#59A14F", width=2)))
    fig.add_trace(go.Scatter(x=df["data"], y=residuos, mode="markers", name="Variações pontuais", marker=dict(color=RMC_RELATORIO_CORES["pico"], size=7)))
    fig = _plotly_layout_executivo(fig, titulo, unidade, height=420)
    return fig


def _fig_histograma(df_hist: pd.DataFrame, titulo: str, unidade: str):
    if df_hist is None or df_hist.empty:
        return None
    fig = go.Figure()
    fig.add_trace(go.Histogram(x=df_hist["uso"], name="Frequência", marker=dict(color=RMC_RELATORIO_CORES["historico"], line=dict(color="white", width=1)), opacity=0.9))
    fig.update_layout(title=titulo, xaxis_title=unidade, yaxis_title="Frequência", height=340, template="plotly_white", plot_bgcolor="white", paper_bgcolor="white")
    return fig


def _fig_uso_medio_hora(df_hist: pd.DataFrame, titulo: str, unidade: str):
    if df_hist is None or df_hist.empty:
        return None
    df = df_hist.copy()
    df["hora"] = pd.to_datetime(df["data"], errors="coerce").dt.hour
    if df["hora"].isna().all():
        return None
    grp = df.groupby("hora", as_index=False)["uso"].mean()
    fig = go.Figure()
    fig.add_trace(go.Bar(x=grp["hora"], y=grp["uso"], name="Uso médio", marker=dict(color=RMC_RELATORIO_CORES["media"])))
    fig.update_layout(title=titulo, xaxis_title="Hora do dia", yaxis_title=unidade, height=340, template="plotly_white", plot_bgcolor="white", paper_bgcolor="white")
    return fig


def _classificar_recomendacao_analise(df_hist: pd.DataFrame, df_forecast: pd.DataFrame):
    if df_hist is None or df_hist.empty:
        return "SEM_DADOS", "Não há histórico suficiente para uma recomendação segura. O ideal é ampliar a coleta antes de uma decisão definitiva."
    pct_atual = pd.to_numeric(df_hist["percentual"], errors="coerce").dropna()
    pct_fore = pd.to_numeric(df_forecast["percentual"], errors="coerce").dropna() if df_forecast is not None and not df_forecast.empty else pd.Series(dtype=float)
    max_atual = float(pct_atual.max()) if not pct_atual.empty else 0.0
    q3 = float(pct_atual.quantile(0.75)) if len(pct_atual) else 0.0
    max_fore = float(pct_fore.max()) if not pct_fore.empty else max_atual
    if q3 >= 90 or max_fore >= 95:
        return "CRÍTICO", "O recurso está muito próximo do limite. Existe chance relevante de impacto no serviço, então a ação deve ser tratada com prioridade alta."
    if q3 >= 80 or max_fore >= 80:
        return "RISCO", "O recurso já entrou ou tende a entrar na faixa de atenção. Recomenda-se acompanhar de perto e avaliar ampliação ou ajuste do ambiente."
    if q3 < 20 and max_atual < 50:
        return "OTIMIZAÇÃO", "O uso ficou bem abaixo da capacidade disponível. Há indícios de sobra de recurso e pode valer revisar o dimensionamento para ganhar eficiência."
    return "OK", "O comportamento observado está dentro de uma faixa confortável. Neste momento, o cenário indica continuidade do monitoramento sem ação urgente."


def _extrair_resumo_executivo(df_hist: pd.DataFrame, df_forecast: pd.DataFrame):
    pct_hist = pd.to_numeric(df_hist["percentual"], errors="coerce").dropna() if df_hist is not None and not df_hist.empty else pd.Series(dtype=float)
    pct_fore = pd.to_numeric(df_forecast["percentual"], errors="coerce").dropna() if df_forecast is not None and not df_forecast.empty else pd.Series(dtype=float)
    media = float(pct_hist.mean()) if not pct_hist.empty else None
    pico = float(pct_hist.max()) if not pct_hist.empty else None
    q3 = float(pct_hist.quantile(0.75)) if not pct_hist.empty else None
    max_fore = float(pct_fore.max()) if not pct_fore.empty else None
    hist_cross = int((pct_hist >= 80).sum()) if not pct_hist.empty else 0
    fore_cross = int((pct_fore >= 80).sum()) if not pct_fore.empty else 0
    return {
        "media_pct": round(media, 2) if media is not None else None,
        "pico_pct": round(pico, 2) if pico is not None else None,
        "q3_pct": round(q3, 2) if q3 is not None else None,
        "max_forecast_pct": round(max_fore, 2) if max_fore is not None else None,
        "pontos_hist_acima_80": hist_cross,
        "pontos_forecast_acima_80": fore_cross,
    }


def _montar_contexto_analise_completa(numero_solicitacao: str, nome_vm: str, recurso: str, capacidade: float, unidade: str, origem_capacidade: str, origem_serie: str, historico_dias: int, df_hist: pd.DataFrame, df_forecast: pd.DataFrame, describe_df: pd.DataFrame, status: str, recomendacao: str, solicitante: str = "Solicitante"):
    return {
        "numero_solicitacao": numero_solicitacao,
        "vm": nome_vm,
        "solicitante": solicitante or "Solicitante",
        "recurso": recurso,
        "capacidade": capacidade,
        "unidade": unidade,
        "origem_capacidade": origem_capacidade,
        "origem_serie": origem_serie,
        "historico_dias": historico_dias,
        "margem_atencao_pct": 80,
        "status": status,
        "recomendacao_modelo": recomendacao,
        "resumo_executivo_numerico": _extrair_resumo_executivo(df_hist, df_forecast),
        "estatistica_descritiva": describe_df.to_dict(orient="records") if describe_df is not None and not describe_df.empty else [],
        "historico_amostra": df_hist.tail(20).to_dict(orient="records") if df_hist is not None and not df_hist.empty else [],
        "forecast": df_forecast.to_dict(orient="records") if df_forecast is not None and not df_forecast.empty else [],
        "diretiva_prompt": "Escreva como consultor de negócios. Explique termos técnicos em linguagem simples. Evite jargões sem explicação. O leitor pode ser leigo em capacity planning e em TI.",
    }


def _ia_json_extract(texto: str):
    """Extrai o primeiro JSON válido da resposta da IA, tolerando texto extra."""
    raw = str(texto or "").strip()
    if not raw:
        raise ValueError("resposta vazia")
    # Remove cercas de código, caso a IA desobedeça.
    raw = re.sub(r"^```(?:json)?", "", raw.strip(), flags=re.I).strip()
    raw = re.sub(r"```$", "", raw.strip()).strip()
    start = raw.find("{")
    end = raw.rfind("}")
    if start < 0 or end <= start:
        raise ValueError("JSON não encontrado")
    return json.loads(raw[start:end+1])


def _normalizar_status_txt(status: str) -> str:
    s = str(status or "").strip().upper().replace("CRITICO", "CRÍTICO").replace("OTIMIZACAO", "OTIMIZAÇÃO")
    return s or "N/C"


def _texto_tem_contradicao_recurso(contexto: dict, texto: str) -> bool:
    """Bloqueia respostas da IA que contradizem a classificação calculada."""
    status = _normalizar_status_txt(contexto.get("status"))
    recurso = str(contexto.get("recurso", "recurso"))
    low = str(texto or "").lower()
    if not low.strip():
        return True
    # A IA não pode analisar o produto RMC Copilot como se fosse o servidor.
    if "capacidade do rmc copilot" in low or "uso do rmc copilot" in low:
        return True
    # Recurso deve aparecer ou texto precisa ser genérico do servidor.
    if status == "OTIMIZAÇÃO":
        proibidos = ["crise", "urgente", "sobrecarga", "exaustão", "travamento", "lentidão extrema", "aumentar imediatamente", "aumento imediato", "abaixo da demanda"]
        if any(p in low for p in proibidos):
            return True
    if status in ["CRÍTICO", "RISCO"]:
        proibidos = ["sem risco", "não apresenta risco", "não há risco", "não exige atenção", "apenas monitoramento", "muito confortável"]
        if any(p in low for p in proibidos):
            return True
    return False


def _montar_texto_controlado_recurso(contexto: dict) -> dict:
    resumo = contexto.get("resumo_executivo_numerico", {}) or {}
    recurso = str(contexto.get("recurso", "recurso"))
    vm = str(contexto.get("vm", "servidor"))
    status = _normalizar_status_txt(contexto.get("status"))
    capacidade = contexto.get("capacidade", "N/D")
    unidade = contexto.get("unidade", "")
    media = resumo.get("media_pct", "N/D")
    pico = resumo.get("pico_pct", "N/D")
    q3 = resumo.get("q3_pct", "N/D")
    max_fore = resumo.get("max_forecast_pct", "N/D")
    recomendacao = contexto.get("recomendacao_modelo", "Manter monitoramento.")
    hist_dias = contexto.get("historico_dias", "N/D")
    origem = contexto.get("origem_serie", "N/D")
    # Texto base controlado, sem invenção, no estilo do PDF de referência.
    resumo_exec = (
        f"Após analisar os dados de utilização de {recurso} do servidor {vm} no período de {hist_dias} dias, "
        f"a situação foi classificada como {status}. A utilização média foi de {media}% e o maior pico observado foi de {pico}%. "
        f"A faixa de atenção adotada é 80% da capacidade, pois acima desse ponto a margem operacional fica menor para variações de carga, rotinas agendadas e crescimento do ambiente."
    )
    if status == "OTIMIZAÇÃO":
        resumo_exec += " Os números indicam folga relevante no recurso, portanto a recomendação deve priorizar revisão de dimensionamento antes de qualquer aumento."
    elif status in ["CRÍTICO", "RISCO"]:
        resumo_exec += " Os números indicam consumo elevado ou tendência de aproximação do limite, portanto o recurso deve ser tratado com atenção antes que afete a operação."
    else:
        resumo_exec += " O comportamento observado não indica necessidade imediata de aumento, mas o monitoramento deve continuar."
    return {
        "objetivo": f"Avaliar o comportamento de {recurso} do servidor {vm}, explicar os gráficos de forma clara e indicar se há necessidade de aumentar, reduzir, manter ou acompanhar o recurso.",
        "resumo_executivo": resumo_exec,
        "historico_previsao": f"O gráfico de comportamento histórico e previsão compara o uso observado com a capacidade considerada ({capacidade} {unidade}), a faixa de atenção de 80% e a projeção. A projeção parte do último ponto real para demonstrar continuidade entre o histórico e a tendência futura.",
        "decomposicao": "A decomposição separa a tendência principal das oscilações pontuais. Isso evita que a decisão seja tomada apenas por um pico isolado quando o comportamento geral é estável, ou ajuda a confirmar risco quando a tendência se mantém alta.",
        "histograma": f"A distribuição do uso mostra onde o servidor permaneceu na maior parte do tempo. Nesta análise, a média foi {media}% e o Q3 foi {q3}%, ajudando a entender se o consumo típico está próximo ou distante da faixa de atenção.",
        "uso_horario": "O uso médio por horário ajuda a identificar se existe concentração de carga em determinados períodos do dia. Essa leitura é útil para diferenciar consumo constante de picos causados por rotinas específicas.",
        "estatistica": f"A estatística resume o comportamento observado: média {media}%, Q3 {q3}%, pico {pico}% e maior projeção {max_fore}%. Esses números devem ser lidos em conjunto com a faixa de atenção de 80%.",
        "conclusao": f"Conclusão para {recurso}: {recomendacao}",
        "recomendacao": recomendacao,
        "observacao_dados": f"Fonte da série: {origem}. Quando a série absoluta não está disponível, o consumo é estimado a partir do percentual e da capacidade atual, e essa condição deve ser considerada na interpretação.",
    }


def _gerar_relatorio_analise_completa_ia(contexto: dict):
    """Gera o texto da Análise Individual por IA, com guarda-corpo contra contradições e sem Markdown bruto obrigatório."""
    prompt = f"""
Você é a IA do RMC Copilot. Gere a análise individual de VM/servidor para UM recurso.

IMPORTANTE:
- Isto NÃO é FUP. É Análise Individual de VM/Servidor.
- A IA só analisa, explica e recomenda. Não toma ação operacional.
- Use somente os dados do contexto. Não invente aplicação, causa, horário, processo, percentual ou capacidade.
- Não analise o produto RMC Copilot; analise o servidor/VM informado no contexto.
- Se o status calculado for OTIMIZAÇÃO, não fale em crise, urgência, sobrecarga ou aumento imediato.
- Se o status calculado for CRÍTICO ou RISCO, não diga que está sem risco ou que basta encerrar.
- Use linguagem de negócio, simples, para leitor leigo em capacity planning.
- Responda SOMENTE JSON válido, sem Markdown, sem ```.

Formato JSON obrigatório:
{{
  "objetivo": "texto",
  "resumo_executivo": "texto",
  "historico_previsao": "texto",
  "decomposicao": "texto",
  "histograma": "texto",
  "uso_horario": "texto",
  "estatistica": "texto",
  "conclusao": "texto",
  "recomendacao": "texto",
  "observacao_dados": "texto"
}}

Contexto:
{json.dumps(contexto, ensure_ascii=False, default=str, indent=2)}
"""
    origem = "IA"
    try:
        resposta = _chamar_llm_rmc(prompt)
        dados = _ia_json_extract(resposta)
        obrig = ["objetivo", "resumo_executivo", "historico_previsao", "decomposicao", "histograma", "uso_horario", "estatistica", "conclusao", "recomendacao", "observacao_dados"]
        if not isinstance(dados, dict) or any(not str(dados.get(k, "")).strip() for k in obrig):
            raise ValueError("JSON incompleto")
        unido = "\n".join(str(dados.get(k, "")) for k in obrig)
        if _texto_tem_contradicao_recurso(contexto, unido):
            raise ValueError("resposta da IA contraditória ou fora do escopo")
    except Exception as exc:
        dados = _montar_texto_controlado_recurso(contexto)
        origem = f"IA controlada ({exc})"

    # Monta em Markdown simples porque os conversores já transformam em Word/PDF/HTML real.
    texto = f"""# Relatório de Utilização de {contexto.get('recurso')}: Servidor {contexto.get('vm')}

## 1. Objetivo da análise
{dados.get('objetivo')}

## 2. Resumo Executivo
{dados.get('resumo_executivo')}

## 3. Análise Detalhada dos Gráficos

### A. Comportamento Histórico e Previsão
{dados.get('historico_previsao')}

### B. Decomposição da Série Temporal
{dados.get('decomposicao')}

### C. Distribuição do Uso / Histograma
{dados.get('histograma')}

### D. Uso Médio por Hora
{dados.get('uso_horario')}

## 4. Análise Estatística
{dados.get('estatistica')}

## 5. Conclusão e Recomendação Final
{dados.get('conclusao')}

{dados.get('recomendacao')}

Observação sobre os dados: {dados.get('observacao_dados')}
"""
    return texto, origem


def _fig_to_html(fig):
    if fig is None:
        return ""
    return fig.to_html(full_html=False, include_plotlyjs=False, config={"displayModeBar": False, "responsive": True})



def _md_strip_inline(texto: str) -> str:
    """Remove marcações Markdown sem remover o conteúdo."""
    s = str(texto or "")
    s = re.sub(r"\[([^\]]+)\]\(([^\)]+)\)", r"\1", s)
    s = re.sub(r"\*\*([^*]+)\*\*", r"\1", s)
    s = re.sub(r"__([^_]+)__", r"\1", s)
    s = re.sub(r"`([^`]+)`", r"\1", s)
    s = re.sub(r"^\s*>\s?", "", s)
    return s.strip()


def _md_inline_to_html(texto: str) -> str:
    s = html.escape(str(texto or ""))
    s = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", s)
    s = re.sub(r"__([^_]+)__", r"<strong>\1</strong>", s)
    s = re.sub(r"`([^`]+)`", r"<code>\1</code>", s)
    s = re.sub(r"\[([^\]]+)\]\(([^\)]+)\)", r"\1", s)
    return s


def _is_md_table_sep(line: str) -> bool:
    parts = [p.strip() for p in str(line).strip().strip("|").split("|")]
    return bool(parts) and all(re.fullmatch(r":?-{3,}:?", p or "") for p in parts)


def _parse_md_table(lines, start_idx):
    if start_idx + 1 >= len(lines):
        return None, start_idx
    if "|" not in lines[start_idx] or not _is_md_table_sep(lines[start_idx + 1]):
        return None, start_idx
    rows = []
    i = start_idx
    while i < len(lines) and "|" in lines[i]:
        if not _is_md_table_sep(lines[i]):
            rows.append([_md_strip_inline(c) for c in lines[i].strip().strip("|").split("|")])
        i += 1
    return rows, i


def _markdown_to_html_relatorio(texto: str) -> str:
    """Converte um subconjunto seguro de Markdown para HTML real, sem expor ##/**/pipes."""
    lines = str(texto or "").splitlines()
    out = []
    in_ul = False
    in_ol = False

    def close_lists():
        nonlocal in_ul, in_ol
        if in_ul:
            out.append("</ul>")
            in_ul = False
        if in_ol:
            out.append("</ol>")
            in_ol = False

    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            close_lists()
            i += 1
            continue
        table, next_i = _parse_md_table(lines, i)
        if table:
            close_lists()
            out.append('<table class="md-table">')
            for r_idx, row in enumerate(table):
                tag = "th" if r_idx == 0 else "td"
                out.append("<tr>" + "".join(f"<{tag}>{html.escape(str(c))}</{tag}>" for c in row) + "</tr>")
            out.append("</table>")
            i = next_i
            continue
        h = re.match(r"^(#{1,6})\s+(.*)$", line)
        if h:
            close_lists()
            level = min(4, len(h.group(1)) + 1)
            out.append(f"<h{level}>{_md_inline_to_html(h.group(2))}</h{level}>")
            i += 1
            continue
        bullet = re.match(r"^[-*•]\s+(.*)$", line)
        if bullet:
            if in_ol:
                out.append("</ol>")
                in_ol = False
            if not in_ul:
                out.append("<ul>")
                in_ul = True
            out.append(f"<li>{_md_inline_to_html(bullet.group(1))}</li>")
            i += 1
            continue
        num = re.match(r"^\d+[\.)]\s+(.*)$", line)
        if num:
            if in_ul:
                out.append("</ul>")
                in_ul = False
            if not in_ol:
                out.append("<ol>")
                in_ol = True
            out.append(f"<li>{_md_inline_to_html(num.group(1))}</li>")
            i += 1
            continue
        close_lists()
        out.append(f"<p>{_md_inline_to_html(line)}</p>")
        i += 1
    close_lists()
    return "\n".join(out)


def _docx_add_inline_runs(paragraph, texto: str):
    """Adiciona runs no Word tratando **negrito** sem deixar marcação visível."""
    s = str(texto or "")
    pos = 0
    for m in re.finditer(r"\*\*([^*]+)\*\*|__([^_]+)__", s):
        if m.start() > pos:
            paragraph.add_run(_md_strip_inline(s[pos:m.start()]))
        run = paragraph.add_run(_md_strip_inline(m.group(1) or m.group(2)))
        run.bold = True
        pos = m.end()
    if pos < len(s):
        paragraph.add_run(_md_strip_inline(s[pos:]))


def _add_markdown_to_docx(doc, texto: str):
    """Renderiza Markdown simples como elementos nativos do Word."""
    lines = str(texto or "").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        table, next_i = _parse_md_table(lines, i)
        if table:
            tbl = doc.add_table(rows=1, cols=len(table[0]))
            tbl.style = "Table Grid"
            for c, val in enumerate(table[0]):
                tbl.rows[0].cells[c].text = str(val)
            for row in table[1:]:
                cells = tbl.add_row().cells
                for c, val in enumerate(row[:len(cells)]):
                    cells[c].text = str(val)
            i = next_i
            continue
        h = re.match(r"^(#{1,6})\s+(.*)$", line)
        if h:
            doc.add_heading(_md_strip_inline(h.group(2)), level=min(3, len(h.group(1))))
            i += 1
            continue
        bullet = re.match(r"^[-*•]\s+(.*)$", line)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            _docx_add_inline_runs(p, bullet.group(1))
            i += 1
            continue
        num = re.match(r"^\d+[\.)]\s+(.*)$", line)
        if num:
            p = doc.add_paragraph(style="List Number")
            _docx_add_inline_runs(p, num.group(1))
            i += 1
            continue
        p = doc.add_paragraph()
        _docx_add_inline_runs(p, line)
        i += 1


def _md_inline_to_reportlab(texto: str) -> str:
    s = html.escape(str(texto or ""))
    s = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", s)
    s = re.sub(r"__([^_]+)__", r"<b>\1</b>", s)
    s = re.sub(r"`([^`]+)`", r"\1", s)
    s = re.sub(r"\[([^\]]+)\]\(([^\)]+)\)", r"\1", s)
    return s


def _markdown_to_pdf_flowables(texto: str, styles):
    """Renderiza Markdown simples como Paragraph/Table no ReportLab."""
    from reportlab.lib import colors
    from reportlab.platypus import Paragraph, Spacer, Table, TableStyle
    flow = []
    lines = str(texto or "").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            flow.append(Spacer(1, 4))
            i += 1
            continue
        table, next_i = _parse_md_table(lines, i)
        if table:
            tbl = Table(table, hAlign="LEFT")
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DCEBFA")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1F3B73")),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#A7B7C7")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
            ]))
            flow.append(tbl)
            flow.append(Spacer(1, 6))
            i = next_i
            continue
        h = re.match(r"^(#{1,6})\s+(.*)$", line)
        if h:
            flow.append(Paragraph(_md_inline_to_reportlab(h.group(2)), styles["H2RMC"]))
            i += 1
            continue
        bullet = re.match(r"^[-*•]\s+(.*)$", line)
        if bullet:
            flow.append(Paragraph("• " + _md_inline_to_reportlab(bullet.group(1)), styles["BodyRMC"]))
            i += 1
            continue
        num = re.match(r"^(\d+[\.)])\s+(.*)$", line)
        if num:
            flow.append(Paragraph(html.escape(num.group(1)) + " " + _md_inline_to_reportlab(num.group(2)), styles["BodyRMC"]))
            i += 1
            continue
        flow.append(Paragraph(_md_inline_to_reportlab(line), styles["BodyRMC"]))
        i += 1
    return flow


def _gerar_html_relatorio_ilustrado(relatorio_md: str, contexto: dict, figs: dict, describe_df: pd.DataFrame, df_hist: pd.DataFrame = None, df_forecast: pd.DataFrame = None):
    import base64
    meta = contexto.get("resumo_executivo_numerico", {}) or {}
    tables_html = describe_df.to_html(index=False, border=0) if describe_df is not None and not describe_df.empty else "<p>Sem estatística descritiva disponível.</p>"
    imagens = _criar_figuras_png_relatorio(df_hist, df_forecast, figs, contexto.get("recurso", ""), contexto.get("unidade", "")) if df_hist is not None else {}
    nomes = [
        ("01_comparacao_previsao", "Comparação e previsão"),
        ("02_media_movel", "Histórico com média móvel"),
        ("03_decomposicao", "Decomposição da série temporal"),
        ("04_histograma", "Distribuição do uso"),
        ("05_uso_medio_hora", "Uso médio por horário"),
    ]
    fig_blocks = []
    for key, titulo in nomes:
        img = imagens.get(key)
        if img:
            b64 = base64.b64encode(img).decode("ascii")
            fig_blocks.append(f"<h3>{html.escape(titulo)}</h3><img class='grafico' src='data:image/png;base64,{b64}' />")
    if not fig_blocks:
        # Último recurso: renderiza Plotly se não houver imagem estática.
        fig_blocks = [f"<h3>{html.escape(nome)}</h3>{_fig_to_html(fig)}" for nome, fig in (figs or {}).items() if fig is not None]
    return f"""
<!DOCTYPE html>
<html lang="pt-br">
<head>
  <meta charset="utf-8" />
  <title>Relatório ilustrado - {html.escape(str(contexto.get('vm')))}</title>
  <script src="https://cdn.plot.ly/plotly-2.35.2.min.js"></script>
  <style>
    body {{ font-family: Arial, sans-serif; margin: 24px; color: #24324a; background: #f6f8fc; }}
    .card {{ background: #fff; border-radius: 14px; padding: 18px 20px; margin-bottom: 18px; box-shadow: 0 2px 10px rgba(0,0,0,.06); }}
    .title {{ color: #1F3B73; }}
    .grid {{ display: grid; grid-template-columns: repeat(4, minmax(180px, 1fr)); gap: 12px; }}
    .metric {{ background: #eef4ff; border-radius: 12px; padding: 12px; }}
    .metric small {{ display:block; color:#5d6a7a; }}
    .metric strong {{ font-size: 22px; color:#1F3B73; }}
    .content {{ line-height: 1.5; }}
    .grafico {{ width: 100%; max-width: 980px; display:block; margin: 8px auto 18px auto; border: 1px solid #e5e7eb; border-radius: 8px; }}
    table {{ border-collapse: collapse; width: 100%; background: white; }}
    th, td {{ padding: 8px 10px; border-bottom: 1px solid #e5e7eb; text-align: left; }}
    th {{ background:#DCEBFA; color:#1F3B73; }}
    h2, h3 {{ color: #1F3B73; }}
  </style>
</head>
<body>
  <div class="card">
    <h1 class="title">Relatório de análise individual de VM</h1>
    <p><strong>Solicitação:</strong> {html.escape(str(contexto.get('numero_solicitacao')))}<br>
    <strong>Servidor/VM:</strong> {html.escape(str(contexto.get('vm')))}<br>
    <strong>Recurso:</strong> {html.escape(str(contexto.get('recurso')))}<br>
    <strong>Histórico considerado:</strong> {html.escape(str(contexto.get('historico_dias')))} dias<br>
    <strong>Fonte da série:</strong> {html.escape(str(contexto.get('origem_serie')))}<br>
    <strong>Faixa de atenção:</strong> 80% da capacidade</p>
  </div>
  <div class="card grid">
    <div class="metric"><small>Média de uso</small><strong>{meta.get('media_pct', 'N/D')}%</strong></div>
    <div class="metric"><small>Pico observado</small><strong>{meta.get('pico_pct', 'N/D')}%</strong></div>
    <div class="metric"><small>Pico projetado</small><strong>{meta.get('max_forecast_pct', 'N/D')}%</strong></div>
    <div class="metric"><small>Classificação</small><strong>{html.escape(str(contexto.get('status')))}</strong></div>
  </div>
  <div class="card">
    <h2>Texto gerado pela IA</h2>
    <div class="content">{_markdown_to_html_relatorio(relatorio_md)}</div>
  </div>
  <div class="card">
    <h2>Figuras de apoio</h2>
    {''.join(fig_blocks)}
  </div>
  <div class="card">
    <h2>Estatística descritiva</h2>
    {tables_html}
  </div>
</body>
</html>
"""


def _gerar_pacote_relatorio_ia(numero_solicitacao: str, nome_vm: str, recurso: str, relatorio_md: str, html_relatorio: str, describe_df: pd.DataFrame, df_hist: pd.DataFrame, df_forecast: pd.DataFrame):
    mem = io.BytesIO()
    slug = f"{_normalizar_nome_arquivo(numero_solicitacao)}_{_normalizar_nome_arquivo(nome_vm)}_{_normalizar_nome_arquivo(recurso)}"
    with zipfile.ZipFile(mem, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(f"relatorio_{slug}.md", relatorio_md)
        zf.writestr(f"relatorio_ilustrado_{slug}.html", html_relatorio)
        if describe_df is not None and not describe_df.empty:
            zf.writestr(f"estatistica_{slug}.csv", describe_df.to_csv(index=False, sep=';', encoding='utf-8-sig'))
        if df_hist is not None and not df_hist.empty:
            zf.writestr(f"historico_{slug}.csv", df_hist.to_csv(index=False, sep=';', encoding='utf-8-sig'))
        if df_forecast is not None and not df_forecast.empty:
            zf.writestr(f"forecast_{slug}.csv", df_forecast.to_csv(index=False, sep=';', encoding='utf-8-sig'))
    mem.seek(0)
    return mem.getvalue()



def _criar_figuras_png_relatorio(df_hist: pd.DataFrame, df_forecast: pd.DataFrame, figs_plotly: dict, recurso: str, unidade: str) -> dict:
    """
    Gera figuras estáticas PNG coloridas para Word/PDF.
    A projeção é conectada ao último ponto histórico para manter continuidade visual.
    """
    imagens = {}
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import matplotlib.dates as mdates

        def _save_current():
            bio = io.BytesIO()
            plt.tight_layout()
            plt.savefig(bio, format="png", dpi=170, bbox_inches="tight", facecolor="white")
            plt.close()
            bio.seek(0)
            return bio.getvalue()

        hist_ok = df_hist is not None and not df_hist.empty
        fore_ok = df_forecast is not None and not df_forecast.empty
        if hist_ok:
            dfh = df_hist.sort_values("data").copy()
            dff = _forecast_continuo_com_historico(dfh, df_forecast) if fore_ok else pd.DataFrame()

            # 1 - Comparação e projeção
            plt.figure(figsize=(9, 4.8))
            ax = plt.gca()
            ax.axhspan(0, 80, color="#2CA02C", alpha=0.07)
            ax.axhspan(80, 90, color="#FFC107", alpha=0.12)
            ax.axhspan(90, max(float(dfh.get("capacidade", pd.Series([100])).max()), 100) * 1.05, color="#D62728", alpha=0.08)
            ax.plot(dfh["data"], dfh["uso"], color="#1F77B4", linewidth=2.8, marker="o", label="Uso observado")
            ax.fill_between(dfh["data"], pd.to_numeric(dfh["uso"], errors="coerce"), color="#1F77B4", alpha=0.10)
            if "capacidade" in dfh.columns:
                ax.plot(dfh["data"], dfh["capacidade"], color="#2CA02C", linewidth=2, label="Capacidade atual")
            if "threshold_80" in dfh.columns:
                ax.plot(dfh["data"], dfh["threshold_80"], color="#D62728", linewidth=2, linestyle="--", label="Faixa de atenção (80%)")
            if dff is not None and not dff.empty:
                ax.plot(dff["data"], dff["uso"], color="#F28E2B", linewidth=2.8, linestyle=":", marker="o", label="Projeção")
                last = dfh.iloc[-1]
                ax.scatter([last["data"]], [last["uso"]], color="#F28E2B", marker="D", s=70, zorder=5, label="Início da projeção")
            ax.set_title(f"Comparação e projeção - {recurso}", fontsize=13, color="#24324A", weight="bold")
            ax.set_xlabel("Data")
            ax.set_ylabel(unidade)
            ax.grid(True, color="#E6ECF5")
            ax.legend(loc="best", frameon=True)
            imagens["01_comparacao_previsao"] = _save_current()

            # 2 - Média móvel
            dfh["media_movel"] = pd.to_numeric(dfh["uso"], errors="coerce").rolling(window=min(7, max(2, len(dfh)//5)), min_periods=1).mean()
            plt.figure(figsize=(9, 4.4))
            ax = plt.gca()
            ax.plot(dfh["data"], dfh["uso"], color="#1F77B4", linewidth=2.3, marker="o", label="Uso observado")
            ax.plot(dfh["data"], dfh["media_movel"], color="#F28E2B", linewidth=3, label="Tendência suavizada")
            ax.fill_between(dfh["data"], pd.to_numeric(dfh["media_movel"], errors="coerce"), color="#F28E2B", alpha=0.10)
            ax.set_title(f"Histórico com média móvel - {recurso}", fontsize=13, color="#24324A", weight="bold")
            ax.set_xlabel("Data")
            ax.set_ylabel(unidade)
            ax.grid(True, color="#E6ECF5")
            ax.legend()
            imagens["02_media_movel"] = _save_current()

            # 3 - Decomposição simples
            y = pd.to_numeric(dfh["uso"], errors="coerce").ffill().bfill()
            tendencia = y.rolling(window=min(14, max(2, len(dfh)//4)), min_periods=1).mean()
            oscilacoes = y - tendencia
            residuos = y - tendencia - oscilacoes.rolling(window=min(7, max(2, len(dfh)//5)), min_periods=1).mean()
            fig, axes = plt.subplots(4, 1, figsize=(9, 7), sharex=True, facecolor="white")
            series_cfg = [
                (y, "Série original", "#1F77B4"),
                (tendencia, "Tendência", "#F28E2B"),
                (oscilacoes, "Oscilações", "#59A14F"),
                (residuos, "Variações pontuais", "#9467BD"),
            ]
            for ax, (serie, titulo_ax, cor) in zip(axes, series_cfg):
                if titulo_ax == "Variações pontuais":
                    ax.scatter(dfh["data"], serie, color=cor, s=20)
                else:
                    ax.plot(dfh["data"], serie, color=cor, linewidth=2)
                ax.set_title(titulo_ax, fontsize=10, color="#24324A")
                ax.grid(True, color="#E6ECF5")
            imagens["03_decomposicao"] = _save_current()

            # 4 - Histograma
            plt.figure(figsize=(9, 4.4))
            ax = plt.gca()
            ax.hist(pd.to_numeric(dfh["uso"], errors="coerce").dropna(), bins=18, color="#1F77B4", edgecolor="white", alpha=0.9)
            ax.set_title(f"Distribuição do uso - {recurso}", fontsize=13, color="#24324A", weight="bold")
            ax.set_xlabel(unidade)
            ax.set_ylabel("Frequência")
            ax.grid(True, axis="y", color="#E6ECF5")
            imagens["04_histograma"] = _save_current()

            # 5 - Uso por hora
            dfh["hora"] = pd.to_datetime(dfh["data"], errors="coerce").dt.hour
            if not dfh["hora"].isna().all():
                grp = dfh.groupby("hora", as_index=False)["uso"].mean()
                plt.figure(figsize=(9, 4.4))
                ax = plt.gca()
                ax.bar(grp["hora"], grp["uso"], color="#4E79A7", alpha=0.92)
                ax.set_title(f"Uso médio por horário - {recurso}", fontsize=13, color="#24324A", weight="bold")
                ax.set_xlabel("Hora do dia")
                ax.set_ylabel(unidade)
                ax.grid(True, axis="y", color="#E6ECF5")
                imagens["05_uso_medio_hora"] = _save_current()
    except Exception:
        try:
            for nome, fig in (figs_plotly or {}).items():
                if fig is not None:
                    imagens[_normalizar_nome_arquivo(nome)] = fig.to_image(format="png", scale=2)
        except Exception:
            return imagens
    return imagens


def _relatorio_exec_sections(contexto: dict, relatorio_ia: str):
    resumo = contexto.get("resumo_executivo_numerico", {}) or {}
    recurso = str(contexto.get("recurso", "recurso"))
    status = str(contexto.get("status", "N/C"))
    capacidade = contexto.get("capacidade", "N/C")
    unidade = contexto.get("unidade", "")
    origem_serie = contexto.get("origem_serie", "N/C")
    dias = contexto.get("historico_dias", "N/C")
    media = resumo.get("media_pct", "N/D")
    pico = resumo.get("pico_pct", "N/D")
    q3 = resumo.get("q3_pct", "N/D")
    max_fore = resumo.get("max_forecast_pct", "N/D")
    recomendacao = contexto.get("recomendacao_modelo", "")

    objetivo = (
        f"Este relatório avalia o comportamento de {recurso} do servidor {contexto.get('vm')} "
        f"no período de {dias} dias. O objetivo é indicar, em linguagem simples, se o ambiente "
        "tem folga, se está entrando em uma faixa de atenção ou se exige ação imediata."
    )
    resumo_exec = (
        f"A análise classificou o cenário como {status}. No período avaliado, a média de utilização "
        f"foi de {media}% e o maior pico observado foi de {pico}%. A faixa de atenção adotada é 80% "
        "da capacidade, pois acima desse ponto o ambiente começa a perder margem de segurança para "
        "crescimentos inesperados, rotinas de manutenção e variações de carga."
    )
    nota_dado = (
        f"A capacidade considerada foi {capacidade} {unidade}. A origem da série foi: {origem_serie}. "
        "Quando a série absoluta não está disponível, o relatório usa a capacidade atual e o percentual "
        "de uso para estimar o consumo em unidade de negócio. Essa condição é informada para evitar "
        "interpretação equivocada."
    )
    graficos = [
        ("A. Comportamento histórico e projeção",
         "Este gráfico compara o uso observado com a capacidade disponível, a faixa de atenção de 80% e a tendência projetada. Ele mostra se o servidor está trabalhando com folga ou se a curva se aproxima de um ponto de risco."),
        ("B. Histórico com média móvel",
         "A média móvel reduz oscilações pontuais e ajuda a enxergar a tendência principal. Para o solicitante, ela responde se o comportamento do servidor está melhorando, piorando ou permanecendo estável."),
        ("C. Comportamento ao longo do tempo",
         "Esta visão separa tendência e variações pontuais. Isso evita decisões baseadas em um pico isolado quando o comportamento geral é saudável, ou ajuda a confirmar risco quando o crescimento é contínuo."),
        ("D. Distribuição do uso",
         "O histograma mostra onde o servidor passa a maior parte do tempo. Se a maior concentração estiver perto do limite, o cenário exige atenção mesmo que existam poucos picos extremos."),
        ("E. Uso médio por horário",
         "Este gráfico identifica se existe horário mais crítico. Isso ajuda a diferenciar uso constante de períodos específicos de maior carga, como rotinas de negócio, integrações, backup ou processamento agendado."),
    ]
    estat = (
        f"A estatística reforça a leitura visual. A média indica o consumo típico; o Q3, aqui em {q3}%, "
        "mostra uma faixa onde o ambiente passa boa parte do tempo; e o pico mostra o pior momento observado. "
        f"A maior projeção identificada foi {max_fore}% quando havia forecast disponível."
    )
    conclusao = (
        f"Conclusão: {recomendacao} "
        "Essa recomendação deve ser entendida como apoio à decisão. Se o relatório indicar risco, a ação recomendada "
        "é avaliar ajuste de capacidade ou acompanhar em FUP. Se indicar otimização, a recomendação é revisar a alocação "
        "antes de ampliar recursos."
    )
    return {
        "objetivo": objetivo,
        "resumo_exec": resumo_exec,
        "nota_dado": nota_dado,
        "graficos": graficos,
        "estatistica": estat,
        "conclusao": conclusao,
        "parecer_ia": relatorio_ia or "",
    }


def _gerar_docx_relatorio_executivo(contexto: dict, relatorio_ia: str, imagens: dict, describe_df: pd.DataFrame) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.65)
    sec.right_margin = Inches(0.65)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(12)

    header = sec.header.paragraphs[0]
    header.text = "PÚBLICO"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer = sec.footer.paragraphs[0]
    footer.text = "RMC Copilot | Relatório executivo de capacidade"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title = doc.add_heading(f"Relatório de Utilização de {contexto.get('recurso')}: Servidor {contexto.get('vm')}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for label, value in [
        ("Para", contexto.get("solicitante", "Solicitante")),
        ("De", "Francisco Alves"),
        ("Solicitação", contexto.get("numero_solicitacao", "")),
        ("Data", pd.Timestamp.today().strftime("%d/%m/%Y")),
        ("Recurso", contexto.get("recurso", "")),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(str(value))

    sections = _relatorio_exec_sections(contexto, relatorio_ia)

    doc.add_heading("1. Objetivo da análise", level=2)
    doc.add_paragraph(sections["objetivo"])

    doc.add_heading("2. Resumo Executivo", level=2)
    doc.add_paragraph(sections["resumo_exec"])
    doc.add_paragraph(sections["nota_dado"])

    doc.add_heading("3. Análise Detalhada dos Gráficos", level=2)
    fig_order = [
        ("01_comparacao_previsao", sections["graficos"][0]),
        ("02_media_movel", sections["graficos"][1]),
        ("03_decomposicao", sections["graficos"][2]),
        ("04_histograma", sections["graficos"][3]),
        ("05_uso_medio_hora", sections["graficos"][4]),
    ]
    for key, (subtitulo, explicacao) in fig_order:
        doc.add_heading(subtitulo, level=3)
        if key in imagens:
            doc.add_picture(io.BytesIO(imagens[key]), width=Inches(6.4))
        doc.add_paragraph(explicacao)

    doc.add_heading("4. Análise Estatística", level=2)
    doc.add_paragraph(sections["estatistica"])
    if describe_df is not None and not describe_df.empty:
        table = doc.add_table(rows=1, cols=len(describe_df.columns))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, col in enumerate(describe_df.columns):
            hdr[i].text = str(col)
        for _, row in describe_df.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(describe_df.columns):
                cells[i].text = str(row[col])
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "F2F6FC")
                tcPr.append(shd)

    doc.add_heading("5. Conclusão e Recomendação Final", level=2)
    doc.add_paragraph(sections["conclusao"])

    doc.add_heading("6. Parecer gerado pela IA", level=2)
    _add_markdown_to_docx(doc, (sections["parecer_ia"][:7000] if sections["parecer_ia"] else "Parecer da IA não disponível nesta execução."))

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def _gerar_pdf_relatorio_executivo(contexto: dict, relatorio_ia: str, imagens: dict, describe_df: pd.DataFrame) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle, PageBreak

    bio = io.BytesIO()
    doc = SimpleDocTemplate(
        bio,
        pagesize=A4,
        rightMargin=40,
        leftMargin=40,
        topMargin=42,
        bottomMargin=36,
    )
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="Small", parent=styles["Normal"], fontSize=8, leading=10))
    styles.add(ParagraphStyle(name="BodyRMC", parent=styles["Normal"], fontSize=10, leading=14))
    styles.add(ParagraphStyle(name="H1RMC", parent=styles["Heading1"], fontSize=16, leading=20, textColor=colors.HexColor("#1F3B73")))
    styles.add(ParagraphStyle(name="H2RMC", parent=styles["Heading2"], fontSize=12, leading=16, textColor=colors.HexColor("#1F3B73")))

    story = []
    story.append(Paragraph("PÚBLICO", styles["Small"]))
    story.append(Paragraph(f"Relatório de Utilização de {contexto.get('recurso')}: Servidor {contexto.get('vm')}", styles["H1RMC"]))
    for label, value in [
        ("Para", contexto.get("solicitante", "Solicitante")),
        ("De", "Francisco Alves"),
        ("Solicitação", contexto.get("numero_solicitacao", "")),
        ("Data", pd.Timestamp.today().strftime("%d/%m/%Y")),
        ("Recurso", contexto.get("recurso", "")),
    ]:
        story.append(Paragraph(f"<b>{label}:</b> {html.escape(str(value))}", styles["BodyRMC"]))
    story.append(Spacer(1, 10))

    sections = _relatorio_exec_sections(contexto, relatorio_ia)

    story.append(Paragraph("1. Objetivo da análise", styles["H2RMC"]))
    story.append(Paragraph(html.escape(sections["objetivo"]), styles["BodyRMC"]))
    story.append(Paragraph("2. Resumo Executivo", styles["H2RMC"]))
    story.append(Paragraph(html.escape(sections["resumo_exec"]), styles["BodyRMC"]))
    story.append(Paragraph(html.escape(sections["nota_dado"]), styles["BodyRMC"]))

    story.append(Paragraph("3. Análise Detalhada dos Gráficos", styles["H2RMC"]))
    fig_order = [
        ("01_comparacao_previsao", sections["graficos"][0]),
        ("02_media_movel", sections["graficos"][1]),
        ("03_decomposicao", sections["graficos"][2]),
        ("04_histograma", sections["graficos"][3]),
        ("05_uso_medio_hora", sections["graficos"][4]),
    ]
    for key, (subtitulo, explicacao) in fig_order:
        story.append(Paragraph(html.escape(subtitulo), styles["H2RMC"]))
        if key in imagens:
            story.append(Image(io.BytesIO(imagens[key]), width=6.3*inch, height=(4.5*inch if key == "03_decomposicao" else 3.3*inch)))
        story.append(Paragraph(html.escape(explicacao), styles["BodyRMC"]))
        story.append(Spacer(1, 8))

    story.append(Paragraph("4. Análise Estatística", styles["H2RMC"]))
    story.append(Paragraph(html.escape(sections["estatistica"]), styles["BodyRMC"]))
    if describe_df is not None and not describe_df.empty:
        data = [list(map(str, describe_df.columns))]
        for _, row in describe_df.iterrows():
            data.append([str(row[c]) for c in describe_df.columns])
        tbl = Table(data, hAlign="LEFT")
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DCEBFA")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1F3B73")),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#A7B7C7")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
        ]))
        story.append(tbl)

    story.append(Paragraph("5. Conclusão e Recomendação Final", styles["H2RMC"]))
    story.append(Paragraph(html.escape(sections["conclusao"]), styles["BodyRMC"]))
    story.append(Paragraph("6. Parecer gerado pela IA", styles["H2RMC"]))
    story.extend(_markdown_to_pdf_flowables((sections["parecer_ia"] or "Parecer da IA não disponível nesta execução.")[:6000], styles))

    doc.build(story)
    bio.seek(0)
    return bio.getvalue()


def _gerar_word_pdf_relatorio_executivo(contexto: dict, relatorio_ia: str, figs_plotly: dict, describe_df: pd.DataFrame, df_hist: pd.DataFrame, df_forecast: pd.DataFrame):
    imagens = _criar_figuras_png_relatorio(df_hist, df_forecast, figs_plotly, contexto.get("recurso", ""), contexto.get("unidade", ""))
    docx_bytes = _gerar_docx_relatorio_executivo(contexto, relatorio_ia, imagens, describe_df)
    pdf_bytes = _gerar_pdf_relatorio_executivo(contexto, relatorio_ia, imagens, describe_df)
    return docx_bytes, pdf_bytes

def _render_analise_completa_recurso(numero_solicitacao: str, nome_vm: str, vm_df: pd.DataFrame, recurso: str, solicitante: str = "Solicitante"):
    st.markdown(f"### Análise completa — {recurso}")
    capacidade, unidade, origem_capacidade = _inferir_capacidade_recurso_vm(vm_df, recurso)
    run_id = st.session_state.get("resource_run_id_ativo") or st.session_state.get("execution_id_ativo")
    dias = int(st.session_state.get("historico_dias_ativo", 90) or 90)
    df_hist = _buscar_timeseries_vm_duckdb(nome_vm, recurso, capacidade, unidade, run_id=run_id, dias=dias)
    origem_serie = "DuckDB/vROps"
    if df_hist is None or df_hist.empty:
        df_hist = _serie_sintetica_percentual_vm(vm_df, recurso, capacidade, unidade)
        origem_serie = "consolidado percentual + capacidade atual"
    if df_hist is None or df_hist.empty:
        st.warning(f"Sem dados suficientes para montar a análise completa de {recurso}.")
        return
    df_forecast = _calcular_forecast_linear(df_hist, capacidade, unidade)
    if df_forecast is None:
        df_forecast = pd.DataFrame()
    c1, c2, c3, c4 = st.columns(4)
    pct = pd.to_numeric(df_hist["percentual"], errors="coerce").dropna()
    uso = pd.to_numeric(df_hist["uso"], errors="coerce").dropna()
    c1.metric("Capacidade considerada", f"{capacidade:.2f} {unidade}")
    c2.metric("Faixa de atenção", f"{(capacidade*0.8 if unidade != '%' else 80):.2f} {unidade}")
    c3.metric("Uso médio", f"{(uso.mean() if not uso.empty else 0):.2f} {unidade}")
    c4.metric("Pico observado", f"{(pct.max() if not pct.empty else 0):.2f}%")
    st.caption(f"Fonte da série: {origem_serie} | Origem da capacidade: {origem_capacidade}")
    figs = {
        "Comparação e projeção": _fig_comparacao_previsao(df_hist, df_forecast, f"Comparação e projeção — {recurso}", unidade),
        "Histórico com média móvel": _fig_media_movel(df_hist, f"Histórico com média móvel — {recurso}", unidade),
        "Comportamento ao longo do tempo": _fig_decomposicao_simples(df_hist, f"Comportamento ao longo do tempo — {recurso}", unidade),
        "Distribuição do uso": _fig_histograma(df_hist, f"Distribuição do uso — {recurso}", unidade),
        "Uso médio por horário": _fig_uso_medio_hora(df_hist, f"Uso médio por horário — {recurso}", unidade),
    }
    for fig in figs.values():
        if fig is not None:
            st.plotly_chart(fig, width="stretch")
    describe_df = _describe_analise(df_hist, unidade)
    if not describe_df.empty:
        st.markdown("#### Estatística descritiva")
        st.dataframe(describe_df, width="stretch", hide_index=True)
    status, recomendacao = _classificar_recomendacao_analise(df_hist, df_forecast)
    st.markdown(f"#### Conclusão automática: **{status}**")
    st.write(recomendacao)
    contexto = _montar_contexto_analise_completa(
        numero_solicitacao=numero_solicitacao,
        nome_vm=nome_vm,
        recurso=recurso,
        capacidade=capacidade,
        unidade=unidade,
        origem_capacidade=origem_capacidade,
        origem_serie=origem_serie,
        historico_dias=dias,
        df_hist=df_hist,
        df_forecast=df_forecast,
        describe_df=describe_df,
        status=status,
        recomendacao=recomendacao,
        solicitante=solicitante,
    )
    with st.expander(f"Relatório da Análise Individual gerado por IA — {recurso}", expanded=False):
        st.info("A IA escreve o relatório em linguagem acessível para leitores não técnicos e o pacote inclui um HTML ilustrado com as figuras desta análise.")
        if st.button(f"Gerar relatório da Análise Individual — {recurso}", key=f"btn_relatorio_completo_{_normalizar_nome_arquivo(numero_solicitacao)}_{_normalizar_nome_arquivo(nome_vm)}_{_normalizar_nome_arquivo(recurso)}"):
            relatorio, origem = _gerar_relatorio_analise_completa_ia(contexto)
            html_relatorio = _gerar_html_relatorio_ilustrado(relatorio, contexto, figs, describe_df, df_hist=df_hist, df_forecast=df_forecast)
            pacote = _gerar_pacote_relatorio_ia(numero_solicitacao, nome_vm, recurso, relatorio, html_relatorio, describe_df, df_hist, df_forecast)
            st.markdown(relatorio)
            st.download_button(
                f"Baixar relatório {recurso} (Markdown)",
                data=relatorio.encode("utf-8"),
                file_name=f"relatorio_analise_completa_{_normalizar_nome_arquivo(numero_solicitacao)}_{_normalizar_nome_arquivo(nome_vm)}_{_normalizar_nome_arquivo(recurso)}.md",
                mime="text/markdown",
                key=f"download_relatorio_md_{_normalizar_nome_arquivo(numero_solicitacao)}_{_normalizar_nome_arquivo(nome_vm)}_{_normalizar_nome_arquivo(recurso)}",
            )
            try:
                docx_bytes, pdf_bytes = _gerar_word_pdf_relatorio_executivo(contexto, relatorio, figs, describe_df, df_hist, df_forecast)
                st.download_button(
                    f"Baixar relatório executivo {recurso} (Word)",
                    data=docx_bytes,
                    file_name=f"relatorio_executivo_{_normalizar_nome_arquivo(numero_solicitacao)}_{_normalizar_nome_arquivo(nome_vm)}_{_normalizar_nome_arquivo(recurso)}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"download_relatorio_word_{_normalizar_nome_arquivo(numero_solicitacao)}_{_normalizar_nome_arquivo(nome_vm)}_{_normalizar_nome_arquivo(recurso)}",
                )
                st.download_button(
                    f"Baixar relatório executivo {recurso} (PDF)",
                    data=pdf_bytes,
                    file_name=f"relatorio_executivo_{_normalizar_nome_arquivo(numero_solicitacao)}_{_normalizar_nome_arquivo(nome_vm)}_{_normalizar_nome_arquivo(recurso)}.pdf",
                    mime="application/pdf",
                    key=f"download_relatorio_pdf_{_normalizar_nome_arquivo(numero_solicitacao)}_{_normalizar_nome_arquivo(nome_vm)}_{_normalizar_nome_arquivo(recurso)}",
                )
            except Exception as exc:
                st.warning(f"Não foi possível gerar Word/PDF nesta execução. Instale as dependências do relatório executivo. Detalhe: {exc}")

            st.download_button(
                f"Baixar relatório {recurso} ilustrado (HTML)",
                data=html_relatorio.encode("utf-8"),
                file_name=f"relatorio_ilustrado_{_normalizar_nome_arquivo(numero_solicitacao)}_{_normalizar_nome_arquivo(nome_vm)}_{_normalizar_nome_arquivo(recurso)}.html",
                mime="text/html",
                key=f"download_relatorio_html_{_normalizar_nome_arquivo(numero_solicitacao)}_{_normalizar_nome_arquivo(nome_vm)}_{_normalizar_nome_arquivo(recurso)}",
            )
            st.download_button(
                f"Baixar pacote completo {recurso} (ZIP)",
                data=pacote,
                file_name=f"pacote_relatorio_{_normalizar_nome_arquivo(numero_solicitacao)}_{_normalizar_nome_arquivo(nome_vm)}_{_normalizar_nome_arquivo(recurso)}.zip",
                mime="application/zip",
                key=f"download_relatorio_zip_{_normalizar_nome_arquivo(numero_solicitacao)}_{_normalizar_nome_arquivo(nome_vm)}_{_normalizar_nome_arquivo(recurso)}",
            )
            st.caption(f"Fonte do texto: {origem}")



def _analise_total_status(resultados: list):
    statuses = [str(r.get("status", "")).upper() for r in resultados]
    if any(s == "CRÍTICO" or s == "CRITICO" for s in statuses):
        return "CRÍTICO"
    if any(s == "RISCO" for s in statuses):
        return "RISCO"
    if any(s == "OTIMIZAÇÃO" or s == "OTIMIZACAO" for s in statuses):
        return "OTIMIZAÇÃO"
    if any(s == "SEM_DADOS" for s in statuses):
        return "ATENÇÃO"
    return "OK"


def _texto_recomendacao_total(status_total: str, resultados: list):
    criticos = [r for r in resultados if str(r.get("status", "")).upper() in ["CRÍTICO", "CRITICO"]]
    riscos = [r for r in resultados if str(r.get("status", "")).upper() == "RISCO"]
    otmz = [r for r in resultados if str(r.get("status", "")).upper() in ["OTIMIZAÇÃO", "OTIMIZACAO"]]
    if criticos:
        nomes = ", ".join(r.get("recurso", "recurso") for r in criticos)
        return f"Ação prioritária no(s) recurso(s): {nomes}. O relatório total não pode ser tratado como OK enquanto houver recurso crítico. Recomenda-se tratar o gargalo antes de encerrar a solicitação."
    if riscos:
        nomes = ", ".join(r.get("recurso", "recurso") for r in riscos)
        return f"Há risco relevante em {nomes}. Recomenda-se planejar ajuste ou manter monitoramento próximo, pois a margem de segurança pode ser consumida no período projetado."
    if otmz:
        nomes = ", ".join(r.get("recurso", "recurso") for r in otmz)
        return f"Não há risco crítico consolidado, mas há oportunidade de otimização em {nomes}. Recomenda-se avaliar rightsizing antes de qualquer expansão."
    return "CPU, memória e disco não apresentam risco consolidado relevante no período avaliado. Recomenda-se manter monitoramento periódico."


def _coletar_resultado_recurso_total(numero_solicitacao: str, nome_vm: str, vm_df: pd.DataFrame, recurso: str, solicitante: str):
    capacidade, unidade, origem_capacidade = _inferir_capacidade_recurso_vm(vm_df, recurso)
    run_id = st.session_state.get("resource_run_id_ativo") or st.session_state.get("execution_id_ativo")
    dias = int(st.session_state.get("historico_dias_ativo", 90) or 90)
    df_hist = _buscar_timeseries_vm_duckdb(nome_vm, recurso, capacidade, unidade, run_id=run_id, dias=dias)
    origem_serie = "DuckDB/vROps"
    if df_hist is None or df_hist.empty:
        df_hist = _serie_sintetica_percentual_vm(vm_df, recurso, capacidade, unidade)
        origem_serie = "consolidado percentual + capacidade atual"
    if df_hist is None or df_hist.empty:
        df_hist = pd.DataFrame()
    df_forecast = _calcular_forecast_linear(df_hist, capacidade, unidade) if df_hist is not None and not df_hist.empty else pd.DataFrame()
    describe_df = _describe_analise(df_hist, unidade) if df_hist is not None and not df_hist.empty else pd.DataFrame()
    status, recomendacao = _classificar_recomendacao_analise(df_hist, df_forecast)
    resumo = _extrair_resumo_executivo(df_hist, df_forecast) if df_hist is not None and not df_hist.empty else {}
    contexto = _montar_contexto_analise_completa(
        numero_solicitacao=numero_solicitacao,
        nome_vm=nome_vm,
        recurso=recurso,
        capacidade=capacidade,
        unidade=unidade,
        origem_capacidade=origem_capacidade,
        origem_serie=origem_serie,
        historico_dias=dias,
        df_hist=df_hist,
        df_forecast=df_forecast,
        describe_df=describe_df,
        status=status,
        recomendacao=recomendacao,
        solicitante=solicitante,
    )
    figs = {
        "Comparação e projeção": _fig_comparacao_previsao(df_hist, df_forecast, f"Comparação e projeção — {recurso}", unidade),
        "Histórico com média móvel": _fig_media_movel(df_hist, f"Histórico com média móvel — {recurso}", unidade),
    }
    return {
        "recurso": recurso,
        "status": status,
        "recomendacao": recomendacao,
        "resumo": resumo,
        "capacidade": capacidade,
        "unidade": unidade,
        "origem_serie": origem_serie,
        "origem_capacidade": origem_capacidade,
        "df_hist": df_hist,
        "df_forecast": df_forecast,
        "describe_df": describe_df,
        "contexto": contexto,
        "figs": figs,
    }


def _fig_total_consolidado(resultados: list):
    rows = []
    for r in resultados:
        resumo = r.get("resumo", {}) or {}
        rows.append({
            "recurso": r.get("recurso", "N/C"),
            "media": resumo.get("media_pct") or 0,
            "pico": resumo.get("pico_pct") or 0,
            "forecast": resumo.get("max_forecast_pct") or 0,
        })
    df = pd.DataFrame(rows)
    fig = go.Figure()
    if df.empty:
        return None
    fig.add_trace(go.Bar(x=df["recurso"], y=df["media"], name="Média", marker=dict(color=RMC_RELATORIO_CORES["media"])))
    fig.add_trace(go.Bar(x=df["recurso"], y=df["pico"], name="Pico observado", marker=dict(color=RMC_RELATORIO_CORES["pico"])))
    fig.add_trace(go.Bar(x=df["recurso"], y=df["forecast"], name="Pico projetado", marker=dict(color=RMC_RELATORIO_CORES["forecast"])))
    ymax = max(105, float(df[["media", "pico", "forecast"]].max().max()) + 8)
    fig.add_hrect(y0=0, y1=80, fillcolor=RMC_RELATORIO_CORES["fundo_ok"], line_width=0)
    fig.add_hrect(y0=80, y1=90, fillcolor=RMC_RELATORIO_CORES["fundo_atencao"], line_width=0)
    fig.add_hrect(y0=90, y1=ymax, fillcolor=RMC_RELATORIO_CORES["fundo_risco"], line_width=0)
    fig.add_hline(y=80, line_dash="dash", line_color=RMC_RELATORIO_CORES["threshold"], annotation_text="Faixa de atenção 80%", annotation_position="top left")
    fig.update_layout(
        title=dict(text="Visão consolidada dos recursos", x=0.02, font=dict(size=18, color=RMC_RELATORIO_CORES["texto"])),
        template="plotly_white", barmode="group", height=430,
        yaxis=dict(title="Percentual de uso (%)", range=[0, ymax], gridcolor=RMC_RELATORIO_CORES["grid"]),
        xaxis=dict(title="Recurso"), plot_bgcolor="white", paper_bgcolor="white",
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
    )
    return fig


def _fig_total_consolidado_png(resultados: list) -> bytes:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import numpy as np
    rows = []
    for r in resultados:
        resumo = r.get("resumo", {}) or {}
        rows.append([r.get("recurso", "N/C"), resumo.get("media_pct") or 0, resumo.get("pico_pct") or 0, resumo.get("max_forecast_pct") or 0])
    df = pd.DataFrame(rows, columns=["Recurso", "Média", "Pico observado", "Pico projetado"])
    bio = io.BytesIO()
    if df.empty:
        return b""
    x = np.arange(len(df))
    width = 0.24
    fig, ax = plt.subplots(figsize=(9, 4.8), facecolor="white")
    ymax = max(105, float(df[["Média", "Pico observado", "Pico projetado"]].max().max()) + 8)
    ax.axhspan(0, 80, color="#2CA02C", alpha=0.07)
    ax.axhspan(80, 90, color="#FFC107", alpha=0.12)
    ax.axhspan(90, ymax, color="#D62728", alpha=0.08)
    ax.bar(x - width, df["Média"], width, label="Média", color="#4E79A7")
    ax.bar(x, df["Pico observado"], width, label="Pico observado", color="#9467BD")
    ax.bar(x + width, df["Pico projetado"], width, label="Pico projetado", color="#F28E2B")
    ax.axhline(80, color="#D62728", linestyle="--", linewidth=2, label="Faixa de atenção 80%")
    ax.set_title("Visão consolidada dos recursos", fontsize=13, weight="bold", color="#24324A")
    ax.set_ylabel("Percentual de uso (%)")
    ax.set_xticks(x)
    ax.set_xticklabels(df["Recurso"])
    ax.set_ylim(0, ymax)
    ax.grid(True, axis="y", color="#E6ECF5")
    ax.legend(loc="best", frameon=True)
    plt.tight_layout()
    plt.savefig(bio, format="png", dpi=170, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    bio.seek(0)
    return bio.getvalue()


def _montar_contexto_total_ia(numero_solicitacao: str, nome_vm: str, solicitante: str, resultados: list):
    status_total = _analise_total_status(resultados)
    recomendacao_total = _texto_recomendacao_total(status_total, resultados)
    recursos = []
    for r in resultados:
        resumo = r.get("resumo", {}) or {}
        recursos.append({
            "recurso": r.get("recurso"),
            "status": r.get("status"),
            "media_pct": resumo.get("media_pct"),
            "pico_pct": resumo.get("pico_pct"),
            "q3_pct": resumo.get("q3_pct"),
            "max_forecast_pct": resumo.get("max_forecast_pct"),
            "pontos_hist_acima_80": resumo.get("pontos_hist_acima_80"),
            "pontos_forecast_acima_80": resumo.get("pontos_forecast_acima_80"),
            "recomendacao_calculada": r.get("recomendacao"),
            "origem_serie": r.get("origem_serie"),
        })
    return {
        "numero_solicitacao": numero_solicitacao,
        "vm": nome_vm,
        "solicitante": solicitante or "Solicitante",
        "historico_dias": int(st.session_state.get("historico_dias_ativo", 90) or 90),
        "status_total": status_total,
        "recomendacao_total": recomendacao_total,
        "recursos": recursos,
        "regra_de_consolidacao": "Se qualquer recurso estiver CRÍTICO, a situação consolidada é CRÍTICO. Se qualquer recurso estiver em RISCO, a situação consolidada é RISCO. Nunca classificar o relatório total como OK quando CPU, memória ou disco estiver crítico.",
    }


def _texto_tem_contradicao_total(contexto_total: dict, texto: str) -> bool:
    status = _normalizar_status_txt(contexto_total.get("status_total"))
    low = str(texto or "").lower()
    if not low.strip():
        return True
    if status in ["CRÍTICO", "RISCO"]:
        proibidos = ["sem risco consolidado", "não apresentam risco", "não há risco", "situação consolidada: ok", "pode encerrar", "apenas manter monitoramento"]
        if any(p in low for p in proibidos):
            return True
    # Cada recurso crítico precisa aparecer no texto.
    criticos = [str(r.get("recurso", "")).lower() for r in contexto_total.get("recursos", []) if _normalizar_status_txt(r.get("status")) == "CRÍTICO"]
    if criticos and not all(c in low for c in criticos):
        return True
    return False


def _montar_texto_controlado_total(contexto_total: dict) -> dict:
    vm = contexto_total.get("vm")
    status = _normalizar_status_txt(contexto_total.get("status_total"))
    recursos = contexto_total.get("recursos", []) or []
    criticos = [r for r in recursos if _normalizar_status_txt(r.get("status")) == "CRÍTICO"]
    riscos = [r for r in recursos if _normalizar_status_txt(r.get("status")) == "RISCO"]
    otmz = [r for r in recursos if _normalizar_status_txt(r.get("status")) == "OTIMIZAÇÃO"]
    resumo = f"A análise consolidada da VM {vm} classifica a situação geral como {status}. "
    if criticos:
        nomes = ", ".join(r.get("recurso", "recurso") for r in criticos)
        resumo += f"O principal ponto de atenção é {nomes}, pois pelo menos um recurso foi classificado como crítico. O relatório total não deve ser tratado como OK enquanto houver recurso crítico."
    elif riscos:
        nomes = ", ".join(r.get("recurso", "recurso") for r in riscos)
        resumo += f"Há risco relevante em {nomes}. A recomendação é planejar ajuste ou acompanhar de perto a evolução do consumo."
    elif otmz:
        nomes = ", ".join(r.get("recurso", "recurso") for r in otmz)
        resumo += f"Não há gargalo crítico consolidado, mas há oportunidade de otimização em {nomes}."
    else:
        resumo += "CPU, memória e disco estão dentro de uma faixa operacional aceitável, mantendo-se o monitoramento periódico."
    leitura = []
    for r in recursos:
        leitura.append(f"{r.get('recurso')}: status {r.get('status')}, média {r.get('media_pct')}%, pico {r.get('pico_pct')}% e maior projeção {r.get('max_forecast_pct')}%.")
    return {
        "resumo_executivo": resumo,
        "principal_ponto_atencao": contexto_total.get("recomendacao_total", ""),
        "leitura_por_recurso": " ".join(leitura),
        "impacto_operacional": "A leitura consolidada ajuda o solicitante a entender se o servidor tem folga, se exige ampliação ou se existe oportunidade de revisão de recursos. A IA apenas analisa e recomenda; nenhuma ação operacional é executada automaticamente.",
        "recomendacao_final": contexto_total.get("recomendacao_total", ""),
    }


def _gerar_relatorio_total_ia(contexto_total: dict):
    prompt = f"""
Você é a IA do RMC Copilot, especialista em VMware e capacity planning.

Gere a análise executiva TOTAL da Análise Individual de VM, consolidando CPU, Memória e Disco.

IMPORTANTE:
- Isto NÃO é FUP.
- A IA só analisa, explica e recomenda. Não toma ação operacional.
- Use somente os números do contexto. Não invente dados, causa, aplicação, horário ou percentual.
- Se algum recurso estiver CRÍTICO, o relatório total deve ser CRÍTICO e deve citar esse recurso.
- Nunca diga que o relatório total está OK, sem risco ou encerrado se qualquer recurso estiver CRÍTICO ou em RISCO.
- Explique para uma pessoa leiga em capacity planning.
- Responda SOMENTE JSON válido, sem Markdown.

Formato obrigatório:
{{
  "resumo_executivo": "texto",
  "principal_ponto_atencao": "texto",
  "leitura_por_recurso": "texto",
  "impacto_operacional": "texto",
  "recomendacao_final": "texto"
}}

Contexto:
{json.dumps(contexto_total, ensure_ascii=False, default=str, indent=2)}
"""
    origem = "IA"
    try:
        resposta = _chamar_llm_rmc(prompt)
        dados = _ia_json_extract(resposta)
        obrig = ["resumo_executivo", "principal_ponto_atencao", "leitura_por_recurso", "impacto_operacional", "recomendacao_final"]
        if not isinstance(dados, dict) or any(not str(dados.get(k, "")).strip() for k in obrig):
            raise ValueError("JSON incompleto")
        unido = "\n".join(str(dados.get(k, "")) for k in obrig)
        if _texto_tem_contradicao_total(contexto_total, unido):
            raise ValueError("resposta da IA contraditória ao status consolidado")
    except Exception as exc:
        dados = _montar_texto_controlado_total(contexto_total)
        origem = f"IA controlada ({exc})"

    texto = f"""# Relatório Executivo Total de Capacidade: Servidor {contexto_total.get('vm')}

## 1. Resumo executivo consolidado
{dados.get('resumo_executivo')}

## 2. Principal ponto de atenção
{dados.get('principal_ponto_atencao')}

## 3. Leitura por recurso
{dados.get('leitura_por_recurso')}

## 4. Impacto operacional
{dados.get('impacto_operacional')}

## 5. Recomendação final
{dados.get('recomendacao_final')}
"""
    return texto, origem


def _gerar_html_total_executivo(contexto_total: dict, relatorio_ia: str, resultados: list, fig_total):
    resumo_rows = "".join([
        f"<tr><td>{html.escape(str(r.get('recurso')))}</td><td>{html.escape(str(r.get('status')))}</td><td>{(r.get('resumo') or {}).get('media_pct', 'N/D')}</td><td>{(r.get('resumo') or {}).get('pico_pct', 'N/D')}</td><td>{(r.get('resumo') or {}).get('max_forecast_pct', 'N/D')}</td></tr>"
        for r in resultados
    ])
    recursos_html = []
    for r in resultados:
        figs = r.get("figs") or {}
        recursos_html.append(f"<h2>{html.escape(str(r.get('recurso')))}</h2><p><strong>Status:</strong> {html.escape(str(r.get('status')))}<br>{html.escape(str(r.get('recomendacao')))}</p>")
        for nome, fig in figs.items():
            if fig is not None:
                recursos_html.append(f"<h3>{html.escape(nome)}</h3>{_fig_to_html(fig)}")
    return f"""
<!DOCTYPE html>
<html lang="pt-br">
<head>
<meta charset="utf-8" />
<title>Relatório total - {html.escape(str(contexto_total.get('vm')))}</title>
<script src="https://cdn.plot.ly/plotly-2.35.2.min.js"></script>
<style>
body {{ font-family: Arial, sans-serif; margin: 24px; color:#24324A; background:#F6F8FC; }}
.card {{ background:#fff; border-radius:14px; padding:18px 22px; margin-bottom:18px; box-shadow:0 2px 10px rgba(0,0,0,.06); }}
h1,h2,h3 {{ color:#1F3B73; }}
table {{ border-collapse:collapse; width:100%; }} th,td {{ border-bottom:1px solid #E5E7EB; padding:8px 10px; text-align:left; }} th {{ background:#DCEBFA; }}
</style>
</head>
<body>
<div class="card"><h1>Relatório Executivo Total de Capacidade</h1>
<p><strong>Servidor/VM:</strong> {html.escape(str(contexto_total.get('vm')))}<br>
<strong>Solicitação:</strong> {html.escape(str(contexto_total.get('numero_solicitacao')))}<br>
<strong>Solicitante:</strong> {html.escape(str(contexto_total.get('solicitante')))}<br>
<strong>Situação consolidada:</strong> {html.escape(str(contexto_total.get('status_total')))}</p></div>
<div class="card"><h2>Análise consolidada gerada pela IA</h2>{_markdown_to_html_relatorio(relatorio_ia)}</div>
<div class="card"><h2>Gráfico consolidado</h2>{_fig_to_html(fig_total)}</div>
<div class="card"><h2>Quadro-resumo por recurso</h2><table><tr><th>Recurso</th><th>Status</th><th>Média %</th><th>Pico %</th><th>Forecast máx. %</th></tr>{resumo_rows}</table></div>
<div class="card"><h2>Detalhamento visual por recurso</h2>{''.join(recursos_html)}</div>
</body></html>
"""


def _gerar_docx_total_executivo(contexto_total: dict, relatorio_ia: str, resultados: list, imagem_total: bytes, imagens_recursos: dict) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.55); sec.bottom_margin = Inches(0.55); sec.left_margin = Inches(0.65); sec.right_margin = Inches(0.65)
    doc.styles["Normal"].font.name = "Arial"; doc.styles["Normal"].font.size = Pt(10)
    sec.header.paragraphs[0].text = "PÚBLICO"
    sec.footer.paragraphs[0].text = "RMC Copilot | Relatório executivo total de capacidade"
    sec.footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(f"Relatório Executivo Total de Capacidade: Servidor {contexto_total.get('vm')}", level=1)
    for label, value in [("Para", contexto_total.get("solicitante")), ("De", "Francisco Alves"), ("Solicitação", contexto_total.get("numero_solicitacao")), ("Data", pd.Timestamp.today().strftime("%d/%m/%Y")), ("Situação consolidada", contexto_total.get("status_total"))]:
        p = doc.add_paragraph(); p.add_run(f"{label}: ").bold = True; p.add_run(str(value))
    doc.add_heading("1. Análise consolidada gerada pela IA", level=2)
    _add_markdown_to_docx(doc, relatorio_ia)
    doc.add_heading("2. Gráfico consolidado", level=2)
    if imagem_total:
        doc.add_picture(io.BytesIO(imagem_total), width=Inches(6.4))
    doc.add_heading("3. Quadro-resumo por recurso", level=2)
    table = doc.add_table(rows=1, cols=5); table.style = "Table Grid"
    for i, h in enumerate(["Recurso", "Status", "Média %", "Pico %", "Forecast máx. %"]): table.rows[0].cells[i].text = h
    for r in resultados:
        resumo = r.get("resumo") or {}; cells = table.add_row().cells
        vals = [r.get("recurso"), r.get("status"), resumo.get("media_pct"), resumo.get("pico_pct"), resumo.get("max_forecast_pct")]
        for i, v in enumerate(vals): cells[i].text = str(v)
    doc.add_heading("4. Detalhamento visual por recurso", level=2)
    for r in resultados:
        doc.add_heading(str(r.get("recurso")), level=3)
        doc.add_paragraph(f"Status: {r.get('status')} | {r.get('recomendacao')}")
        for key, img in (imagens_recursos.get(r.get("recurso"), {}) or {}).items():
            if key in ["01_comparacao_previsao", "02_media_movel"] and img:
                doc.add_picture(io.BytesIO(img), width=Inches(6.2))
    bio = io.BytesIO(); doc.save(bio); bio.seek(0); return bio.getvalue()


def _gerar_pdf_total_executivo(contexto_total: dict, relatorio_ia: str, resultados: list, imagem_total: bytes, imagens_recursos: dict) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
    bio = io.BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=A4, rightMargin=40, leftMargin=40, topMargin=42, bottomMargin=36)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="BodyRMC", parent=styles["Normal"], fontSize=10, leading=14))
    styles.add(ParagraphStyle(name="H1RMC", parent=styles["Heading1"], fontSize=16, leading=20, textColor=colors.HexColor("#1F3B73")))
    styles.add(ParagraphStyle(name="H2RMC", parent=styles["Heading2"], fontSize=12, leading=16, textColor=colors.HexColor("#1F3B73")))
    story = [Paragraph("PÚBLICO", styles["BodyRMC"]), Paragraph(f"Relatório Executivo Total de Capacidade: Servidor {html.escape(str(contexto_total.get('vm')))}", styles["H1RMC"])]
    for label, value in [("Para", contexto_total.get("solicitante")), ("De", "Francisco Alves"), ("Solicitação", contexto_total.get("numero_solicitacao")), ("Data", pd.Timestamp.today().strftime("%d/%m/%Y")), ("Situação consolidada", contexto_total.get("status_total"))]:
        story.append(Paragraph(f"<b>{label}:</b> {html.escape(str(value))}", styles["BodyRMC"]))
    story.append(Spacer(1, 8))
    story.append(Paragraph("1. Análise consolidada gerada pela IA", styles["H2RMC"]))
    story.extend(_markdown_to_pdf_flowables(relatorio_ia, styles))
    story.append(Paragraph("2. Gráfico consolidado", styles["H2RMC"]))
    if imagem_total:
        story.append(Image(io.BytesIO(imagem_total), width=6.3*inch, height=3.4*inch))
    story.append(Paragraph("3. Quadro-resumo por recurso", styles["H2RMC"]))
    rows = [["Recurso", "Status", "Média %", "Pico %", "Forecast máx. %"]]
    for r in resultados:
        resumo = r.get("resumo") or {}; rows.append([str(r.get("recurso")), str(r.get("status")), str(resumo.get("media_pct")), str(resumo.get("pico_pct")), str(resumo.get("max_forecast_pct"))])
    tbl = Table(rows, hAlign="LEFT")
    tbl.setStyle(TableStyle([("BACKGROUND", (0,0), (-1,0), colors.HexColor("#DCEBFA")), ("GRID", (0,0), (-1,-1), 0.25, colors.HexColor("#A7B7C7")), ("FONTSIZE", (0,0), (-1,-1), 8)]))
    story.append(tbl)
    story.append(Paragraph("4. Detalhamento visual por recurso", styles["H2RMC"]))
    for r in resultados:
        story.append(Paragraph(str(r.get("recurso")), styles["H2RMC"]))
        story.append(Paragraph(html.escape(f"Status: {r.get('status')} | {r.get('recomendacao')}"), styles["BodyRMC"]))
        imgs = imagens_recursos.get(r.get("recurso"), {}) or {}
        if imgs.get("01_comparacao_previsao"):
            story.append(Image(io.BytesIO(imgs["01_comparacao_previsao"]), width=6.2*inch, height=3.2*inch))
    doc.build(story); bio.seek(0); return bio.getvalue()


def _render_relatorio_total_ia(numero_solicitacao: str, nome_vm: str, vm_df: pd.DataFrame, solicitante: str):
    st.markdown("### Relatório consolidado da Análise Individual — IA")
    st.warning("Este relatório consolida CPU, Memória e Disco. Se qualquer recurso estiver crítico, a situação consolidada também será crítica.")
    resultados = [_coletar_resultado_recurso_total(numero_solicitacao, nome_vm, vm_df, r, solicitante) for r in ["CPU", "Memória", "Disco"]]
    contexto_total = _montar_contexto_total_ia(numero_solicitacao, nome_vm, solicitante, resultados)
    fig_total = _fig_total_consolidado(resultados)
    if fig_total is not None:
        st.plotly_chart(fig_total, width="stretch")
    quadro = []
    for r in resultados:
        resumo = r.get("resumo") or {}
        quadro.append({"Recurso": r.get("recurso"), "Status": r.get("status"), "Média %": resumo.get("media_pct"), "Pico %": resumo.get("pico_pct"), "Forecast máx. %": resumo.get("max_forecast_pct")})
    st.dataframe(pd.DataFrame(quadro), width="stretch", hide_index=True)
    if st.button("Gerar relatório consolidado pela IA", key=f"btn_relatorio_total_ia_{_normalizar_nome_arquivo(numero_solicitacao)}_{_normalizar_nome_arquivo(nome_vm)}"):
        relatorio_ia, origem = _gerar_relatorio_total_ia(contexto_total)
        st.markdown(relatorio_ia)
        imagem_total = _fig_total_consolidado_png(resultados)
        imagens_recursos = {r.get("recurso"): _criar_figuras_png_relatorio(r.get("df_hist"), r.get("df_forecast"), r.get("figs"), r.get("recurso"), r.get("unidade")) for r in resultados}
        html_total = _gerar_html_total_executivo(contexto_total, relatorio_ia, resultados, fig_total)
        try:
            docx_total = _gerar_docx_total_executivo(contexto_total, relatorio_ia, resultados, imagem_total, imagens_recursos)
            pdf_total = _gerar_pdf_total_executivo(contexto_total, relatorio_ia, resultados, imagem_total, imagens_recursos)
            st.download_button("Baixar relatório total executivo (Word)", data=docx_total, file_name=f"relatorio_total_{_normalizar_nome_arquivo(numero_solicitacao)}_{_normalizar_nome_arquivo(nome_vm)}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"download_total_word_{_normalizar_nome_arquivo(numero_solicitacao)}_{_normalizar_nome_arquivo(nome_vm)}")
            st.download_button("Baixar relatório total executivo (PDF)", data=pdf_total, file_name=f"relatorio_total_{_normalizar_nome_arquivo(numero_solicitacao)}_{_normalizar_nome_arquivo(nome_vm)}.pdf", mime="application/pdf", key=f"download_total_pdf_{_normalizar_nome_arquivo(numero_solicitacao)}_{_normalizar_nome_arquivo(nome_vm)}")
        except Exception as exc:
            st.warning(f"Não foi possível gerar Word/PDF do relatório total: {exc}")
        st.download_button("Baixar relatório total executivo (HTML)", data=html_total.encode("utf-8"), file_name=f"relatorio_total_{_normalizar_nome_arquivo(numero_solicitacao)}_{_normalizar_nome_arquivo(nome_vm)}.html", mime="text/html", key=f"download_total_html_{_normalizar_nome_arquivo(numero_solicitacao)}_{_normalizar_nome_arquivo(nome_vm)}")
        st.caption(f"Fonte do texto consolidado: {origem}")

def _render_analise_completa_vm(numero_solicitacao: str, nome_vm: str, vm_df: pd.DataFrame, recurso_selecionado: str, solicitante: str = "Solicitante"):
    st.markdown('<div class="section-title">Análise Individual de VM — modelo SOL1809645</div>', unsafe_allow_html=True)
    st.caption("Esta visão tenta usar série real do DuckDB/vROps. Quando ela não existir, a IA usa o percentual consolidado combinado com a capacidade atual e informa isso no relatório.")
    opcoes_dias = [30, 60, 90, 180, 365]
    atual = int(st.session_state.get("historico_dias_ativo", 90) or 90)
    idx = opcoes_dias.index(atual) if atual in opcoes_dias else 2
    col_a, col_b = st.columns([1, 2])
    dias_sel = col_a.selectbox("Período do histórico", opcoes_dias, index=idx, key=f"hist_analise_completa_{_normalizar_nome_arquivo(numero_solicitacao)}_{_normalizar_nome_arquivo(nome_vm)}")
    st.session_state["historico_dias_ativo"] = int(dias_sel)
    col_b.caption("Escolha o período que fará parte dos gráficos e do relatório entregue ao solicitante.")
    recursos = ["CPU", "Memória", "Disco"]
    if recurso_selecionado and recurso_selecionado not in ["Todos", "Geral"]:
        r = str(recurso_selecionado)
        if "cpu" in r.lower():
            recursos = ["CPU"]
        elif "mem" in r.lower():
            recursos = ["Memória"]
        elif "disk" in r.lower() or "disco" in r.lower():
            recursos = ["Disco"]
    tabs = st.tabs(["Relatório total"] + recursos)
    with tabs[0]:
        _render_relatorio_total_ia(numero_solicitacao, nome_vm, vm_df, solicitante=solicitante)
    for tab, recurso in zip(tabs[1:], recursos):
        with tab:
            _render_analise_completa_recurso(numero_solicitacao, nome_vm, vm_df, recurso, solicitante=solicitante)


# =============================================================================
# HOTFIX 15F.10.23 — LLM assistida, threshold explícito e downloads CSV
# =============================================================================
def _to_csv_bytes(df: pd.DataFrame) -> bytes:
    if df is None:
        df = pd.DataFrame()
    return df.to_csv(index=False, sep=";", encoding="utf-8-sig").encode("utf-8-sig")


def _normalizar_nome_arquivo(texto: str) -> str:
    s = str(texto or "arquivo").strip().lower()
    s = re.sub(r"[^a-z0-9A-Z_-]+", "_", s)
    return s.strip("_") or "arquivo"


def _colunas_forecast_por_recurso(recurso: str):
    recurso_l = str(recurso or "").lower()
    if "cpu" in recurso_l:
        return [
            ("Atual", "cpu_p95_pct"),
            ("30 dias", "cpu_forecast_30d"),
            ("60 dias", "cpu_forecast_60d"),
            ("90 dias", "cpu_forecast_90d"),
        ]
    if "mem" in recurso_l:
        return [
            ("Atual", "mem_p95_pct"),
            ("30 dias", "mem_forecast_30d"),
            ("60 dias", "mem_forecast_60d"),
            ("90 dias", "mem_forecast_90d"),
        ]
    return [
        ("Atual", "last_used_pct"),
        ("30 dias", "forecast_30d"),
        ("60 dias", "forecast_60d"),
        ("90 dias", "forecast_90d"),
    ]


def _adicionar_threshold_fup(df: pd.DataFrame, recurso: str, limite: float) -> pd.DataFrame:
    if df is None or df.empty:
        return df

    out = df.copy()
    colunas = _colunas_forecast_por_recurso(recurso)

    for _, col in colunas:
        if col in out.columns:
            out[col] = pd.to_numeric(out[col], errors="coerce")

    momentos = []
    maximos = []
    confiancas = []

    for _, row in out.iterrows():
        momento = "N/C"
        max_val = None
        valores_validos = 0

        for label, col in colunas:
            if col not in out.columns:
                continue
            val = row.get(col)
            if pd.notna(val):
                valores_validos += 1
                try:
                    fval = float(val)
                    if max_val is None or fval > max_val:
                        max_val = fval
                    if momento == "N/C" and fval >= float(limite):
                        momento = label
                except Exception:
                    pass

        if valores_validos >= 3:
            conf = "Média"
        elif valores_validos >= 1:
            conf = "Baixa"
        else:
            conf = "Sem dados"

        momentos.append(momento)
        maximos.append(max_val)
        confiancas.append(conf)

    out["threshold_%"] = float(limite)
    out["ultrapassa_threshold_em"] = momentos
    out["maior_previsao_%"] = maximos
    out["confianca_forecast"] = confiancas
    return out


def _render_download_csv(label: str, df: pd.DataFrame, filename: str, key: str):
    if df is None or df.empty:
        return
    st.download_button(
        label=label,
        data=_to_csv_bytes(df),
        file_name=filename,
        mime="text/csv",
        key=key,
    )


def _resumo_threshold_textual(df: pd.DataFrame, recurso: str, limite: float) -> str:
    if df is None or df.empty:
        return f"{recurso}: nenhum item acima de {limite:.0f}%."
    total = len(df)
    atual = int((df.get("ultrapassa_threshold_em", pd.Series(dtype=str)).astype(str) == "Atual").sum()) if "ultrapassa_threshold_em" in df.columns else 0
    em30 = int((df.get("ultrapassa_threshold_em", pd.Series(dtype=str)).astype(str) == "30 dias").sum()) if "ultrapassa_threshold_em" in df.columns else 0
    em60 = int((df.get("ultrapassa_threshold_em", pd.Series(dtype=str)).astype(str) == "60 dias").sum()) if "ultrapassa_threshold_em" in df.columns else 0
    em90 = int((df.get("ultrapassa_threshold_em", pd.Series(dtype=str)).astype(str) == "90 dias").sum()) if "ultrapassa_threshold_em" in df.columns else 0
    return f"{recurso}: {total} itens; Atual={atual}; 30d={em30}; 60d={em60}; 90d={em90}; limite={limite:.0f}%."


def _montar_contexto_relatorio_fup(vm_nome: str, recurso: str, origem: str, limite: float, vm_df: pd.DataFrame, tabela_forecast: pd.DataFrame = None, df_particoes: pd.DataFrame = None) -> dict:
    contexto = {
        "vm": str(vm_nome or ""),
        "recurso": str(recurso or ""),
        "origem": str(origem or ""),
        "threshold_percentual": float(limite) if limite is not None else None,
        "run_id": str(st.session_state.get("resource_run_id_ativo") or st.session_state.get("execution_id_ativo") or ""),
        "historico_dias": int(st.session_state.get("historico_dias_ativo", 90) or 90),
        "forecast_horizons": st.session_state.get("forecast_horizons_ativo", [30, 60, 90]),
        "numero_solicitacao": str(st.session_state.get("solicitacao_analise_individual_ativa", "") or ""),
        "vm_linha": {},
        "forecast": [],
        "particoes": [],
    }

    try:
        if vm_df is not None and not vm_df.empty:
            contexto["vm_linha"] = {
                k: (None if pd.isna(v) else v)
                for k, v in vm_df.iloc[0].to_dict().items()
                if k in [
                    "cluster", "host", "vm", "cpu_p95_pct", "mem_p95_pct", "disk_p95_pct",
                    "cpu_forecast_30d", "cpu_forecast_60d", "cpu_forecast_90d",
                    "mem_forecast_30d", "mem_forecast_60d", "mem_forecast_90d",
                    "prioridade_final", "acao_final", "recomendacao_final", "criticidade_futura"
                ]
            }
    except Exception:
        pass

    try:
        if tabela_forecast is not None and not tabela_forecast.empty:
            contexto["forecast"] = tabela_forecast.head(20).to_dict(orient="records")
    except Exception:
        pass

    try:
        if df_particoes is not None and not df_particoes.empty:
            cols = [c for c in ["cluster", "host", "datastore", "vm", "partition_name", "last_used_pct", "forecast_30d", "forecast_60d", "forecast_90d", "ultrapassa_threshold_em", "maior_previsao_%", "status_fup", "acao_fup"] if c in df_particoes.columns]
            contexto["particoes"] = df_particoes[cols].head(30).to_dict(orient="records")
    except Exception:
        pass

    return contexto


def _relatorio_deterministico_fup(contexto: dict) -> str:
    vm = contexto.get("vm", "N/C")
    recurso = contexto.get("recurso", "N/C")
    limite = contexto.get("threshold_percentual", "N/C")
    origem = contexto.get("origem", "FUP")
    linha = contexto.get("vm_linha", {}) or {}

    partes = []
    partes.append(f"# Relatório demonstrativo FUP — {vm}")
    partes.append("")
    partes.append(f"**Origem:** {origem}")
    partes.append(f"**Recurso:** {recurso}")
    partes.append(f"**Threshold:** {limite}%")
    partes.append(f"**Run:** {contexto.get('run_id', 'N/C')}")
    partes.append("")
    partes.append("## Leitura executiva")
    partes.append(
        "A VM foi selecionada no FUP por apresentar uso atual ou tendência de capacidade acima do threshold definido. "
        "Quando o histórico individual não é suficiente para forecast estatístico robusto, a análise deve ser tratada como assistida e com confiança reduzida."
    )
    partes.append("")
    partes.append("## Indicadores principais")
    for k in ["cluster", "host", "cpu_p95_pct", "mem_p95_pct", "disk_p95_pct", "cpu_forecast_60d", "mem_forecast_60d", "prioridade_final", "acao_final", "criticidade_futura"]:
        if k in linha:
            partes.append(f"- **{k}:** {linha.get(k)}")

    partes.append("")
    partes.append("## Observação sobre forecast")
    partes.append(
        "A LLM não deve inventar valores de forecast. Ela interpreta os sinais disponíveis, aponta lacunas de dados, "
        "explica o risco e recomenda acompanhamento. O número final deve continuar vindo do histórico real, quando existir."
    )

    particoes = contexto.get("particoes", []) or []
    if particoes:
        partes.append("")
        partes.append("## Partições relevantes")
        for p in particoes[:10]:
            partes.append(
                f"- {p.get('partition_name', 'N/C')} | atual={p.get('last_used_pct', 'N/C')} | "
                f"30d={p.get('forecast_30d', 'N/C')} | 60d={p.get('forecast_60d', 'N/C')} | "
                f"90d={p.get('forecast_90d', 'N/C')} | threshold={p.get('ultrapassa_threshold_em', 'N/C')}"
            )

    partes.append("")
    partes.append("## Recomendação")
    partes.append(
        "Manter a VM no FUP, validar crescimento com nova coleta e priorizar ação se o threshold já foi ultrapassado em 'Atual' "
        "ou se a previsão de 30/60 dias indicar saturação."
    )
    return "\n".join(partes)



def _carregar_config_llm_rmc() -> dict:
    """
    Carrega a configuração persistente da LLM do RMC Copilot.
    Prioridade:
    1. Variáveis de ambiente
    2. config/rmc_llm.env.json criado pelo instalador
    3. st.secrets["rmc_llm"], quando existir
    """
    cfg = {
        "provider": os.environ.get("RMC_LLM_PROVIDER", ""),
        "model": os.environ.get("RMC_LLM_MODEL", ""),
        "base_url": os.environ.get("RMC_LLM_BASE_URL", "") or os.environ.get("OLLAMA_HOST", ""),
        "api_key": os.environ.get("RMC_LLM_API_KEY", ""),
    }

    try:
        config_path = PROJECT_ROOT / "config" / "rmc_llm.env.json"
        if config_path.exists():
            file_cfg = json.loads(config_path.read_text(encoding="utf-8"))
            for k in ["provider", "model", "base_url", "api_key"]:
                if not cfg.get(k) and file_cfg.get(k):
                    cfg[k] = str(file_cfg.get(k))
    except Exception:
        pass

    try:
        if hasattr(st, "secrets") and "rmc_llm" in st.secrets:
            sec = st.secrets["rmc_llm"]
            for k in ["provider", "model", "base_url", "api_key"]:
                if not cfg.get(k) and sec.get(k):
                    cfg[k] = str(sec.get(k))
    except Exception:
        pass

    if not cfg.get("provider") and cfg.get("model"):
        cfg["provider"] = "ollama"
    if cfg.get("provider") == "ollama" and not cfg.get("base_url"):
        cfg["base_url"] = "http://localhost:11434"
    if not cfg.get("model"):
        cfg["model"] = "rmc-copilot-gemma3:1b"

    cfg["provider"] = str(cfg.get("provider", "")).strip().lower()
    return cfg


def _chamar_llm_rmc(prompt: str) -> str:
    """
    Adaptador para a LLM do RMC Copilot.
    Usa env vars, config/rmc_llm.env.json ou st.secrets.
    """
    cfg = _carregar_config_llm_rmc()
    provider = str(cfg.get("provider", "")).strip().lower()
    model = str(cfg.get("model") or "rmc-copilot-gemma3:1b").strip()

    if provider == "ollama":
        base_url = str(cfg.get("base_url") or "http://localhost:11434").strip()
        url = base_url.rstrip("/") + "/api/generate"
        payload = {
            "model": model,
            "prompt": prompt,
            "stream": False,
            "options": {"temperature": 0.2},
        }
        data = json.dumps(payload).encode("utf-8")
        req = urllib.request.Request(url, data=data, headers={"Content-Type": "application/json"}, method="POST")
        with urllib.request.urlopen(req, timeout=120) as resp:
            obj = json.loads(resp.read().decode("utf-8", errors="ignore"))
        return str(obj.get("response", "")).strip()

    if provider == "openai_compatible":
        base_url = str(cfg.get("base_url") or "").rstrip("/")
        api_key = str(cfg.get("api_key") or "")
        if not base_url:
            raise RuntimeError("RMC_LLM_BASE_URL/base_url não configurada.")
        url = base_url + "/chat/completions"
        payload = {
            "model": model,
            "messages": [
                {"role": "system", "content": "Você é a LLM do RMC Copilot, especializada em capacity planning VMware. Não invente números."},
                {"role": "user", "content": prompt},
            ],
            "temperature": 0.2,
        }
        headers = {"Content-Type": "application/json"}
        if api_key:
            headers["Authorization"] = "Bearer " + api_key
        req = urllib.request.Request(url, data=json.dumps(payload).encode("utf-8"), headers=headers, method="POST")
        with urllib.request.urlopen(req, timeout=120) as resp:
            obj = json.loads(resp.read().decode("utf-8", errors="ignore"))
        return str(obj["choices"][0]["message"]["content"]).strip()

    raise RuntimeError("LLM não configurada. Rode scripts/133_install_rmc_llm_starter.ps1 ou configure RMC_LLM_PROVIDER.")


def _gerar_relatorio_llm_ou_fallback(contexto: dict) -> tuple[str, str]:
    prompt = f"""
Use o papel de LLM do RMC Copilot para gerar um relatório demonstrativo de FUP.

Regras:
- Não invente forecast numérico.
- Quando o histórico for insuficiente, declare limitação e confiança baixa.
- Destaque o momento em que o threshold é ultrapassado: Atual, 30 dias, 60 dias ou 90 dias.
- Explique risco, evidências e recomendação.
- Linguagem executiva e técnica, em português do Brasil.

Contexto JSON:
{json.dumps(contexto, ensure_ascii=False, default=str, indent=2)}
"""
    try:
        txt = _chamar_llm_rmc(prompt)
        if txt.strip():
            return txt, "LLM"
    except Exception as exc:
        fallback = _relatorio_deterministico_fup(contexto)
        fallback += "\n\n---\n"
        fallback += f"**Observação:** LLM não executada nesta sessão. Motivo: {exc}\n"
        return fallback, "Fallback"

    return _relatorio_deterministico_fup(contexto), "Fallback"


def _render_relatorio_vm_fup(vm_nome: str, recurso: str, origem: str, limite: float, vm_df: pd.DataFrame, tabela_forecast: pd.DataFrame = None, df_particoes: pd.DataFrame = None):
    contexto = _montar_contexto_relatorio_fup(vm_nome, recurso, origem, limite, vm_df, tabela_forecast, df_particoes)
    with st.expander("Relatório demonstrativo assistido por LLM", expanded=False):
        st.caption(
            "Usa a nossa LLM se RMC_LLM_PROVIDER estiver configurado. "
            "Sem configuração, gera relatório determinístico sem inventar forecast."
        )
        if st.button("Gerar relatório da VM selecionada", key=f"btn_relatorio_llm_{_normalizar_nome_arquivo(vm_nome)}_{_normalizar_nome_arquivo(recurso)}"):
            relatorio, origem_relatorio = _gerar_relatorio_llm_ou_fallback(contexto)
            st.markdown(relatorio)
            st.download_button(
                "Baixar relatório Markdown",
                data=relatorio.encode("utf-8"),
                file_name=f"relatorio_fup_{_normalizar_nome_arquivo(vm_nome)}.md",
                mime="text/markdown",
                key=f"download_relatorio_llm_{_normalizar_nome_arquivo(vm_nome)}_{_normalizar_nome_arquivo(recurso)}",
            )
            st.caption(f"Fonte do relatório: {origem_relatorio}")


def _render_relatorio_fup_geral(fup_cpu: pd.DataFrame, fup_mem: pd.DataFrame, fup_part: pd.DataFrame, limite: float):
    contexto = {
        "tipo": "FUP geral",
        "threshold_percentual": float(limite),
        "run_id": str(st.session_state.get("resource_run_id_ativo") or st.session_state.get("execution_id_ativo") or ""),
        "resumo": [
            _resumo_threshold_textual(fup_cpu, "CPU", limite),
            _resumo_threshold_textual(fup_mem, "Memória", limite),
            _resumo_threshold_textual(fup_part, "Partições", limite),
        ],
        "top_cpu": fup_cpu.head(20).to_dict(orient="records") if fup_cpu is not None and not fup_cpu.empty else [],
        "top_mem": fup_mem.head(20).to_dict(orient="records") if fup_mem is not None and not fup_mem.empty else [],
        "top_particoes": fup_part.head(20).to_dict(orient="records") if fup_part is not None and not fup_part.empty else [],
    }

    with st.expander("Relatório demonstrativo geral do FUP", expanded=False):
        st.caption("Gera um relatório consolidado do FUP com threshold e momento de ultrapassagem.")
        if st.button("Gerar relatório geral do FUP", key="btn_relatorio_fup_geral_15f23"):
            relatorio, origem_relatorio = _gerar_relatorio_llm_ou_fallback(contexto)
            st.markdown(relatorio)
            st.download_button(
                "Baixar relatório geral Markdown",
                data=relatorio.encode("utf-8"),
                file_name="relatorio_fup_geral.md",
                mime="text/markdown",
                key="download_relatorio_fup_geral_15f23",
            )
            st.caption(f"Fonte do relatório: {origem_relatorio}")


# =============================================================================
# MARCO_15F_10_22_FUP_CLICK_VM
# Estado estável para voltar depois: FUP com clique na VM abrindo histórico + previsão.
# Não altera Capacity Dashboard.
# =============================================================================
def _selecionar_vm_por_grid(evento, df_origem: pd.DataFrame):
    idx = obter_linha_grid(evento)
    if idx is None or df_origem is None or df_origem.empty:
        return None
    try:
        if idx < 0 or idx >= len(df_origem):
            return None
        if "vm" not in df_origem.columns:
            return None
        vm = df_origem.iloc[int(idx)].get("vm")
        if vm is None or pd.isna(vm):
            return None
        return str(vm)
    except Exception:
        return None


def _render_fup_historico_vm(vm_nome: str, df_contexto: pd.DataFrame, recurso: str = "Todos", origem: str = "FUP", limite: float = 80.0):
    if not vm_nome:
        return

    vm_df = df_contexto[df_contexto["vm"].astype(str) == str(vm_nome)].copy() if "vm" in df_contexto.columns else pd.DataFrame()
    if vm_df.empty and "vm" in df_contexto.columns:
        termo = str(vm_nome).strip().upper()
        vm_df = df_contexto[df_contexto["vm"].astype(str).str.upper().str.contains(termo, na=False)].copy()

    if vm_df.empty:
        st.warning(f"Não encontrei a VM selecionada no contexto filtrado: {vm_nome}")
        return

    nome_vm = str(vm_df.iloc[0]["vm"])
    execution_id_grafico = st.session_state.get("resource_run_id_ativo") or st.session_state.get("execution_id_ativo")

    st.markdown(f'<div class="section-title">Histórico e previsão — VM: {nome_vm}</div>', unsafe_allow_html=True)
    st.caption(f"Origem: {origem} | Recurso: {recurso} | Threshold={limite:.0f}% | Clique em outra VM na tabela para trocar o gráfico.")

    render_resumo_recursos(
        vm_df.iloc[[0]],
        f"Recursos da VM — {nome_vm}",
        caminho_banco=CAMINHO_BANCO,
        tipo_escopo="VM",
        nome_escopo=nome_vm,
        run_id=st.session_state.get("resource_run_id_ativo"),
    )

    fig = criar_grafico_historico_forecast_escopo_duckdb(
        df_scope=vm_df.iloc[[0]],
        df_hist=None,
        tipo_escopo="VM",
        nome_escopo=nome_vm,
        recurso_selecionado=recurso,
        caminho_banco=CAMINHO_BANCO,
        run_id=execution_id_grafico,
        dias=int(st.session_state.get("historico_dias_ativo", 90) or 90),
    )

    if fig is not None:
        st.plotly_chart(fig, width="stretch")
        st.caption(f"Histórico real DuckDB + previsão 30/60/90 | run_id={execution_id_grafico or 'última coleta OK'}")
    else:
        st.warning("Não há histórico/forecast suficiente para montar o gráfico desta VM.")

    tabela = criar_tabela_forecast_escopo(vm_df.iloc[[0]], recurso)
    if tabela is not None and not tabela.empty:
        tabela = _adicionar_threshold_fup(tabela, recurso, limite)
        st.markdown("#### Forecast e momento de ultrapassagem do threshold")
        st.dataframe(formatar_dataframe_visual_2_casas(renomear_colunas_visual(tabela)), width="stretch", hide_index=True)
        _render_download_csv(
            "Baixar forecast da VM em CSV",
            tabela,
            f"fup_forecast_{_normalizar_nome_arquivo(nome_vm)}.csv",
            key=f"download_forecast_vm_{_normalizar_nome_arquivo(nome_vm)}_{_normalizar_nome_arquivo(recurso)}",
        )

    df_part_relatorio = pd.DataFrame()
    if recurso in ["Todos", "Disco"]:
        st.markdown(f'<div class="section-title">Partições de disco — VM: {nome_vm}</div>', unsafe_allow_html=True)
        fig_part, df_part, msg_part = criar_grafico_historico_particoes_duckdb(
            df_scope=vm_df.iloc[[0]],
            tipo_escopo="VM",
            nome_escopo=nome_vm,
            caminho_banco=CAMINHO_BANCO,
            run_id=execution_id_grafico,
            dias=int(st.session_state.get("historico_dias_ativo", 90) or 90),
            top_n=30,
            forecast_horizons=st.session_state.get("forecast_horizons_ativo", [30, 60, 90]),
        )
        if msg_part:
            st.caption(msg_part)
        if fig_part is not None:
            st.plotly_chart(fig_part, width="stretch")
        if df_part is not None and not df_part.empty:
            df_part = _enriquecer_particoes_fup(df_part, df_contexto)
            df_part = _adicionar_threshold_fup(df_part, "Disco", limite)
            df_part_relatorio = df_part.copy()
            st.dataframe(formatar_dataframe_visual_2_casas(renomear_colunas_visual(df_part)), width="stretch", hide_index=True)
            _render_download_csv(
                "Baixar partições da VM em CSV",
                df_part,
                f"fup_particoes_{_normalizar_nome_arquivo(nome_vm)}.csv",
                key=f"download_particoes_vm_{_normalizar_nome_arquivo(nome_vm)}",
            )

    _render_relatorio_vm_fup(
        vm_nome=nome_vm,
        recurso=recurso,
        origem=origem,
        limite=limite,
        vm_df=vm_df.iloc[[0]],
        tabela_forecast=tabela if tabela is not None and not tabela.empty else pd.DataFrame(),
        df_particoes=df_part_relatorio,
    )


def _render_grid_fup_vm(df_grid: pd.DataFrame, colunas: list, key: str, recurso: str, df_contexto: pd.DataFrame, origem: str, limite: float = 80.0):
    df_grid = df_grid.copy().reset_index(drop=True)
    evento = st.dataframe(
        formatar_dataframe_visual_2_casas(renomear_colunas_visual(df_grid[colunas])),
        width="stretch",
        hide_index=True,
        on_select="rerun",
        selection_mode="single-row",
        key=key,
    )

    _render_download_csv(
        f"Baixar {origem} em CSV",
        df_grid[colunas],
        f"{_normalizar_nome_arquivo(origem)}.csv",
        key=f"download_{key}",
    )

    vm_nome = _selecionar_vm_por_grid(evento, df_grid)
    if vm_nome:
        _render_fup_historico_vm(vm_nome, df_contexto=df_contexto, recurso=recurso, origem=origem, limite=limite)
    else:
        st.caption("Clique em uma VM na tabela acima para abrir o gráfico de histórico, previsão e relatório.")


def render_pagina_fup():
    """FUP consolidado: CPU, memória e partições acima do limite, com clique na VM para histórico + forecast."""
    df_analise_v4 = carregar_dados_dashboard()
    df_filtrado, _ = filtrar_dataframe(df_analise_v4)
    df_filtrado = preparar_coluna_host(df_filtrado)

    st.markdown('<div class="section-title">FUP — CPU, Memória e Partições acima de 80%</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="section-subtitle">Todas as VMs com CPU, memória ou partições acima do limite. Clique em uma VM para abrir o histórico real e a previsão.</div>',
        unsafe_allow_html=True,
    )

    st.info("MARCO 15F.10.22: FUP estabilizado como ponto de retorno. Capacity Dashboard não foi alterado neste pacote.")

    col_lim_1, col_lim_2, col_lim_3 = st.columns([1, 1, 2])
    with col_lim_1:
        limite_slider = st.slider("Controle deslizante (%)", min_value=70, max_value=99, value=80, step=1, key="limite_fup_slider_15f22")
    with col_lim_2:
        limite = st.number_input("Digite a porcentagem (%)", min_value=1.0, max_value=100.0, value=float(limite_slider), step=1.0, key="limite_fup_number_15f22")
    with col_lim_3:
        clusters_disponiveis = sorted(df_filtrado["cluster"].dropna().astype(str).unique().tolist()) if "cluster" in df_filtrado.columns else []
        clusters_sel = st.multiselect("Filtrar por cluster", clusters_disponiveis, default=clusters_disponiveis, key="fup_cluster_filter_15f22")

    if clusters_sel and "cluster" in df_filtrado.columns:
        df_filtrado = df_filtrado[df_filtrado["cluster"].astype(str).isin(clusters_sel)].copy()

    df = df_filtrado.copy()
    for col in ["cpu_p95_pct", "mem_p95_pct", "cpu_forecast_30d", "cpu_forecast_60d", "cpu_forecast_90d", "mem_forecast_30d", "mem_forecast_60d", "mem_forecast_90d"]:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce")

    mask_cpu = pd.Series(False, index=df.index)
    mask_mem = pd.Series(False, index=df.index)
    for col in [c for c in ["cpu_p95_pct", "cpu_forecast_30d", "cpu_forecast_60d", "cpu_forecast_90d"] if c in df.columns]:
        mask_cpu = mask_cpu | (df[col].fillna(0) >= float(limite))
    for col in [c for c in ["mem_p95_pct", "mem_forecast_30d", "mem_forecast_60d", "mem_forecast_90d"] if c in df.columns]:
        mask_mem = mask_mem | (df[col].fillna(0) >= float(limite))

    fup_cpu = _adicionar_threshold_fup(df[mask_cpu].copy(), "CPU", limite)
    fup_mem = _adicionar_threshold_fup(df[mask_mem].copy(), "Memória", limite)

    execution_id_grafico = st.session_state.get("resource_run_id_ativo") or st.session_state.get("execution_id_ativo")
    _, df_particoes, msg_particoes = criar_grafico_historico_particoes_duckdb(
        df_scope=df,
        tipo_escopo="Global",
        nome_escopo="Filtro atual",
        caminho_banco=CAMINHO_BANCO,
        run_id=execution_id_grafico,
        dias=int(st.session_state.get("historico_dias_ativo", 90) or 90),
        top_n=5000,
        forecast_horizons=st.session_state.get("forecast_horizons_ativo", [30, 60, 90]),
    )

    fup_part = pd.DataFrame()
    filtro_particao = ""
    excluir_particoes = []
    if df_particoes is not None and not df_particoes.empty:
        fup_part = _enriquecer_particoes_fup(df_particoes.copy(), df)
        for col in ["last_used_pct", "forecast_30d", "forecast_60d", "forecast_90d"]:
            if col in fup_part.columns:
                fup_part[col] = pd.to_numeric(fup_part[col], errors="coerce")
        cols_alerta = [c for c in ["last_used_pct", "forecast_30d", "forecast_60d", "forecast_90d"] if c in fup_part.columns]
        mascara_part = pd.Series(False, index=fup_part.index)
        for col in cols_alerta:
            mascara_part = mascara_part | (fup_part[col].fillna(0) >= float(limite))
        fup_part = fup_part[mascara_part].copy()

        colp1, colp2 = st.columns([1, 2])
        with colp1:
            filtro_particao = st.text_input("Filtrar partição contém", "", key="fup_particao_contem_15f22")
        with colp2:
            nomes_part = sorted(fup_part["partition_name"].dropna().astype(str).unique().tolist()) if "partition_name" in fup_part.columns else []
            excluir_particoes = st.multiselect("Excluir partições", nomes_part, default=[], key="fup_excluir_particoes_15f22")
        if filtro_particao.strip() and "partition_name" in fup_part.columns:
            fup_part = fup_part[fup_part["partition_name"].astype(str).str.contains(filtro_particao.strip(), case=False, na=False)].copy()
        if excluir_particoes and "partition_name" in fup_part.columns:
            fup_part = fup_part[~fup_part["partition_name"].astype(str).isin(excluir_particoes)].copy()

        fup_part = _adicionar_threshold_fup(fup_part, "Disco", limite)

    _render_relatorio_fup_geral(fup_cpu, fup_mem, fup_part, limite)

    c1, c2, c3, c4 = st.columns(4)
    with c1:
        render_kpi_card("CPU", _fmt_card_count(len(fup_cpu)), "kpi-orange", f">= {limite:.0f}%")
    with c2:
        render_kpi_card("Memória", _fmt_card_count(len(fup_mem)), "kpi-yellow", f">= {limite:.0f}%")
    with c3:
        render_kpi_card("Partições", _fmt_card_count(len(fup_part)), "kpi-red", f">= {limite:.0f}%")
    with c4:
        total_vms = set()
        for dfx in [fup_cpu, fup_mem, fup_part]:
            if dfx is not None and not dfx.empty and "vm" in dfx.columns:
                total_vms.update(dfx["vm"].dropna().astype(str).tolist())
        render_kpi_card("VMs afetadas", _fmt_card_count(len(total_vms)), "kpi-blue", "FUP")

    tab_cpu, tab_mem, tab_part = st.tabs(["CPU", "Memória", "Partições"])

    with tab_cpu:
        st.markdown("#### FUP CPU")
        if fup_cpu.empty:
            st.info("Nenhuma VM acima do limite para CPU.")
        else:
            cols = [c for c in ["cluster", "host", "vm", "cpu_p95_pct", "cpu_forecast_30d", "cpu_forecast_60d", "cpu_forecast_90d", "threshold_%", "ultrapassa_threshold_em", "maior_previsao_%", "confianca_forecast", "prioridade_final", "acao_final", "recomendacao_final"] if c in fup_cpu.columns]
            _render_grid_fup_vm(fup_cpu[cols], cols, key="grid_fup_cpu_click_vm_15f22", recurso="CPU", df_contexto=df, origem="FUP CPU", limite=limite)

    with tab_mem:
        st.markdown("#### FUP Memória")
        if fup_mem.empty:
            st.info("Nenhuma VM acima do limite para Memória.")
        else:
            cols = [c for c in ["cluster", "host", "vm", "mem_p95_pct", "mem_forecast_30d", "mem_forecast_60d", "mem_forecast_90d", "threshold_%", "ultrapassa_threshold_em", "maior_previsao_%", "confianca_forecast", "prioridade_final", "acao_final", "recomendacao_final"] if c in fup_mem.columns]
            _render_grid_fup_vm(fup_mem[cols], cols, key="grid_fup_mem_click_vm_15f22", recurso="Memória", df_contexto=df, origem="FUP Memória", limite=limite)

    with tab_part:
        st.markdown("#### FUP Partições")
        if msg_particoes:
            st.caption(msg_particoes)
        if fup_part.empty:
            st.info("Nenhuma partição acima do limite ou após os filtros aplicados.")
        else:
            cols = [c for c in ["cluster", "host", "datastore", "vm", "partition_name", "last_used_pct", "forecast_30d", "forecast_60d", "forecast_90d", "threshold_%", "ultrapassa_threshold_em", "maior_previsao_%", "confianca_forecast", "status_fup", "acao_fup", "evidencia_fup"] if c in fup_part.columns]
            _render_grid_fup_vm(fup_part[cols], cols, key="grid_fup_part_click_vm_15f22", recurso="Disco", df_contexto=df, origem="FUP Partições", limite=limite)


def render_pagina_otmz():
    """OTMZ separado: powered off, órfãos e snapshots."""
    df_analise_v4 = carregar_dados_dashboard()
    df_filtrado, _ = filtrar_dataframe(df_analise_v4)
    df_filtrado = preparar_coluna_host(df_filtrado)

    st.markdown('<div class="section-title">OTMZ — Otimização</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-subtitle">VMs powered off, discos órfãos e snapshots acima de 30 dias.</div>', unsafe_allow_html=True)

    df = df_filtrado.copy()
    power_cols = [c for c in ["power_state", "Power State", "status_power", "vm_power_state"] if c in df.columns]
    if power_cols:
        col = power_cols[0]
        poweredoff = df[df[col].astype(str).str.upper().str.contains("POWERED|OFF|DESLIG", na=False)].copy()
    else:
        poweredoff = pd.DataFrame()

    c1, c2, c3 = st.columns(3)
    with c1:
        render_kpi_card("Powered off", _fmt_card_count(len(poweredoff)), "kpi-yellow", "VMs")
    with c2:
        render_kpi_card("Discos órfãos", "N/C", "kpi-orange", "Coleta")
    with c3:
        render_kpi_card("Snapshots >30d", "N/C", "kpi-red", "Coleta")

    st.markdown("#### VMs powered off")
    if poweredoff.empty:
        st.info("A base atual não trouxe VMs powered off ou não possui coluna de power state.")
    else:
        cols = [c for c in ["cluster", "host", "vm", power_cols[0], "categoria_vm", "prioridade_final", "acao_final"] if c in poweredoff.columns]
        st.dataframe(formatar_dataframe_visual_2_casas(renomear_colunas_visual(poweredoff[cols])), width="stretch", hide_index=True)

    st.markdown("#### Discos órfãos")
    st.info("Seção separada criada. Falta incluir coleta específica de VMDK órfão no vROps/vCenter.")

    st.markdown("#### Snapshots acima de 30 dias")
    st.info("Seção separada criada. Falta incluir coleta específica de snapshots e idade.")


def render_pagina_pac():
    """PAC separado para forecast anual."""
    df_analise_v4 = carregar_dados_dashboard()
    df_filtrado, _ = filtrar_dataframe(df_analise_v4)
    df_filtrado = preparar_coluna_host(df_filtrado)

    st.markdown('<div class="section-title">PAC — Previsão anual de capacidade</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="section-subtitle">Previsão dos próximos 365 dias para hosts e clusters, baseada nos 365 dias anteriores.</div>',
        unsafe_allow_html=True,
    )
    st.info("Ambiente PAC separado. O próximo refinamento deve calcular forecast anual direto do histórico DuckDB por host/cluster.")

    resumo_cluster = criar_resumo_cluster(df_filtrado)
    if not resumo_cluster.empty:
        st.markdown("#### Clusters para PAC")
        st.dataframe(formatar_dataframe_visual_2_casas(renomear_colunas_visual(resumo_cluster)), width="stretch", hide_index=True)

    g1, g2 = st.columns(2)
    with g1:
        st.plotly_chart(criar_grafico_top_clusters(resumo_cluster), width="stretch")
    with g2:
        st.plotly_chart(criar_grafico_risco_futuro(df_filtrado), width="stretch")


def main():
    aplicar_estilo_bv()
    render_header()

    st.sidebar.markdown("## BV | RMC Copilot")
    st.sidebar.caption("Capacity Planning VMware")
    st.sidebar.markdown("---")
    st.sidebar.markdown("### Navegação")

    pagina = st.sidebar.radio(
        "Página",
        [
            "Capacity Dashboard",
            "Análise Individual de Recursos",
            "FUP",
            "OTMZ",
            "PAC",
            "Comparação entre execuções",
        ],
        label_visibility="collapsed",
    )

    st.sidebar.markdown("---")

    if pagina == "Capacity Dashboard":
        render_capacity_dashboard()
    elif pagina == "Análise Individual de Recursos":
        render_analise_individual_recursos()
    elif pagina == "FUP":
        render_pagina_fup()
    elif pagina == "OTMZ":
        render_pagina_otmz()
    elif pagina == "PAC":
        render_pagina_pac()
    elif pagina == "Comparação entre execuções":
        render_pagina_comparacao()
    else:
        render_capacity_dashboard()


if __name__ == "__main__":
    main()
