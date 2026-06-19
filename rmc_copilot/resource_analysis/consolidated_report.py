from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Any

from .models import ResourceAnalysisRequest, ResourceStats, sanitize_name

BV_BLUE = "0033A0"
BV_GREEN = "78BE20"
BV_LIGHT_GRAY = "F2F4F7"
BV_DARK = "1F2937"


@dataclass
class ResourceAnalysisItem:
    req: ResourceAnalysisRequest
    stats: ResourceStats
    narrative: Dict[str, str]
    chart_paths: Dict[str, str | Path]


def _fmt(value, unit: str = "") -> str:
    if value is None or value == "":
        return ""
    try:
        return f"{float(value):.2f}{(' ' + unit) if unit else ''}"
    except Exception:
        return str(value)


def _fmt_pct(value) -> str:
    if value is None or value == "":
        return ""
    try:
        return f"{float(value):.2f}%"
    except Exception:
        return str(value)


def _is_blank(value) -> bool:
    if value is None:
        return True
    text = str(value).strip()
    return text == "" or text.lower() in {"nan", "none", "null"}


def _fmt_optional_capacity(value, unit: str = "") -> str:
    if _is_blank(value):
        return "Não aplicável"
    try:
        return f"{float(value):.2f}{(' ' + unit) if unit else ''}"
    except Exception:
        return str(value)


def _human_action(value: str) -> str:
    text = str(value or "").strip().replace("_", " ")
    mapping = {
        "AVALIAR REDUÇÃO RECURSO": "AVALIAR REDUÇÃO",
        "AVALIAR REDUCAO RECURSO": "AVALIAR REDUÇÃO",
        "AUMENTAR RECURSO": "AUMENTAR RECURSO",
        "MANTER MONITORAMENTO": "MANTER MONITORAMENTO",
    }
    return mapping.get(text.upper(), text)


def _human_diagnosis(value: str) -> str:
    text = str(value or "").strip().replace("_", " ")
    return text


def _rel(path: str | Path, base: str | Path) -> str:
    try:
        return Path(path).resolve().relative_to(Path(base).resolve()).as_posix()
    except Exception:
        try:
            import os
            return os.path.relpath(str(path), str(base)).replace("\\", "/")
        except Exception:
            return Path(path).as_posix()


def _chart_items(chart_paths: Dict[str, str | Path]) -> List[tuple[str, Path]]:
    order = [
        ("A. Comparação e Previsão", "comparacao_previsao"),
        ("B. Histórico com Média Móvel", "media_movel"),
        ("C. Decomposição da Série Temporal", "decomposicao"),
        ("D. Distribuição de Uso", "histograma"),
        ("E. Uso Médio por Hora", "uso_por_hora"),
    ]
    items: List[tuple[str, Path]] = []
    for label, key in order:
        path = chart_paths.get(key)
        if path and Path(path).exists():
            items.append((label, Path(path)))
    return items


def _severity_rank(diagnosis: str) -> int:
    d = str(diagnosis or "").upper()
    if "CRIT" in d:
        return 0
    if "ATEN" in d or "RISCO" in d:
        return 1
    if "SUPER" in d or "SUB" in d:
        return 2
    if "OK" in d:
        return 3
    return 4


def _consolidated_name(items: List[ResourceAnalysisItem]) -> str:
    if not items:
        return "relatorio_consolidado"
    req0 = items[0].req
    return f"{req0.safe_solicitacao}_{req0.safe_vm}_analise_recursos_CONSOLIDADO"


def build_consolidated_markdown(items: List[ResourceAnalysisItem], out_dir: str | Path) -> Path:
    if not items:
        raise ValueError("Nenhuma análise recebida para consolidação.")
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    req0 = items[0].req
    generated_at = datetime.now().strftime("%d/%m/%Y %H:%M")
    title = f"Relatório Consolidado de Análise Individual de Recursos — {req0.vm}"
    md_path = out_dir / f"{_consolidated_name(items)}.md"

    worst = sorted(items, key=lambda x: (_severity_rank(x.stats.diagnosis), -float(x.stats.forecast_90_pct or 0)))[0]
    summary_rows = []
    for item in items:
        s = item.stats
        summary_rows.append(
            f"| {item.req.resource_title} | {_human_diagnosis(s.diagnosis)} | {_human_action(s.recommendation_action)} | {_fmt_pct(s.mean_pct)} | {_fmt_pct(s.p95_pct)} | {_fmt_pct(s.forecast_90_pct)} | {_fmt_optional_capacity(s.recommended_capacity, s.unit)} |"
        )

    header = f"""
<div style="border-top: 12px solid #{BV_BLUE}; border-bottom: 5px solid #{BV_GREEN}; padding: 18px 22px; background: #{BV_LIGHT_GRAY}; font-family: Arial, Helvetica, sans-serif;">
  <div style="font-size: 26px; font-weight: 700; color: #{BV_BLUE};">BV</div>
  <div style="font-size: 20px; font-weight: 700; color: #{BV_DARK}; margin-top: 4px;">{title}</div>
  <div style="font-size: 12px; color: #{BV_DARK}; margin-top: 10px;">Classificação: <strong>{req0.classificacao}</strong></div>
</div>
""".strip()

    content = [header, "", f"# {title}", ""]
    meta = [
        ("Solicitação", req0.solicitacao),
        ("Servidor / VM", req0.vm),
        ("Período histórico", f"{req0.periodo_dias} dias"),
        ("Solicitante", req0.solicitante or ""),
        ("Analista", req0.analista or ""),
        ("Classificação", req0.classificacao),
        ("Data de geração", generated_at),
    ]
    content.append("| Campo | Valor |")
    content.append("|:--|:--|")
    for k, v in meta:
        if v:
            content.append(f"| {k} | {v} |")
    indice = ["## Índice", "", "1. Resumo Executivo Consolidado", "2. Análises por Recurso"]
    for idx, item in enumerate(items, start=1):
        indice.append(f"   2.{idx}. {item.req.resource_title}")
    indice.append("3. Observações")
    content.extend([
        "",
        *indice,
        "",
        "---",
        "",
        "## 1. Resumo Executivo Consolidado",
        "",
        f"Foram avaliados {len(items)} recurso(s) da VM **{req0.vm}** para a solicitação **{req0.solicitacao}**. "
        f"O ponto de maior atenção identificado foi **{worst.req.resource_title}**, com diagnóstico **{worst.stats.diagnosis}** e ação recomendada **{worst.stats.recommendation_action}**.",
        "",
        "| Recurso | Diagnóstico | Ação | Média | P95 | Forecast 90d | Capacidade sugerida |",
        "|:--|:--|:--|--:|--:|--:|--:|",
        *summary_rows,
        "",
        "## 2. Análises por Recurso",
        "",
    ])

    for idx, item in enumerate(items, start=1):
        req, s, narrative, charts = item.req, item.stats, item.narrative, item.chart_paths
        content.extend([
            f"### 2.{idx}. {req.resource_title}",
            "",
            f"**Diagnóstico:** {_human_diagnosis(s.diagnosis)}  ",
            f"**Ação recomendada:** {_human_action(s.recommendation_action)}  ",
            f"**Uso médio:** {_fmt(s.mean, s.unit)} ({_fmt_pct(s.mean_pct)})  ",
            f"**P95:** {_fmt(s.p95, s.unit)} ({_fmt_pct(s.p95_pct)})  ",
            f"**Forecast 90 dias:** {_fmt(s.forecast_90, s.unit)} ({_fmt_pct(s.forecast_90_pct)})  ",
            "",
            "#### Resumo Executivo",
            "",
            narrative.get("resumo_executivo", ""),
            "",
            "#### Análise Técnica dos Gráficos",
            "",
            narrative.get("analise_graficos", ""),
            "",
        ])
        for label, path in _chart_items(charts):
            content.extend([f"##### {label}", "", f"![{label}]({_rel(path, out_dir)})", ""])
        content.extend([
            "#### Análise Estatística",
            "",
            narrative.get("analise_estatistica", ""),
            "",
            "| Métrica | Valor |",
            "|:--|--:|",
            f"| Capacidade total | {_fmt(s.capacity, s.unit)} |",
            f"| Margem de segurança ({s.threshold_pct:.0f}%) | {_fmt(s.threshold_value, s.unit)} |",
            f"| Uso mínimo | {_fmt(s.minimum, s.unit)} |",
            f"| Uso médio | {_fmt(s.mean, s.unit)} ({_fmt_pct(s.mean_pct)}) |",
            f"| Mediana | {_fmt(s.median, s.unit)} ({_fmt_pct(s.median_pct)}) |",
            f"| P95 | {_fmt(s.p95, s.unit)} ({_fmt_pct(s.p95_pct)}) |",
            f"| Máximo | {_fmt(s.maximum, s.unit)} ({_fmt_pct(s.maximum_pct)}) |",
            f"| Forecast 30 dias | {_fmt(s.forecast_30, s.unit)} ({_fmt_pct(s.forecast_30_pct)}) |",
            f"| Forecast 60 dias | {_fmt(s.forecast_60, s.unit)} ({_fmt_pct(s.forecast_60_pct)}) |",
            f"| Forecast 90 dias | {_fmt(s.forecast_90, s.unit)} ({_fmt_pct(s.forecast_90_pct)}) |",
            f"| Capacidade sugerida | {_fmt_optional_capacity(s.recommended_capacity, s.unit)} |",
            f"| Variação sugerida | {_fmt_optional_capacity(s.recommended_delta, s.unit)} |",
            "",
            "#### Conclusão e Recomendação",
            "",
            narrative.get("conclusao_recomendacao", ""),
            "",
            "---",
            "",
        ])

    content.extend([
        "## 3. Observações",
        "",
        "- A LLM/Data+RAG não calcula os números: ela apenas transforma os indicadores calculados pelo motor estatístico em texto executivo.",
        f"- A margem de segurança usada foi de {req0.threshold_pct:.0f}% da capacidade total.",
        "- O relatório consolidado reúne as análises individuais em um único documento para facilitar anexos em solicitação, FUP ou comunicação executiva.",
        "",
        "---",
        "",
        req0.classificacao,
    ])
    md_path.write_text("\n".join(content), encoding="utf-8")
    return md_path


def build_consolidated_txt(items: List[ResourceAnalysisItem], out_dir: str | Path) -> Path:
    if not items:
        raise ValueError("Nenhuma análise recebida para consolidação.")
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    req0 = items[0].req
    txt_path = out_dir / f"{_consolidated_name(items)}.txt"
    lines = [
        f"Relatório Consolidado de Análise Individual de Recursos — {req0.vm}",
        f"Solicitação: {req0.solicitacao}",
        f"VM: {req0.vm}",
        f"Período histórico: {req0.periodo_dias} dias",
        f"Solicitante: {req0.solicitante or ''}",
        f"Analista: {req0.analista or ''}",
        f"Classificação: {req0.classificacao}",
        "",
    ]
    for item in items:
        s = item.stats
        lines.extend([
            f"== {item.req.resource_title} ==",
            f"Diagnóstico: {_human_diagnosis(s.diagnosis)}",
            f"Ação: {_human_action(s.recommendation_action)}",
            f"Uso médio: {_fmt(s.mean, s.unit)} ({_fmt_pct(s.mean_pct)})",
            f"P95: {_fmt(s.p95, s.unit)} ({_fmt_pct(s.p95_pct)})",
            f"Forecast 90 dias: {_fmt(s.forecast_90, s.unit)} ({_fmt_pct(s.forecast_90_pct)})",
            "Resumo executivo:",
            item.narrative.get("resumo_executivo", ""),
            "Conclusão e recomendação:",
            item.narrative.get("conclusao_recomendacao", ""),
            "",
        ])
    txt_path.write_text("\n".join(lines), encoding="utf-8")
    return txt_path


def build_consolidated_docx(items: List[ResourceAnalysisItem], out_dir: str | Path) -> Path:
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError as exc:
        raise RuntimeError("Pacote python-docx não instalado. Rode: pip install python-docx") from exc

    def shade(cell, fill: str):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    def set_text(cell, text: str, bold: bool = False):
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(str(text or ""))
        r.bold = bold
        r.font.name = "Arial"
        r.font.size = Pt(9)

    if not items:
        raise ValueError("Nenhuma análise recebida para consolidação.")
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    req0 = items[0].req
    docx_path = out_dir / f"{_consolidated_name(items)}.docx"
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.55)
    sec.right_margin = Inches(0.55)

    header = doc.add_table(rows=2, cols=1)
    shade(header.rows[0].cells[0], BV_BLUE)
    p = header.rows[0].cells[0].paragraphs[0]
    run = p.add_run("BV | RMC Copilot")
    run.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.size = Pt(16)
    shade(header.rows[1].cells[0], BV_LIGHT_GRAY)
    p = header.rows[1].cells[0].paragraphs[0]
    r = p.add_run(f"Relatório Consolidado de Análise Individual de Recursos — {req0.vm}\nClassificação: {req0.classificacao}")
    r.bold = True
    r.font.color.rgb = RGBColor(31, 41, 55)

    title = doc.add_heading(f"Relatório Consolidado de Análise Individual de Recursos — {req0.vm}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = [
        ("Solicitação", req0.solicitacao),
        ("Servidor / VM", req0.vm),
        ("Período histórico", f"{req0.periodo_dias} dias"),
        ("Solicitante", req0.solicitante or ""),
        ("Analista", req0.analista or ""),
        ("Classificação", req0.classificacao),
        ("Data de geração", datetime.now().strftime("%d/%m/%Y %H:%M")),
    ]
    mt = doc.add_table(rows=1, cols=2)
    mt.style = "Table Grid"
    set_text(mt.rows[0].cells[0], "Campo", True)
    set_text(mt.rows[0].cells[1], "Valor", True)
    for k, v in meta:
        if v:
            row = mt.add_row().cells
            set_text(row[0], k, True)
            set_text(row[1], v)

    doc.add_heading("Índice", level=2)
    doc.add_paragraph("1. Resumo Executivo Consolidado")
    doc.add_paragraph("2. Análises por Recurso")
    for idx, item in enumerate(items, start=1):
        doc.add_paragraph(f"2.{idx}. {item.req.resource_title}")
    doc.add_paragraph("3. Observações")

    worst = sorted(items, key=lambda x: (_severity_rank(x.stats.diagnosis), -float(x.stats.forecast_90_pct or 0)))[0]
    doc.add_heading("1. Resumo Executivo Consolidado", level=2)
    doc.add_paragraph(
        f"Foram avaliados {len(items)} recurso(s) da VM {req0.vm} para a solicitação {req0.solicitacao}. "
        f"O ponto de maior atenção identificado foi {worst.req.resource_title}, com diagnóstico {_human_diagnosis(worst.stats.diagnosis)} "
        f"e ação recomendada {worst.stats.recommendation_action}."
    )
    st = doc.add_table(rows=1, cols=7)
    st.style = "Table Grid"
    headers = ["Recurso", "Diagnóstico", "Ação", "Média", "P95", "Forecast 90d", "Capacidade sugerida"]
    for cell, h in zip(st.rows[0].cells, headers):
        set_text(cell, h, True)
    for item in items:
        s = item.stats
        row = st.add_row().cells
        values = [item.req.resource_title, _human_diagnosis(s.diagnosis), _human_action(s.recommendation_action), _fmt_pct(s.mean_pct), _fmt_pct(s.p95_pct), _fmt_pct(s.forecast_90_pct), _fmt_optional_capacity(s.recommended_capacity, s.unit)]
        for cell, v in zip(row, values):
            set_text(cell, v)

    doc.add_heading("2. Análises por Recurso", level=2)
    for idx, item in enumerate(items, start=1):
        req, s, narrative, charts = item.req, item.stats, item.narrative, item.chart_paths
        doc.add_heading(f"2.{idx}. {req.resource_title}", level=3)
        doc.add_paragraph(f"Diagnóstico: {_human_diagnosis(s.diagnosis)} | Ação: {_human_action(s.recommendation_action)}")
        doc.add_heading("Resumo Executivo", level=4)
        doc.add_paragraph(narrative.get("resumo_executivo", ""))
        doc.add_heading("Análise Técnica dos Gráficos", level=4)
        doc.add_paragraph(narrative.get("analise_graficos", ""))
        for label, path in _chart_items(charts):
            doc.add_paragraph(label).runs[0].bold = True
            try:
                doc.add_picture(str(path), width=Inches(6.3))
            except Exception as exc:
                doc.add_paragraph(f"[Falha ao inserir gráfico: {path} — {exc}]")
        doc.add_heading("Análise Estatística", level=4)
        doc.add_paragraph(narrative.get("analise_estatistica", ""))
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        set_text(tbl.rows[0].cells[0], "Métrica", True)
        set_text(tbl.rows[0].cells[1], "Valor", True)
        rows = [
            ("Capacidade total", _fmt(s.capacity, s.unit)),
            (f"Margem de segurança ({s.threshold_pct:.0f}%)", _fmt(s.threshold_value, s.unit)),
            ("Uso médio", f"{_fmt(s.mean, s.unit)} ({_fmt_pct(s.mean_pct)})"),
            ("P95", f"{_fmt(s.p95, s.unit)} ({_fmt_pct(s.p95_pct)})"),
            ("Máximo", f"{_fmt(s.maximum, s.unit)} ({_fmt_pct(s.maximum_pct)})"),
            ("Forecast 90 dias", f"{_fmt(s.forecast_90, s.unit)} ({_fmt_pct(s.forecast_90_pct)})"),
            ("Capacidade sugerida", _fmt_optional_capacity(s.recommended_capacity, s.unit)),
            ("Variação sugerida", _fmt_optional_capacity(s.recommended_delta, s.unit)),
        ]
        for k, v in rows:
            row = tbl.add_row().cells
            set_text(row[0], k)
            set_text(row[1], v)
        doc.add_heading("Conclusão e Recomendação", level=4)
        doc.add_paragraph(narrative.get("conclusao_recomendacao", ""))

    doc.add_heading("3. Observações", level=2)
    for obs in [
        "A LLM/Data+RAG não calcula os números: ela apenas transforma os indicadores calculados pelo motor estatístico em texto executivo.",
        f"A margem de segurança usada foi de {req0.threshold_pct:.0f}% da capacidade total.",
        "O relatório consolidado reúne as análises individuais em um único documento.",
    ]:
        doc.add_paragraph(obs, style="List Bullet")
    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = req0.classificacao
    doc.save(docx_path)
    return docx_path


def build_consolidated_pdf(items: List[ResourceAnalysisItem], out_dir: str | Path) -> Path:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
    except ImportError as exc:
        raise RuntimeError("Pacote reportlab não instalado. Rode: pip install reportlab") from exc

    if not items:
        raise ValueError("Nenhuma análise recebida para consolidação.")
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    req0 = items[0].req
    pdf_path = out_dir / f"{_consolidated_name(items)}.pdf"
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="BVTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=17, textColor=colors.HexColor(f"#{BV_BLUE}"), spaceAfter=8))
    styles.add(ParagraphStyle(name="BVHeading", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=12, textColor=colors.HexColor(f"#{BV_DARK}"), spaceBefore=8, spaceAfter=5))
    normal = styles["Normal"]
    normal.fontName = "Helvetica"
    normal.fontSize = 8.5
    normal.leading = 11
    normal.splitLongWords = 0
    styles.add(ParagraphStyle(name="BVTableSmall", parent=styles["Normal"], fontName="Helvetica", fontSize=7.1, leading=8.2, splitLongWords=0))
    table_small = styles["BVTableSmall"]

    def clean(text: Any) -> str:
        text = str(text or "")
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")

    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.drawString(1.5 * cm, 1.0 * cm, req0.classificacao)
        canvas.drawRightString(A4[0] - 1.5 * cm, 1.0 * cm, f"Página {doc.page}")
        canvas.restoreState()

    doc = SimpleDocTemplate(str(pdf_path), pagesize=A4, rightMargin=1.35 * cm, leftMargin=1.35 * cm, topMargin=1.2 * cm, bottomMargin=1.4 * cm)
    story = []
    header = Table([
        [Paragraph('<font color="white"><b>BV | RMC Copilot</b></font>', styles["Heading1"])],
        [Paragraph(f"<b>Relatório Consolidado de Análise Individual de Recursos — {clean(req0.vm)}</b><br/>Classificação: <b>{clean(req0.classificacao)}</b>", normal)],
    ], colWidths=[18.3 * cm])
    header.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{BV_BLUE}")),
        ("BACKGROUND", (0, 1), (-1, 1), colors.HexColor(f"#{BV_LIGHT_GRAY}")),
        ("LINEBELOW", (0, 1), (-1, 1), 3, colors.HexColor(f"#{BV_GREEN}")),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
    ]))
    story.extend([header, Spacer(1, 8), Paragraph(f"Relatório Consolidado de Análise Individual de Recursos — {clean(req0.vm)}", styles["BVTitle"])])

    meta = [
        ["Campo", "Valor"],
        ["Solicitação", req0.solicitacao],
        ["Servidor / VM", req0.vm],
        ["Período histórico", f"{req0.periodo_dias} dias"],
        ["Solicitante", req0.solicitante or ""],
        ["Analista", req0.analista or ""],
        ["Classificação", req0.classificacao],
        ["Data de geração", datetime.now().strftime("%d/%m/%Y %H:%M")],
    ]
    meta_table = Table([[Paragraph(clean(c), normal) for c in row] for row in meta], colWidths=[5 * cm, 13.3 * cm])
    meta_table.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{BV_LIGHT_GRAY}")), ("GRID", (0, 0), (-1, -1), 0.25, colors.grey)]))
    story.extend([meta_table, Spacer(1, 8)])
    story.append(Paragraph("Índice", styles["BVHeading"]))
    indice_text = "1. Resumo Executivo Consolidado<br/>2. Análises por Recurso"
    for idx, item in enumerate(items, start=1):
        indice_text += f"<br/>&nbsp;&nbsp;2.{idx}. {clean(item.req.resource_title)}"
    indice_text += "<br/>3. Observações"
    story.extend([Paragraph(indice_text, normal), Spacer(1, 8)])

    worst = sorted(items, key=lambda x: (_severity_rank(x.stats.diagnosis), -float(x.stats.forecast_90_pct or 0)))[0]
    story.append(Paragraph("1. Resumo Executivo Consolidado", styles["BVHeading"]))
    story.append(Paragraph(clean(f"Foram avaliados {len(items)} recurso(s) da VM {req0.vm} para a solicitação {req0.solicitacao}. O ponto de maior atenção identificado foi {worst.req.resource_title}, com diagnóstico {_human_diagnosis(worst.stats.diagnosis)} e ação recomendada {worst.stats.recommendation_action}."), normal))
    summary = [["Recurso", "Diagnóstico", "Ação", "Média", "P95", "Forecast 90d"]]
    for item in items:
        s = item.stats
        summary.append([item.req.resource_title, _human_diagnosis(s.diagnosis), _human_action(s.recommendation_action), _fmt_pct(s.mean_pct), _fmt_pct(s.p95_pct), _fmt_pct(s.forecast_90_pct)])
    summary_table = Table([[Paragraph(clean(c), table_small) for c in row] for row in summary], colWidths=[4.3*cm, 3.2*cm, 4.5*cm, 1.6*cm, 1.6*cm, 2.1*cm])
    summary_table.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{BV_LIGHT_GRAY}")), ("GRID", (0, 0), (-1, -1), 0.25, colors.grey), ("VALIGN", (0, 0), (-1, -1), "TOP")]))
    story.extend([Spacer(1, 6), summary_table, Spacer(1, 8), Paragraph("2. Análises por Recurso", styles["BVHeading"])])

    for idx, item in enumerate(items, start=1):
        req, s, narrative, charts = item.req, item.stats, item.narrative, item.chart_paths
        story.append(Paragraph(f"2.{idx}. {clean(req.resource_title)}", styles["BVHeading"]))
        story.append(Paragraph(clean(f"Diagnóstico: {_human_diagnosis(s.diagnosis)} | Ação: {_human_action(s.recommendation_action)} | Média: {_fmt_pct(s.mean_pct)} | P95: {_fmt_pct(s.p95_pct)} | Forecast 90d: {_fmt_pct(s.forecast_90_pct)}"), normal))
        for heading, text in [("Resumo Executivo", narrative.get("resumo_executivo", "")), ("Análise Técnica dos Gráficos", narrative.get("analise_graficos", ""))]:
            story.append(Paragraph(heading, styles["BVHeading"]))
            story.append(Paragraph(clean(text), normal))
        for n, (label, path) in enumerate(_chart_items(charts), start=1):
            story.append(Paragraph(clean(label), styles["BVHeading"]))
            try:
                story.append(Image(str(path), width=17.2 * cm, height=9.1 * cm, kind="proportional"))
            except Exception as exc:
                story.append(Paragraph(clean(f"Falha ao inserir gráfico: {path} — {exc}"), normal))
            story.append(Spacer(1, 5))
            if n in {2, 4}:
                story.append(PageBreak())
        story.append(Paragraph("Análise Estatística", styles["BVHeading"]))
        story.append(Paragraph(clean(narrative.get("analise_estatistica", "")), normal))
        rows = [["Métrica", "Valor"], ["Capacidade total", _fmt(s.capacity, s.unit)], ["Uso médio", f"{_fmt(s.mean, s.unit)} ({_fmt_pct(s.mean_pct)})"], ["P95", f"{_fmt(s.p95, s.unit)} ({_fmt_pct(s.p95_pct)})"], ["Máximo", f"{_fmt(s.maximum, s.unit)} ({_fmt_pct(s.maximum_pct)})"], ["Forecast 90 dias", f"{_fmt(s.forecast_90, s.unit)} ({_fmt_pct(s.forecast_90_pct)})"], ["Capacidade sugerida", _fmt_optional_capacity(s.recommended_capacity, s.unit)]]
        table = Table([[Paragraph(clean(c), normal) for c in row] for row in rows], colWidths=[7*cm, 10*cm])
        table.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{BV_LIGHT_GRAY}")), ("GRID", (0, 0), (-1, -1), 0.25, colors.grey)]))
        story.extend([table, Paragraph("Conclusão e Recomendação", styles["BVHeading"]), Paragraph(clean(narrative.get("conclusao_recomendacao", "")), normal)])
        if idx < len(items):
            story.append(PageBreak())

    story.append(Paragraph("3. Observações", styles["BVHeading"]))
    story.append(Paragraph(clean("A LLM/Data+RAG não calcula os números: ela apenas transforma os indicadores calculados pelo motor estatístico em texto executivo."), normal))
    doc.build(story, onFirstPage=footer, onLaterPages=footer)
    return pdf_path


def export_consolidated_reports(items: List[ResourceAnalysisItem], out_dir: str | Path, formats: Iterable[str] | str = ("md", "docx", "pdf")) -> Dict[str, Path]:
    from .exporters import parse_formats
    selected = parse_formats(formats)
    out: Dict[str, Path] = {}
    if "md" in selected:
        out["md"] = build_consolidated_markdown(items, out_dir)
    out["txt"] = build_consolidated_txt(items, out_dir)
    if "docx" in selected:
        out["docx"] = build_consolidated_docx(items, out_dir)
    if "pdf" in selected:
        out["pdf"] = build_consolidated_pdf(items, out_dir)
    return out
