from __future__ import annotations

from pathlib import Path
from typing import Dict, Iterable, List, Tuple
from zipfile import ZipFile, ZIP_DEFLATED
import re

from .models import ResourceAnalysisRequest, ResourceStats
from .ptbr_format import period_br

BV_BLUE = "0033A0"
BV_GREEN = "78BE20"
BV_LIGHT_GRAY = "F2F4F7"
BV_DARK = "1F2937"


def _fmt(value, unit: str = "") -> str:
    if value is None or value == "":
        return ""
    try:
        return f"{float(value):.2f}{(' ' + unit) if unit else ''}"
    except Exception:
        return str(value)


def _fmt_pct(value) -> str:
    if value is None or value == "":
        return ""
    try:
        return f"{float(value):.2f}%"
    except Exception:
        return str(value)


def _is_blank(value) -> bool:
    if value is None:
        return True
    text = str(value).strip()
    return text == "" or text.lower() in {"nan", "none", "null"}


def _fmt_optional_capacity(value, unit: str = "") -> str:
    if _is_blank(value):
        return "Não aplicável"
    try:
        return f"{float(value):.2f}{(' ' + unit) if unit else ''}"
    except Exception:
        return str(value)


def _human_action(value: str) -> str:
    text = str(value or "").strip().replace("_", " ")
    mapping = {
        "AVALIAR REDUÇÃO RECURSO": "AVALIAR REDUÇÃO",
        "AVALIAR REDUCAO RECURSO": "AVALIAR REDUÇÃO",
    }
    return mapping.get(text.upper(), text)


def _human_diagnosis(value: str) -> str:
    return str(value or "").strip().replace("_", " ")


def _chart_items(chart_paths: Dict[str, str | Path]) -> List[Tuple[str, Path]]:
    order = [
        ("A. Comparação e Previsão", "comparacao_previsao"),
        ("B. Histórico com Média Móvel", "media_movel"),
        ("C. Decomposição da Série Temporal", "decomposicao"),
        ("D. Distribuição de Uso", "histograma"),
        ("E. Uso Médio por Hora", "uso_por_hora"),
    ]
    items: List[Tuple[str, Path]] = []
    for label, key in order:
        path = chart_paths.get(key)
        if path and Path(path).exists():
            items.append((label, Path(path)))
    return items


def build_plain_text_report(req: ResourceAnalysisRequest, stats: ResourceStats, narrative: Dict[str, str], out_dir: str | Path) -> Path:
    """Gera versão .txt simples, útil para copiar/colar ou baixar todos os textos."""
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    txt_path = out_dir / f"{req.safe_solicitacao}_{req.safe_vm}_{req.safe_resource}_analise_recursos.txt"
    lines = [
        f"Relatório de Análise Individual de Recursos — {req.vm}",
        f"Solicitação: {req.solicitacao}",
        f"Servidor / VM: {req.vm}",
        f"Recurso: {req.resource_title}",
        f"Período histórico: {req.periodo_dias} dias",
        f"Período analisado: {period_br(stats.start, stats.end)}",
        f"Solicitante: {req.solicitante or ''}",
        f"Analista: {req.analista or ''}",
        f"Classificação: {req.classificacao}",
        "",
        "1. Resumo Executivo",
        narrative.get("resumo_executivo", ""),
        "",
        "2. Análise Técnica dos Gráficos",
        narrative.get("analise_graficos", ""),
        "",
        "3. Análise Estatística",
        narrative.get("analise_estatistica", ""),
        "",
        "Indicadores calculados:",
        f"- Capacidade total: {_fmt(stats.capacity, stats.unit)}",
        f"- Margem de segurança ({stats.threshold_pct:.0f}%): {_fmt(stats.threshold_value, stats.unit)}",
        f"- Uso mínimo: {_fmt(stats.minimum, stats.unit)}",
        f"- Uso médio: {_fmt(stats.mean, stats.unit)} ({_fmt_pct(stats.mean_pct)})",
        f"- Mediana: {_fmt(stats.median, stats.unit)} ({_fmt_pct(stats.median_pct)})",
        f"- Q1: {_fmt(stats.q1, stats.unit)}",
        f"- Q3: {_fmt(stats.q3, stats.unit)}",
        f"- P95: {_fmt(stats.p95, stats.unit)} ({_fmt_pct(stats.p95_pct)})",
        f"- Máximo: {_fmt(stats.maximum, stats.unit)} ({_fmt_pct(stats.maximum_pct)})",
        f"- Forecast 30 dias: {_fmt(stats.forecast_30, stats.unit)} ({_fmt_pct(stats.forecast_30_pct)})",
        f"- Forecast 60 dias: {_fmt(stats.forecast_60, stats.unit)} ({_fmt_pct(stats.forecast_60_pct)})",
        f"- Forecast 90 dias: {_fmt(stats.forecast_90, stats.unit)} ({_fmt_pct(stats.forecast_90_pct)})",
        f"- Diagnóstico: {_human_diagnosis(stats.diagnosis)}",
        f"- Ação recomendada: {_human_action(stats.recommendation_action)}",
        f"- Capacidade sugerida: {_fmt_optional_capacity(stats.recommended_capacity, stats.unit)}",
        f"- Variação sugerida: {_fmt_optional_capacity(stats.recommended_delta, stats.unit)}",
        "",
        "4. Conclusão e Recomendação",
        narrative.get("conclusao_recomendacao", ""),
        "",
        "Observações:",
        "- A LLM/Data+RAG não calcula os números; ela apenas transforma os indicadores calculados pelo motor estatístico em texto executivo.",
        f"- A margem de segurança usada foi de {stats.threshold_pct:.0f}% da capacidade total.",
        f"- {stats.confidence_note}",
    ]
    txt_path.write_text("\n".join(lines), encoding="utf-8")
    return txt_path


def build_docx_report(req: ResourceAnalysisRequest, stats: ResourceStats, narrative: Dict[str, str], chart_paths: Dict[str, str | Path], out_dir: str | Path) -> Path:
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError as exc:
        raise RuntimeError("Pacote python-docx não instalado. Rode: pip install python-docx") from exc

    def shade_cell(cell, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.bold = bold
        run.font.name = "Arial"
        if color:
            run.font.color.rgb = RGBColor.from_string(color)

    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    docx_path = out_dir / f"{req.safe_solicitacao}_{req.safe_vm}_{req.safe_resource}_analise_recursos.docx"

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    # Cabeçalho BV
    header_table = doc.add_table(rows=2, cols=1)
    header_table.autofit = True
    shade_cell(header_table.cell(0, 0), BV_BLUE)
    set_cell_text(header_table.cell(0, 0), "BV | RMC Copilot", bold=True, color="FFFFFF")
    header_table.cell(0, 0).paragraphs[0].runs[0].font.size = Pt(18)
    shade_cell(header_table.cell(1, 0), BV_LIGHT_GRAY)
    set_cell_text(header_table.cell(1, 0), f"Relatório de Análise Individual de Recursos — {req.vm}\nClassificação: {req.classificacao}", bold=True, color=BV_DARK)

    doc.add_paragraph()
    title = doc.add_heading(f"Relatório de Análise Individual de Recursos — {req.vm}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = [
        ("Solicitação", req.solicitacao),
        ("Servidor / VM", req.vm),
        ("Recurso", req.resource_title),
        ("Período histórico", f"{req.periodo_dias} dias"),
        ("Período analisado", period_br(stats.start, stats.end)),
        ("Solicitante", req.solicitante or ""),
        ("Analista", req.analista or ""),
        ("Origem dos dados", req.origem),
    ]
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    set_cell_text(t.rows[0].cells[0], "Campo", bold=True)
    set_cell_text(t.rows[0].cells[1], "Valor", bold=True)
    for k, v in meta:
        if v:
            row = t.add_row().cells
            set_cell_text(row[0], k, bold=True)
            set_cell_text(row[1], v)

    sections = [
        ("1. Resumo Executivo", narrative.get("resumo_executivo", "")),
        ("2. Análise Técnica dos Gráficos", narrative.get("analise_graficos", "")),
    ]
    for heading, text in sections:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(text)

    for label, path in _chart_items(chart_paths):
        doc.add_heading(label, level=3)
        try:
            doc.add_picture(str(path), width=Inches(6.3))
        except Exception as exc:
            doc.add_paragraph(f"[Falha ao inserir gráfico: {path} — {exc}]")

    doc.add_heading("3. Análise Estatística", level=2)
    doc.add_paragraph(narrative.get("analise_estatistica", ""))
    stats_rows = [
        ("Capacidade total", _fmt(stats.capacity, stats.unit)),
        (f"Margem de segurança ({stats.threshold_pct:.0f}%)", _fmt(stats.threshold_value, stats.unit)),
        ("Uso mínimo", _fmt(stats.minimum, stats.unit)),
        ("Uso médio", f"{_fmt(stats.mean, stats.unit)} ({_fmt_pct(stats.mean_pct)})"),
        ("Mediana", f"{_fmt(stats.median, stats.unit)} ({_fmt_pct(stats.median_pct)})"),
        ("Q1", _fmt(stats.q1, stats.unit)),
        ("Q3", _fmt(stats.q3, stats.unit)),
        ("P95", f"{_fmt(stats.p95, stats.unit)} ({_fmt_pct(stats.p95_pct)})"),
        ("Uso máximo", f"{_fmt(stats.maximum, stats.unit)} ({_fmt_pct(stats.maximum_pct)})"),
        ("Forecast 30 dias", f"{_fmt(stats.forecast_30, stats.unit)} ({_fmt_pct(stats.forecast_30_pct)})"),
        ("Forecast 60 dias", f"{_fmt(stats.forecast_60, stats.unit)} ({_fmt_pct(stats.forecast_60_pct)})"),
        ("Forecast 90 dias", f"{_fmt(stats.forecast_90, stats.unit)} ({_fmt_pct(stats.forecast_90_pct)})"),
        ("Diagnóstico", _human_diagnosis(stats.diagnosis)),
        ("Ação recomendada", _human_action(stats.recommendation_action)),
        ("Capacidade sugerida", _fmt_optional_capacity(stats.recommended_capacity, stats.unit)),
        ("Variação sugerida", _fmt_optional_capacity(stats.recommended_delta, stats.unit)),
    ]
    stbl = doc.add_table(rows=1, cols=2)
    stbl.style = "Table Grid"
    set_cell_text(stbl.rows[0].cells[0], "Métrica", bold=True)
    set_cell_text(stbl.rows[0].cells[1], "Valor", bold=True)
    for k, v in stats_rows:
        row = stbl.add_row().cells
        set_cell_text(row[0], k)
        set_cell_text(row[1], v)

    doc.add_heading("4. Conclusão e Recomendação", level=2)
    doc.add_paragraph(narrative.get("conclusao_recomendacao", ""))
    doc.add_heading("5. Observações", level=2)
    for obs in [
        "A LLM/Data+RAG não calcula os números: ela apenas transforma os indicadores calculados pelo motor estatístico em texto executivo.",
        f"A margem de segurança usada foi de {stats.threshold_pct:.0f}% da capacidade total.",
        stats.confidence_note,
    ]:
        doc.add_paragraph(obs, style="List Bullet")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = req.classificacao
    footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.save(docx_path)
    return docx_path


def build_pdf_report(req: ResourceAnalysisRequest, stats: ResourceStats, narrative: Dict[str, str], chart_paths: Dict[str, str | Path], out_dir: str | Path) -> Path:
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
    except ImportError as exc:
        raise RuntimeError("Pacote reportlab não instalado. Rode: pip install reportlab") from exc

    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = out_dir / f"{req.safe_solicitacao}_{req.safe_vm}_{req.safe_resource}_analise_recursos.pdf"
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="BVTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=18, textColor=colors.HexColor(f"#{BV_BLUE}"), alignment=TA_LEFT, spaceAfter=8))
    styles.add(ParagraphStyle(name="BVHeading", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=13, textColor=colors.HexColor(f"#{BV_DARK}"), spaceBefore=10, spaceAfter=6))
    styles.add(ParagraphStyle(name="BVSmall", parent=styles["Normal"], fontSize=8, textColor=colors.HexColor(f"#{BV_DARK}")))
    normal = styles["Normal"]
    normal.fontName = "Helvetica"
    normal.fontSize = 9
    normal.leading = 12

    def clean(text: str) -> str:
        text = str(text or "")
        text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        return text.replace("\n", "<br/>")

    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.drawString(1.5 * cm, 1.0 * cm, req.classificacao)
        canvas.drawRightString(A4[0] - 1.5 * cm, 1.0 * cm, f"Página {doc.page}")
        canvas.restoreState()

    doc = SimpleDocTemplate(str(pdf_path), pagesize=A4, rightMargin=1.5 * cm, leftMargin=1.5 * cm, topMargin=1.3 * cm, bottomMargin=1.5 * cm)
    story = []

    header = Table([
        [Paragraph('<font color="white"><b>BV | RMC Copilot</b></font>', styles["Heading1"])],
        [Paragraph(f"<b>Relatório de Análise Individual de Recursos — {clean(req.vm)}</b><br/>Classificação: <b>{clean(req.classificacao)}</b>", styles["BVSmall"])],
    ], colWidths=[18 * cm])
    header.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{BV_BLUE}")),
        ("BACKGROUND", (0, 1), (-1, 1), colors.HexColor(f"#{BV_LIGHT_GRAY}")),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
        ("TOPPADDING", (0, 0), (-1, 0), 8),
        ("BOTTOMPADDING", (0, 1), (-1, 1), 8),
        ("TOPPADDING", (0, 1), (-1, 1), 8),
        ("LINEBELOW", (0, 1), (-1, 1), 3, colors.HexColor(f"#{BV_GREEN}")),
    ]))
    story.extend([header, Spacer(1, 10), Paragraph(f"Relatório de Análise Individual de Recursos — {clean(req.vm)}", styles["BVTitle"])])

    meta = [
        ["Campo", "Valor"],
        ["Solicitação", req.solicitacao],
        ["Servidor / VM", req.vm],
        ["Recurso", req.resource_title],
        ["Período histórico", f"{req.periodo_dias} dias"],
        ["Período analisado", period_br(stats.start, stats.end)],
        ["Solicitante", req.solicitante or ""],
        ["Analista", req.analista or ""],
        ["Origem dos dados", req.origem],
    ]
    meta_table = Table([[Paragraph(clean(c), normal) for c in row] for row in meta], colWidths=[5 * cm, 13 * cm])
    meta_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{BV_LIGHT_GRAY}")),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
    ]))
    story.extend([meta_table, Spacer(1, 8)])

    for heading, text in [
        ("1. Resumo Executivo", narrative.get("resumo_executivo", "")),
        ("2. Análise Técnica dos Gráficos", narrative.get("analise_graficos", "")),
    ]:
        story.append(Paragraph(heading, styles["BVHeading"]))
        story.append(Paragraph(clean(text), normal))
        story.append(Spacer(1, 6))

    for idx, (label, path) in enumerate(_chart_items(chart_paths), start=1):
        story.append(Paragraph(label, styles["BVHeading"]))
        try:
            story.append(Image(str(path), width=17.2 * cm, height=9.3 * cm, kind="proportional"))
        except Exception as exc:
            story.append(Paragraph(clean(f"Falha ao inserir gráfico: {path} — {exc}"), normal))
        story.append(Spacer(1, 6))
        if idx in {2, 4}:
            story.append(PageBreak())

    story.append(Paragraph("3. Análise Estatística", styles["BVHeading"]))
    story.append(Paragraph(clean(narrative.get("analise_estatistica", "")), normal))
    stats_rows = [
        ["Métrica", "Valor"],
        ["Capacidade total", _fmt(stats.capacity, stats.unit)],
        [f"Margem de segurança ({stats.threshold_pct:.0f}%)", _fmt(stats.threshold_value, stats.unit)],
        ["Uso médio", f"{_fmt(stats.mean, stats.unit)} ({_fmt_pct(stats.mean_pct)})"],
        ["P95", f"{_fmt(stats.p95, stats.unit)} ({_fmt_pct(stats.p95_pct)})"],
        ["Uso máximo", f"{_fmt(stats.maximum, stats.unit)} ({_fmt_pct(stats.maximum_pct)})"],
        ["Forecast 30 dias", f"{_fmt(stats.forecast_30, stats.unit)} ({_fmt_pct(stats.forecast_30_pct)})"],
        ["Forecast 60 dias", f"{_fmt(stats.forecast_60, stats.unit)} ({_fmt_pct(stats.forecast_60_pct)})"],
        ["Forecast 90 dias", f"{_fmt(stats.forecast_90, stats.unit)} ({_fmt_pct(stats.forecast_90_pct)})"],
        ["Diagnóstico", _human_diagnosis(stats.diagnosis)],
        ["Ação recomendada", _human_action(stats.recommendation_action)],
        ["Capacidade sugerida", _fmt_optional_capacity(stats.recommended_capacity, stats.unit)],
        ["Variação sugerida", _fmt_optional_capacity(stats.recommended_delta, stats.unit)],
    ]
    table = Table([[Paragraph(clean(c), normal) for c in row] for row in stats_rows], colWidths=[7 * cm, 11 * cm])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{BV_LIGHT_GRAY}")),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
    ]))
    story.extend([table, Spacer(1, 8)])

    story.append(Paragraph("4. Conclusão e Recomendação", styles["BVHeading"]))
    story.append(Paragraph(clean(narrative.get("conclusao_recomendacao", "")), normal))
    story.append(Paragraph("5. Observações", styles["BVHeading"]))
    for obs in [
        "A LLM/Data+RAG não calcula os números: ela apenas transforma os indicadores calculados pelo motor estatístico em texto executivo.",
        f"A margem de segurança usada foi de {stats.threshold_pct:.0f}% da capacidade total.",
        stats.confidence_note,
    ]:
        story.append(Paragraph("• " + clean(obs), normal))

    doc.build(story, onFirstPage=footer, onLaterPages=footer)
    return pdf_path


def parse_formats(formats: str | Iterable[str]) -> List[str]:
    if isinstance(formats, str):
        raw = re.split(r"[,;\s]+", formats.strip())
    else:
        raw = list(formats)
    mapping = {"markdown": "md", "md": "md", "word": "docx", "docx": "docx", "pdf": "pdf"}
    result: List[str] = []
    for item in raw:
        key = str(item or "").strip().lower()
        if not key:
            continue
        if key not in mapping:
            raise ValueError(f"Formato inválido: {item}. Use md, docx ou pdf.")
        fmt = mapping[key]
        if fmt not in result:
            result.append(fmt)
    return result or ["md"]


def export_reports(req: ResourceAnalysisRequest, stats: ResourceStats, narrative: Dict[str, str], chart_paths: Dict[str, str | Path], out_dir: str | Path, formats: str | Iterable[str] = "md") -> Dict[str, Path]:
    """Gera documentos adicionais. Markdown é gerado pelo report_builder; aqui cuidamos de TXT/DOCX/PDF."""
    selected = parse_formats(formats)
    outputs: Dict[str, Path] = {}
    outputs["txt"] = build_plain_text_report(req, stats, narrative, out_dir)
    if "docx" in selected:
        outputs["docx"] = build_docx_report(req, stats, narrative, chart_paths, out_dir)
    if "pdf" in selected:
        outputs["pdf"] = build_pdf_report(req, stats, narrative, chart_paths, out_dir)
    return outputs


def make_zip_bundle(zip_path: str | Path, files: Iterable[str | Path], base_dir: str | Path | None = None) -> Path:
    zip_path = Path(zip_path)
    zip_path.parent.mkdir(parents=True, exist_ok=True)
    base = Path(base_dir) if base_dir else zip_path.parent
    seen = set()
    with ZipFile(zip_path, "w", ZIP_DEFLATED) as zf:
        for file in files:
            path = Path(file)
            if not path.exists() or path.is_dir():
                continue
            resolved = path.resolve()
            if resolved in seen:
                continue
            seen.add(resolved)
            try:
                arcname = path.relative_to(base)
            except Exception:
                arcname = path.name
            zf.write(path, arcname=str(arcname).replace("\\", "/"))
    return zip_path
